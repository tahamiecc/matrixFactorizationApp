"""
Streamlit Web Uygulaması
Matris Faktörizasyon Algoritmaları - Kapsamlı Uygulama
"""

import streamlit as st
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
from sklearn.metrics import mean_squared_error

# Klasik Algoritmalar
from algorithms.svd import SVDRecommender, SVDNoiseReducer
from algorithms.pca import PCAAnalyzer
from algorithms.nmf import NMFImageProcessor, NMFTopicModeler
from algorithms.als import ALSRecommender

# Modern Deep Learning Algoritmalar - PyTorch tabanlı (Python 3.14 uyumlu ✅)
MODERN_AVAILABLE = False
try:
    import torch
    from algorithms.ncf import NCFRecommender
    from algorithms.autoencoder import DenoisingAutoencoder, VAERecommender
    from algorithms.fm import FactorizationMachine, DeepFM
    from algorithms.transformer import TransformerRecommender
    from algorithms.gnn import GNNRecommender
    MODERN_AVAILABLE = True
except ImportError as e:
    st.warning(f"⚠️ Modern algoritmalar için PyTorch gerekli: {e}")
    st.info("Yüklemek için: `pip install torch torch-geometric`")

# Yardımcı fonksiyonlar
from utils.data_loader import (
    generate_sample_data, 
    generate_rating_matrix,
    load_sample_images,
    generate_text_corpus,
    generate_noisy_data,
    load_rating_data_from_file,
    load_rating_matrix_from_file
)
from utils.visualization import (
    plot_ratings_matrix,
    plot_recommendations,
    plot_image_grid,
    plot_topic_words
)

# Sayfa yapılandırması
st.set_page_config(
    page_title="Matris Faktörizasyon Algoritmaları",
    page_icon="🔢",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS
st.markdown("""
<style>
    .main-header {
        font-size: 3rem;
        font-weight: bold;
        text-align: center;
        color: #1f77b4;
        margin-bottom: 2rem;
    }
    .algorithm-card {
        background-color: #f0f2f6;
        padding: 1rem;
        border-radius: 0.5rem;
        margin: 1rem 0;
    }
    .metric-card {
        background-color: #e8f4f8;
        padding: 1rem;
        border-radius: 0.5rem;
        text-align: center;
    }
</style>
""", unsafe_allow_html=True)


def main():
    """Ana uygulama"""
    st.markdown('<h1 class="main-header">🔢 Matris Faktörizasyon Algoritmaları</h1>', 
                unsafe_allow_html=True)
    
    # Sidebar
    st.sidebar.title("📋 Menü")
    menu_items = [
        "🏠 Ana Sayfa",
        "--- Klasik Algoritmalar ---",
        "📊 SVD - Öneri Sistemi",
        "🔇 SVD - Gürültü Temizleme",
        "📈 PCA - Veri Görselleştirme",
        "🖼️ NMF - Görüntü İşleme",
        "📝 NMF - Topic Modeling",
        "⚡ ALS - Öneri Sistemi",
        "--- Modern Algoritmalar ---",
        "🧠 NCF - Neural Collaborative Filtering",
        "🎨 Autoencoder - Gürültü Temizleme",
        "🎯 VAE - Variational Autoencoder",
        "🔗 Factorization Machines",
        "🚀 DeepFM",
        "🔄 Transformer - Sequential Recommendation",
        "🕸️ GNN - Graph Neural Network",
        "--- Karşılaştırma ---",
        "📊 Performans Karşılaştırması",
        "--- AI Asistanı ---",
        "🤖 AI Chat - Veri Asistanı"
    ]
    
    page = st.sidebar.selectbox("Sayfa Seçin", menu_items)
    
    if page == "🏠 Ana Sayfa":
        show_homepage()
    elif page == "📊 SVD - Öneri Sistemi":
        show_svd_recommender()
    elif page == "🔇 SVD - Gürültü Temizleme":
        show_svd_noise_reduction()
    elif page == "📈 PCA - Veri Görselleştirme":
        show_pca_visualization()
    elif page == "🖼️ NMF - Görüntü İşleme":
        show_nmf_image_processing()
    elif page == "📝 NMF - Topic Modeling":
        show_nmf_topic_modeling()
    elif page == "⚡ ALS - Öneri Sistemi":
        show_als_recommender()
    elif page == "🧠 NCF - Neural Collaborative Filtering" and MODERN_AVAILABLE:
        show_ncf_recommender()
    elif page == "🎨 Autoencoder - Gürültü Temizleme" and MODERN_AVAILABLE:
        show_autoencoder_denoising()
    elif page == "🎯 VAE - Variational Autoencoder" and MODERN_AVAILABLE:
        show_vae_recommender()
    elif page == "🔗 Factorization Machines" and MODERN_AVAILABLE:
        show_fm_recommender()
    elif page == "🚀 DeepFM" and MODERN_AVAILABLE:
        show_deepfm_recommender()
    elif page == "🔄 Transformer - Sequential Recommendation" and MODERN_AVAILABLE:
        show_transformer_recommender()
    elif page == "🕸️ GNN - Graph Neural Network" and MODERN_AVAILABLE:
        show_gnn_recommender()
    elif page == "📊 Performans Karşılaştırması":
        show_performance_comparison()
    elif page == "🤖 AI Chat - Veri Asistanı":
        show_ai_chatbot()
    elif page.startswith("---"):
        show_homepage()
    else:
        if not MODERN_AVAILABLE:
            st.error("⚠️ Modern algoritmalar için PyTorch gerekli!")
            st.info("💡 Lütfen şu komutu çalıştırın: `pip install torch torch-geometric`")
        else:
            show_homepage()


def show_homepage():
    """Ana sayfa"""
    st.markdown("""
    ## Hoş Geldiniz! 👋
    
    Bu uygulama, matris faktörizasyon algoritmalarının detaylı kullanımını gösteren 
    kapsamlı bir Python uygulamasıdır.
    
    ### 📚 Algoritmalar
    
    #### 🎯 Klasik Algoritmalar
    
    **1. SVD (Singular Value Decomposition)**
    - Matrisi tekil değerlerine ayırır. Matematiksel olarak en kesin yöntemdir.
    - ✅ Öneri sistemleri, Gürültü temizleme
    
    **2. PCA (Principal Component Analysis)**
    - Veriyi daha düşük boyutlu bir uzaya izdüşürür.
    - ✅ Veri görselleştirme, Özellik seçimi
    
    **3. NMF (Non-negative Matrix Factorization)**
    - Matrisleri sadece pozitif değerlerle ayırır.
    - ✅ Görüntü işleme, Metin madenciliği (Topic Modeling)
    
    **4. ALS (Alternating Least Squares)**
    - Büyük ölçekli sistemlerde paralel çalışmaya çok uygundur.
    - ✅ Büyük ölçekli öneri motorları
    
    #### 🚀 Modern Deep Learning Algoritmalar
    
    **5. NCF (Neural Collaborative Filtering)**
    - SVD'nin Deep Learning versiyonu. Doğrusal olmayan ilişkileri öğrenir.
    - ✅ Netflix, YouTube gibi modern öneri sistemleri
    
    **6. Autoencoder (Denoising & VAE)**
    - SVD ve PCA'in Deep Learning karşılığı.
    - ✅ Gürültü temizleme, Variational Autoencoder ile öneriler
    
    **7. Factorization Machines (FM) & DeepFM**
    - Context-aware öneri sistemi. Yan bilgileri kullanır.
    - ✅ Reklam tıklama tahmini (CTR), Context-aware öneriler
    
    **8. Transformer (BERT4Rec/SASRec)**
    - ChatGPT mimarisinin öneri sistemlerine uyarlanmış hali.
    - ✅ TikTok, YouTube gibi sequential öneriler
    
    **9. GNN (Graph Neural Network)**
    - Veriyi tablo değil, ağ (graph) olarak görür.
    - ✅ Pinterest, Uber Eats, Sosyal ağ tabanlı öneriler
    
    ### 🚀 Kullanım
    
    Sol menüden istediğiniz algoritmayı seçerek başlayabilirsiniz!
    
    **Not**: Modern algoritmalar için TensorFlow/PyTorch gerekebilir.
    """)
    
    # Hızlı istatistikler
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Toplam Algoritma", "9" if MODERN_AVAILABLE else "4")
    with col2:
        st.metric("Klasik Algoritma", "4")
    with col3:
        st.metric("Modern Algoritma", "5" if MODERN_AVAILABLE else "0")
    with col4:
        st.metric("Kullanım Örneği", "8+")
    
    if not MODERN_AVAILABLE:
        st.warning("⚠️ Modern algoritmalar için PyTorch gerekli!")
        st.info("💡 Modern algoritmaları kullanmak için: `pip install torch torch-geometric`")


def get_optimal_model_params(model_name, data_shape=None, n_samples=None, n_features=None):
    """
    Veri boyutuna göre optimal model parametrelerini döndürür
    
    Args:
        model_name: Model adı
        data_shape: (n_users, n_items) veya (n_samples, n_features) tuple
        n_samples: Örnek sayısı
        n_features: Özellik sayısı
        
    Returns:
        Parametreler dictionary'si
    """
    params = {}
    
    if data_shape:
        n_samples = data_shape[0]
        n_features = data_shape[1] if len(data_shape) > 1 else None
    
    if model_name == "svd":
        # SVD: n_components genelde min(n_users, n_items) / 3 ile min(n_users, n_items) / 2 arası
        if data_shape and len(data_shape) == 2:
            n_users, n_items = data_shape
            max_comp = min(n_users, n_items)
            params['n_components'] = max(10, min(30, max_comp // 3))
        else:
            params['n_components'] = 20
    
    elif model_name == "als":
        # ALS: n_factors genelde 20-50 arası, regularization 0.1, iterations 15-20
        if data_shape and len(data_shape) == 2:
            n_users, n_items = data_shape
            max_factors = min(n_users, n_items)
            params['n_factors'] = max(10, min(30, max_factors // 5))
        else:
            params['n_factors'] = 20
        params['regularization'] = 0.1
        params['iterations'] = 15
    
    elif model_name == "pca":
        # PCA: n_components genelde min(n_features, n_samples-1) / 2
        if n_features:
            params['n_components'] = min(20, max(5, n_features // 2))
        else:
            params['n_components'] = 20
    
    elif model_name == "nmf_image":
        # NMF Image: n_components genelde 10-30 arası
        params['n_components'] = 20
    
    elif model_name == "nmf_topic":
        # NMF Topic: n_topics genelde 5-10, max_features 500-1000
        params['n_topics'] = 5
        params['max_features'] = 1000
    
    elif model_name == "autoencoder":
        # AutoEncoder: encoding_dim genelde n_features / 2 ile n_features / 4 arası
        if n_features:
            params['encoding_dim'] = max(10, min(50, n_features // 2))
        else:
            params['encoding_dim'] = 20
        params['epochs'] = 50
        params['noise_factor'] = 0.2
    
    elif model_name == "vae":
        # VAE: latent_dim genelde 50-100, epochs 30-50
        params['latent_dim'] = 50
        params['epochs'] = 30
    
    elif model_name == "ncf":
        # NCF: n_factors 32-64, epochs 10-20, batch_size 128-256
        params['n_factors'] = 50
        params['epochs'] = 10
        params['batch_size'] = 256
        params['dropout'] = 0.2
    
    elif model_name == "transformer":
        # Transformer: d_model 64-128, n_heads 4-8, max_seq_length 50, epochs 10
        params['d_model'] = 128
        params['n_heads'] = 4
        params['max_seq_length'] = 50
        params['epochs'] = 10
    
    elif model_name == "gnn":
        # GNN: embedding_dim 32-128, epochs 30-50
        params['embedding_dim'] = 64
        params['epochs'] = 50
    
    elif model_name == "fm" or model_name == "deepfm":
        # FM/DeepFM: embedding_dim 32-64, epochs 30-50
        params['embedding_dim'] = 64
        params['epochs'] = 30
    
    return params


def get_file_format_selector(model_name="model", include_image=False):
    """
    Dosya formatı seçici widget'ı döndürür
    
    Args:
        model_name: Model adı (key için)
        include_image: Görüntü formatı seçeneğini dahil et (varsayılan: False)
        
    Returns:
        Seçilen format string'i
    """
    format_options = [
        "📊 Excel Formatı (Long Format: user_id, item_id, rating)",
        "📋 Matris Formatı (Rating Matrisi: Satırlar=kullanıcı, Sütunlar=ürün)",
        "📁 Her İkisi (Otomatik Tespit)",
        "🖼️ Fotoğraf/Görüntü (Sadece görüntü işleme modelleri için)"
    ]
    
    # Görüntü formatı dahil edilecekse tüm seçenekleri göster
    if include_image:
        options_to_show = format_options
    else:
        options_to_show = format_options[:3]  # Fotoğraf sadece görüntü modellerinde
    
    selected_format = st.radio(
        "📄 Dosya Formatı Seçin",
        options_to_show,
        key=f"{model_name}_format",
        help="Yükleyeceğiniz dosyanın formatını seçin"
    )
    
    return selected_format


def show_svd_recommender():
    """SVD öneri sistemi"""
    st.header("📊 SVD - Öneri Sistemi")
    
    # Info bölümü
    with st.expander("ℹ️ SVD Öneri Sistemi Hakkında Bilgi", expanded=False):
        st.markdown("""
        ### SVD (Singular Value Decomposition) Nedir?
        
        **SVD**, bir matrisi üç matrisin çarpımına ayıran matematiksel bir yöntemdir:
        - **U**: Sol tekil vektörler matrisi (kullanıcı özellikleri)
        - **Σ**: Tekil değerler matrisi (köşegen matris)
        - **V^T**: Sağ tekil vektörler matrisi (ürün özellikleri)
        
        ### Öneri Sistemlerinde Kullanımı
        
        1. **Rating Matrisi Faktörizasyonu**: Kullanıcı-ürün rating matrisini düşük rank matrislere ayırır
        2. **Latent Faktörler**: Gizli kullanıcı ve ürün özelliklerini keşfeder
        3. **Eksik Rating Tahmini**: Kullanıcıların henüz değerlendirmediği ürünler için rating tahmin eder
        
        ### Parametreler
        
        - **Tekil Değer Sayısı**: Kullanılacak latent faktör sayısı. Daha fazla = daha detaylı ama daha yavaş
        - **Eksik Veri Oranı**: Rating matrisindeki boş hücre oranı (sparsity)
        
        ### Metrikler
        
        - **RMSE (Root Mean Square Error)**: Tahmin hatasının ölçüsü. Düşük değer = daha iyi performans
        - **Tekil Değerler**: Her bileşenin önemini gösterir. Büyük değerler = daha önemli bileşenler
        """)
    
    st.markdown("""
    SVD (Singular Value Decomposition) kullanarak öneri sistemi oluşturma.
    """)
    
    # Veri yükleme seçeneği
    data_source = st.radio(
        "Veri Kaynağı",
        ["📊 Örnek Veri Oluştur", "📁 Dosyadan Yükle"],
        horizontal=True
    )
    
    # Dosya formatı seçimi (sadece dosya yükleme seçildiyse)
    file_format = None
    if data_source == "📁 Dosyadan Yükle":
        file_format = get_file_format_selector("svd", include_image=False)
    
    # Session state ile veriyi koru
    if 'svd_rating_matrix' not in st.session_state:
        st.session_state.svd_rating_matrix = None
        st.session_state.svd_user_mapping = None
        st.session_state.svd_item_mapping = None
    
    rating_matrix = st.session_state.svd_rating_matrix
    user_mapping = st.session_state.svd_user_mapping
    item_mapping = st.session_state.svd_item_mapping
    
    # Varsayılan değerler
    n_users = None
    n_items = None
    n_components = 30
    sparsity = 0.6  # Varsayılan değer
    
    if data_source == "📁 Dosyadan Yükle":
        st.markdown("### 📁 Veri Dosyası Yükle")
        st.info("""
        **Desteklenen Formatlar:** CSV, Excel (.xlsx, .xls)
        
        **Veri Formatı Seçenekleri:**
        1. **Long Format** (Önerilen): Her satır bir kullanıcı-ürün-rating üçlüsü
           - Sütunlar: `user_id`, `item_id`, `rating`
           - Örnek: `user_id,item_id,rating` → `1,5,4.5`
        
        2. **Matrix Format**: Zaten rating matrisi formatında
           - İlk sütun: Kullanıcı ID'leri (index)
           - Diğer sütunlar: Ürün ID'leri
           - Değerler: Rating'ler (NaN = eksik veri)
        """)
        
        file = st.file_uploader(
            "Veri dosyasını seçin",
            type=['csv', 'xlsx', 'xls'],
            help="CSV veya Excel dosyası yükleyin"
        )
        
        if file is not None:
            try:
                # Dosya önizlemesi ve format önerisi
                # Dosya stream'i bir kez okununca tükenir, bu yüzden içeriği hafızaya al
                import io
                file_content = file.read()
                file_bytes = io.BytesIO(file_content)
                
                import pandas as pd
                if file.name.endswith('.csv'):
                    # Delimiter tespiti
                    file_bytes.seek(0)
                    first_line = file_bytes.readline().decode('utf-8', errors='ignore')
                    delimiters = [',', ';', '\t', '|']
                    detected_delimiter = ','
                    max_cols = 0
                    for delim in delimiters:
                        cols = first_line.split(delim)
                        if len(cols) > max_cols:
                            max_cols = len(cols)
                            detected_delimiter = delim
                    
                    file_bytes.seek(0)
                    preview_df = pd.read_csv(file_bytes, nrows=5, sep=detected_delimiter, engine='python')
                    # Toplam satır sayısı için dosyayı tekrar oku
                    file_bytes.seek(0)
                    total_df = pd.read_csv(file_bytes, sep=detected_delimiter, engine='python')
                elif file.name.endswith(('.xlsx', '.xls')):
                    file_bytes.seek(0)
                    preview_df = pd.read_excel(file_bytes, nrows=5)
                    # Toplam satır sayısı için dosyayı tekrar oku
                    file_bytes.seek(0)
                    total_df = pd.read_excel(file_bytes)
                else:
                    preview_df = None
                    total_df = None
                
                if preview_df is not None:
                    with st.expander("👁️ Dosya Önizleme (İlk 5 Satır)", expanded=False):
                        st.dataframe(preview_df, width='stretch')
                        st.info(f"""
                        **Dosya Bilgileri:**
                        - **Satır Sayısı**: {len(total_df) if total_df is not None else 'Bilinmiyor'} (tahmini)
                        - **Sütun Sayısı**: {len(preview_df.columns)}
                        - **Sütun İsimleri**: {', '.join(preview_df.columns.tolist()[:10])}{'...' if len(preview_df.columns) > 10 else ''}
                        
                        **Format Önerisi:**
                        - **3 sütun varsa** → Long Format seçin (user_id, item_id, rating)
                        - **10+ sütun varsa** → Matrix Format seçin (ilk sütun kullanıcı ID, diğerleri ürün ID)
                        """)
                
                # Dosya stream'ini tekrar kullanılabilir hale getir
                file.seek(0)
                
                # Dosya formatına göre veri formatını belirle
                if file_format and file_format.startswith("📊"):
                    data_format = "Long Format (user_id, item_id, rating)"
                elif file_format and file_format.startswith("📋"):
                    data_format = "Matrix Format (Rating Matrisi)"
                elif file_format and file_format.startswith("📁"):
                    # Otomatik tespit - sütun sayısına göre
                    if len(preview_df.columns) == 3:
                        data_format = "Long Format (user_id, item_id, rating)"
                    elif len(preview_df.columns) > 3:
                        data_format = "Matrix Format (Rating Matrisi)"
                    else:
                        data_format = "Long Format (user_id, item_id, rating)"  # Varsayılan
                else:
                    # Eski yöntem (geriye dönük uyumluluk)
                    data_format = st.radio(
                        "Veri Formatı",
                        ["Long Format (user_id, item_id, rating)", "Matrix Format (Rating Matrisi)"],
                        horizontal=True
                    )
                
                if data_format == "Long Format (user_id, item_id, rating)":
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        user_col = st.text_input("Kullanıcı Sütunu", value="", 
                                                help="Kullanıcı ID sütunu adı (boş bırakırsanız otomatik tespit edilir)")
                    with col2:
                        item_col = st.text_input("Ürün Sütunu", value="",
                                                help="Ürün ID sütunu adı (boş bırakırsanız otomatik tespit edilir)")
                    with col3:
                        rating_col = st.text_input("Rating Sütunu", value="",
                                                  help="Rating sütunu adı (boş bırakırsanız otomatik tespit edilir)")
                    
                    if st.button("📥 Veriyi Yükle"):
                        with st.spinner("Veri yükleniyor..."):
                            # Dosya bilgilerini kaydet
                            file_name = file.name
                            file_size = file.size
                            
                            rating_matrix, user_mapping, item_mapping = load_rating_data_from_file(
                                file,
                                user_col=user_col if user_col else None,
                                item_col=item_col if item_col else None,
                                rating_col=rating_col if rating_col else None
                            )
                            
                            # Dosya bilgilerini session state'e kaydet
                            st.session_state.svd_file_name = file_name
                            st.session_state.svd_file_size = file_size
                            st.session_state.svd_user_col = user_col if user_col else "Otomatik tespit edildi"
                            st.session_state.svd_item_col = item_col if item_col else "Otomatik tespit edildi"
                            st.session_state.svd_rating_col = rating_col if rating_col else "Otomatik tespit edildi"
                            
                            # Session state'e kaydet
                            st.session_state.svd_rating_matrix = rating_matrix
                            st.session_state.svd_user_mapping = user_mapping
                            st.session_state.svd_item_mapping = item_mapping
                            
                            # Veri istatistiklerini hesapla
                            from scipy.sparse import issparse
                            if issparse(rating_matrix):
                                n_ratings = rating_matrix.nnz
                                density = rating_matrix.nnz / (rating_matrix.shape[0] * rating_matrix.shape[1]) * 100 if rating_matrix.shape[0] * rating_matrix.shape[1] > 0 else 0
                                ratings_data = rating_matrix.data
                            else:
                                mask = ~np.isnan(rating_matrix)
                                n_ratings = np.sum(mask)
                                density = (1 - np.isnan(rating_matrix).sum() / rating_matrix.size) * 100 if rating_matrix.size > 0 else 0
                                ratings_data = rating_matrix[mask]
                            
                            # Rating istatistikleri - boş array kontrolü
                            if len(ratings_data) == 0:
                                # Eğer hiç rating yoksa varsayılan değerler
                                min_rating = 0.0
                                max_rating = 0.0
                                mean_rating = 0.0
                                median_rating = 0.0
                                st.warning("⚠️ Uyarı: Dosyada hiç rating değeri bulunamadı! Tüm değerler NaN olabilir.")
                            else:
                                min_rating = float(np.min(ratings_data))
                                max_rating = float(np.max(ratings_data))
                                mean_rating = float(np.mean(ratings_data))
                                median_rating = float(np.median(ratings_data))
                            
                            st.success(f"✅ Veri yüklendi! {rating_matrix.shape[0]} kullanıcı, {rating_matrix.shape[1]} ürün")
                            
                            # Detaylı dosya analizi bölümü
                            with st.expander("📋 Dosya Analizi - Kullanılan Veriler", expanded=True):
                                st.markdown("### 📁 Dosya Bilgileri")
                                col1, col2 = st.columns(2)
                                with col1:
                                    st.metric("Dosya Adı", file_name)
                                    st.metric("Dosya Boyutu", f"{file_size / 1024:.2f} KB")
                                with col2:
                                    st.metric("Veri Formatı", "Long Format")
                                    st.metric("Toplam Satır Sayısı", f"{n_ratings:,}")
                                
                                st.markdown("### 📊 Kullanılan Sütunlar")
                                col1, col2, col3 = st.columns(3)
                                with col1:
                                    st.info(f"**Kullanıcı Sütunu:**\n{st.session_state.svd_user_col}")
                                with col2:
                                    st.info(f"**Ürün Sütunu:**\n{st.session_state.svd_item_col}")
                                with col3:
                                    st.info(f"**Rating Sütunu:**\n{st.session_state.svd_rating_col}")
                                
                                st.markdown("### 📈 Veri İstatistikleri")
                                col1, col2, col3, col4 = st.columns(4)
                                with col1:
                                    st.metric("Kullanıcı Sayısı", f"{rating_matrix.shape[0]:,}")
                                with col2:
                                    st.metric("Ürün Sayısı", f"{rating_matrix.shape[1]:,}")
                                with col3:
                                    st.metric("Toplam Rating", f"{n_ratings:,}")
                                with col4:
                                    st.metric("Veri Yoğunluğu", f"{density:.2f}%")
                                
                                st.markdown("### ⭐ Rating Dağılımı")
                                col1, col2, col3, col4 = st.columns(4)
                                with col1:
                                    st.metric("Minimum Rating", f"{min_rating:.2f}")
                                with col2:
                                    st.metric("Maksimum Rating", f"{max_rating:.2f}")
                                with col3:
                                    st.metric("Ortalama Rating", f"{mean_rating:.2f}")
                                with col4:
                                    st.metric("Medyan Rating", f"{median_rating:.2f}")
                                
                                st.markdown("""
                                **📝 Açıklama:**
                                - **Dosya Bilgileri**: Yüklenen dosyanın adı ve boyutu
                                - **Kullanılan Sütunlar**: Veri işlemede kullanılan sütun isimleri
                                - **Veri İstatistikleri**: Matris boyutları ve veri yoğunluğu
                                - **Rating Dağılımı**: Rating değerlerinin istatistiksel özeti
                                
                                Bu veriler, SVD (Singular Value Decomposition) algoritması ile işlenecek ve 
                                kullanıcı-ürün etkileşimlerinden latent faktörler çıkarılacaktır.
                                """)
                            
                            st.rerun()  # Sayfayı yenile
                else:
                    if st.button("📥 Veriyi Yükle"):
                        with st.spinner("Veri yükleniyor..."):
                            # Dosya bilgilerini kaydet
                            file_name = file.name
                            file_size = file.size
                            
                            try:
                                rating_matrix = load_rating_matrix_from_file(file)
                            except Exception as e:
                                st.error(f"❌ Hata: {str(e)}")
                                st.info("""
                                **💡 Matrix Format için:**
                                - İlk sütun kullanıcı ID'leri olmalı (index)
                                - Diğer sütunlar ürün ID'leri olmalı
                                - Değerler rating'ler olmalı (NaN = eksik veri)
                                - CSV dosyasında ilk sütun otomatik olarak index olarak okunur
                                
                                **Örnek Format:**
                                ```
                                user_id,item_1,item_2,item_3,...
                                1,4.5,3.0,5.0,...
                                2,2.5,4.0,NaN,...
                                ```
                                """)
                                st.stop()
                            
                            # Dosya bilgilerini session state'e kaydet
                            st.session_state.svd_file_name = file_name
                            st.session_state.svd_file_size = file_size
                            
                            # Session state'e kaydet
                            st.session_state.svd_rating_matrix = rating_matrix
                            st.session_state.svd_user_mapping = None
                            st.session_state.svd_item_mapping = None
                            
                            # Veri istatistiklerini hesapla
                            from scipy.sparse import issparse
                            if issparse(rating_matrix):
                                n_ratings = rating_matrix.nnz
                                density = rating_matrix.nnz / (rating_matrix.shape[0] * rating_matrix.shape[1]) * 100 if rating_matrix.shape[0] * rating_matrix.shape[1] > 0 else 0
                                ratings_data = rating_matrix.data
                            else:
                                mask = ~np.isnan(rating_matrix)
                                n_ratings = np.sum(mask)
                                density = (1 - np.isnan(rating_matrix).sum() / rating_matrix.size) * 100 if rating_matrix.size > 0 else 0
                                ratings_data = rating_matrix[mask]
                            
                            # Rating istatistikleri - boş array kontrolü
                            if len(ratings_data) == 0:
                                # Eğer hiç rating yoksa varsayılan değerler
                                min_rating = 0.0
                                max_rating = 0.0
                                mean_rating = 0.0
                                median_rating = 0.0
                                st.warning("⚠️ Uyarı: Dosyada hiç rating değeri bulunamadı! Tüm değerler NaN olabilir.")
                            else:
                                min_rating = float(np.min(ratings_data))
                                max_rating = float(np.max(ratings_data))
                                mean_rating = float(np.mean(ratings_data))
                                median_rating = float(np.median(ratings_data))
                            
                            st.success(f"✅ Veri yüklendi! {rating_matrix.shape[0]} kullanıcı, {rating_matrix.shape[1]} ürün")
                            
                            # Detaylı dosya analizi bölümü
                            with st.expander("📋 Dosya Analizi - Kullanılan Veriler", expanded=True):
                                st.markdown("### 📁 Dosya Bilgileri")
                                col1, col2 = st.columns(2)
                                with col1:
                                    st.metric("Dosya Adı", file_name)
                                    st.metric("Dosya Boyutu", f"{file_size / 1024:.2f} KB")
                                with col2:
                                    st.metric("Veri Formatı", "Matrix Format")
                                    st.metric("Toplam Rating", f"{n_ratings:,}")
                                
                                st.markdown("### 📈 Veri İstatistikleri")
                                col1, col2, col3, col4 = st.columns(4)
                                with col1:
                                    st.metric("Kullanıcı Sayısı", f"{rating_matrix.shape[0]:,}")
                                with col2:
                                    st.metric("Ürün Sayısı", f"{rating_matrix.shape[1]:,}")
                                with col3:
                                    st.metric("Toplam Rating", f"{n_ratings:,}")
                                with col4:
                                    st.metric("Veri Yoğunluğu", f"{density:.2f}%")
                                
                                st.markdown("### ⭐ Rating Dağılımı")
                                col1, col2, col3, col4 = st.columns(4)
                                with col1:
                                    st.metric("Minimum Rating", f"{min_rating:.2f}")
                                with col2:
                                    st.metric("Maksimum Rating", f"{max_rating:.2f}")
                                with col3:
                                    st.metric("Ortalama Rating", f"{mean_rating:.2f}")
                                with col4:
                                    st.metric("Medyan Rating", f"{median_rating:.2f}")
                                
                                st.markdown("""
                                **📝 Açıklama:**
                                - **Dosya Bilgileri**: Yüklenen dosyanın adı ve boyutu
                                - **Veri İstatistikleri**: Matris boyutları ve veri yoğunluğu
                                - **Rating Dağılımı**: Rating değerlerinin istatistiksel özeti
                                
                                Bu veriler, SVD (Singular Value Decomposition) algoritması ile işlenecek ve 
                                kullanıcı-ürün etkileşimlerinden latent faktörler çıkarılacaktır.
                                """)
                            
                            st.rerun()  # Sayfayı yenile
                            
            except Exception as e:
                st.error(f"❌ Hata: {str(e)}")
                st.info("💡 Lütfen veri formatını kontrol edin. Örnek format için yukarıdaki bilgi kutusuna bakın.")
    else:
        # Örnek veri oluştur - varsayılan değerleri güncelle
        col1, col2, col3 = st.columns(3)
        with col1:
            n_users = st.slider("Kullanıcı Sayısı", 50, 500, 100)
        with col2:
            n_items = st.slider("Ürün Sayısı", 30, 200, 50)
        with col3:
            n_components = st.slider("Tekil Değer Sayısı", 5, 50, 30, 
                                    help="Daha fazla bileşen = daha çeşitli tahminler ama daha yavaş")
        
        sparsity = st.slider("Eksik Veri Oranı", 0.3, 0.9, 0.6,
                            help="Daha az eksik veri = daha iyi öğrenme")
    
    # Session state'ten güncel değerleri al
    rating_matrix = st.session_state.svd_rating_matrix
    
    # Model parametreleri (veri yüklendiyse)
    if rating_matrix is not None:
        n_users, n_items = rating_matrix.shape
        
        # Veri seti boş mu kontrol et
        if n_users == 0 or n_items == 0:
            st.error("❌ Hata: Veri seti boş! Lütfen geçerli bir veri dosyası yükleyin.")
            n_components = 5  # Varsayılan değer
        else:
            # Optimal parametreleri al
            optimal_params = get_optimal_model_params("svd", data_shape=(n_users, n_items))
            optimal_n_components = optimal_params['n_components']
            
            max_components = min(50, min(n_users, n_items))
            min_components = min(5, max_components)  # min_value max_value'dan küçük olmalı
            default_components = max(min_components, min(max_components, optimal_n_components))
            
            # Eğer max_components çok küçükse veya 0 ise, slider yerine sabit değer kullan
            if max_components <= 0 or max_components < min_components:
                n_components = max(1, max_components)  # En az 1 bileşen
                if max_components <= 0:
                    st.error("❌ Hata: Veri seti çok küçük! Bileşen sayısı ayarlanamıyor.")
                else:
                    st.info(f"⚠️ Veri seti küçük olduğu için bileşen sayısı otomatik olarak {n_components} olarak ayarlandı.")
            else:
                n_components = st.slider(
                    "Tekil Değer Sayısı", 
                    min_components, 
                    max_components, 
                    default_components,
                    help=f"Önerilen değer: {optimal_n_components} (veri boyutuna göre otomatik hesaplandı). Daha fazla bileşen = daha çeşitli tahminler ama daha yavaş"
                )
                if n_components != optimal_n_components:
                    st.info(f"💡 Veri boyutunuza göre önerilen değer: {optimal_n_components}")
    
    if st.button("🚀 Modeli Eğit"):
        if rating_matrix is None and data_source == "📁 Dosyadan Yükle":
            st.warning("⚠️ Lütfen önce veri dosyasını yükleyin!")
        else:
            with st.spinner("Model eğitiliyor..."):
                # Veri yoksa oluştur
                if rating_matrix is None:
                    # sparsity değişkeni sadece örnek veri için tanımlı
                    rating_matrix = generate_rating_matrix(
                        n_users=n_users, 
                        n_items=n_items, 
                        sparsity=sparsity
                    )
                
                # n_users ve n_items'ı güncelle (analiz seçenekleri için)
                n_users, n_items = rating_matrix.shape
            
            # Train-test split (sparse matrix desteği ile)
            from scipy.sparse import issparse
            
            np.random.seed(42)
            
            if issparse(rating_matrix):
                # Sparse matrix için
                rows, cols = rating_matrix.nonzero()
                n_ratings = len(rows)
                test_size = min(int(0.2 * n_ratings), 10000)  # Max 10k test
                test_sample_indices = np.random.choice(n_ratings, size=test_size, replace=False)
                
                test_rows = rows[test_sample_indices]
                test_cols = cols[test_sample_indices]
                # Sparse matrix'ten değerleri al - matrix objesi için np.array kullan
                test_matrix_slice = rating_matrix[test_rows, test_cols]
                # matrix objesi için A property veya np.array kullan
                if hasattr(test_matrix_slice, 'A'):
                    test_values = test_matrix_slice.A.flatten()
                elif hasattr(test_matrix_slice, 'toarray'):
                    test_values = test_matrix_slice.toarray().flatten()
                else:
                    test_values = np.array(test_matrix_slice).flatten()
                
                # Train matrix - test değerlerini çıkar
                train_matrix = rating_matrix.copy()
                train_matrix[test_rows, test_cols] = 0
                train_matrix.eliminate_zeros()
                
                test_indices_tuple = (test_rows, test_cols)
            else:
                # Dense matrix için
                mask = ~np.isnan(rating_matrix)
                n_ratings = np.sum(mask)
                test_size = min(int(0.2 * n_ratings), 10000)  # Max 10k test
                
                valid_indices = np.where(mask)
                test_sample_indices = np.random.choice(
                    len(valid_indices[0]), 
                    size=test_size, 
                    replace=False
                )
                
                test_mask = np.zeros_like(mask, dtype=bool)
                test_mask[valid_indices[0][test_sample_indices], valid_indices[1][test_sample_indices]] = True
                
                train_matrix = rating_matrix.copy()
                train_matrix[test_mask] = np.nan
                
                test_values = rating_matrix[test_mask]
                test_indices_tuple = np.where(test_mask)
            
            # Model eğit (NaN değerleri ortalama ile doldur)
            svd_model = SVDRecommender(n_components=n_components)
            svd_model.fit(train_matrix, fill_na_with_mean=True)
            
            # Değerlendirme (büyük veri setleri için optimize edilmiş)
            # Test değerlerini tahmin et - batch processing ile
            test_predictions = []
            batch_size = 1000  # Her seferde 1000 tahmin
            
            with st.spinner("Test seti üzerinde tahmin yapılıyor..."):
                for i in range(0, len(test_indices_tuple[0]), batch_size):
                    batch_end = min(i + batch_size, len(test_indices_tuple[0]))
                    batch_users = test_indices_tuple[0][i:batch_end]
                    batch_items = test_indices_tuple[1][i:batch_end]
                    
                    batch_preds = [svd_model.predict(u, it) for u, it in zip(batch_users, batch_items)]
                    test_predictions.extend(batch_preds)
            
            test_predictions = np.array(test_predictions)
            
            # RMSE hesapla
            rmse = np.sqrt(mean_squared_error(test_values, test_predictions))
            singular_values = svd_model.get_singular_values()
            
            # Model ve sonuçları session state'e kaydet (analiz seçenekleri için)
            st.session_state.svd_model_trained = True
            st.session_state.svd_model = svd_model
            st.session_state.svd_singular_values = singular_values
            st.session_state.svd_rating_matrix = rating_matrix
            st.session_state.svd_n_users = n_users
            st.session_state.svd_n_items = n_items
            st.session_state.svd_n_components = n_components
            
            # Sonuçlar
            st.subheader("📊 Model Sonuçları")
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Test RMSE", f"{rmse:.4f}")
            with col2:
                st.metric("Kullanılan Bileşen", len(singular_values))
            with col3:
                variance_explained = np.sum(singular_values**2) / np.sum(singular_values**2)
                st.metric("Varyans Açıklama", f"{variance_explained:.2%}")
            
            # Detaylı sonuç açıklaması
            with st.expander("📝 Sonuç Açıklaması - Ne Elde Edildi?", expanded=True):
                st.markdown("### 🔍 Kullanılan Veriler")
                if data_source == "📁 Dosyadan Yükle" and 'svd_file_name' in st.session_state:
                    st.info(f"""
                    **Dosya**: {st.session_state.svd_file_name} ({st.session_state.svd_file_size / 1024:.2f} KB)
                    - **Kullanıcı Sayısı**: {n_users:,}
                    - **Ürün Sayısı**: {n_items:,}
                    - **Toplam Rating**: {n_ratings:,}
                    """)
                else:
                    st.info(f"""
                    **Örnek Veri**:
                    - **Kullanıcı Sayısı**: {n_users:,}
                    - **Ürün Sayısı**: {n_items:,}
                    """)
                
                st.markdown("### ⚙️ Model Parametreleri")
                st.info(f"""
                - **Tekil Değer Sayısı (n_components)**: {n_components}
                - **Eğitim Verisi**: {n_ratings - len(test_values):,} rating
                - **Test Verisi**: {len(test_values):,} rating
                """)
                
                st.markdown("### 📈 Elde Edilen Sonuçlar")
                st.success(f"""
                **Model Başarıyla Eğitildi!**
                
                1. **Test RMSE**: {rmse:.4f}
                   - Bu değer ne kadar düşükse, model o kadar iyi tahmin yapıyor demektir
                   - RMSE, gerçek rating'ler ile tahmin edilen rating'ler arasındaki ortalama hata miktarını gösterir
                   - Örnek: RMSE = {rmse:.4f} → Ortalama {rmse:.2f} puanlık hata var
                
                2. **Kullanılan Bileşen Sayısı**: {len(singular_values)}
                   - Model, veriyi {len(singular_values)} boyutlu latent faktör uzayına indirgedi
                   - Her bileşen, kullanıcı ve ürün özelliklerini temsil eden bir boyuttur
                
                3. **Varyans Açıklama**: {variance_explained:.2%}
                   - Model, verideki bilginin {variance_explained:.1%}'ini koruyor
                   - Yüksek değer = daha az bilgi kaybı
                
                4. **Tekil Değerler**: {len(singular_values)} adet
                   - İlk birkaç tekil değer genellikle en önemli bilgiyi taşır
                   - Düşük tekil değerler genellikle gürültüyü temsil eder
                """)
                
                st.markdown("### 🎯 Ne Yapıldı?")
                st.markdown("""
                **SVD (Singular Value Decomposition) Algoritması** şu adımları izledi:
                
                1. **Veri Hazırlama**: Rating matrisi train ve test setlerine ayrıldı
                2. **Matris Faktörizasyonu**: Rating matrisi üç matrise ayrıldı:
                   - **U**: Kullanıcı latent faktörleri
                   - **Σ**: Tekil değerler (bileşen önemleri)
                   - **V^T**: Ürün latent faktörleri
                3. **Boyut İndirgeme**: Sadece en önemli {n_components} bileşen kullanıldı
                4. **Tahmin**: Eksik rating'ler, latent faktörler kullanılarak tahmin edildi
                5. **Değerlendirme**: Test seti üzerinde RMSE hesaplandı
                
                **Sonuç**: Model, kullanıcı-ürün etkileşimlerinden öğrendiği kalıpları kullanarak 
                yeni rating'leri tahmin edebiliyor. Bu sayede kullanıcılara henüz görmedikleri 
                ürünler için kişiselleştirilmiş öneriler sunulabilir.
                """)
                
                st.markdown("### 💡 Sonraki Adımlar")
                st.info("""
                - **Öneriler**: Aşağıdaki "Kullanıcı Önerileri" bölümünden belirli bir kullanıcı için öneriler görebilirsiniz
                - **Tekil Değerler Analizi**: Grafikten optimal bileşen sayısını belirleyebilirsiniz
                - **Rating Matrisi Görselleştirmesi**: Veri yapısını görsel olarak inceleyebilirsiniz
                - **SVD Analiz Seçenekleri**: Latent matrix, benzerlik analizleri ve tahmini puanlar
                """)
            
            # Tekil değerler grafiği
            st.subheader("📈 Tekil Değerler Analizi")
            with st.expander("ℹ️ Bu Grafik Ne Anlama Geliyor?", expanded=False):
                st.markdown("""
                **Tekil Değerler Grafiği**:
                - Her bileşenin (component) önemini gösterir
                - **Yüksek değerler**: Daha önemli, daha fazla bilgi taşıyan bileşenler
                - **Düşük değerler**: Daha az önemli, gürültü içeren bileşenler
                - Genellikle ilk birkaç bileşen en önemlidir (elbow point)
                - Bu grafik, optimal bileşen sayısını seçmek için kullanılır
                """)
            
            fig, ax = plt.subplots(figsize=(10, 5))
            ax.plot(range(1, len(singular_values) + 1), singular_values, 'o-', linewidth=2, markersize=6)
            ax.set_xlabel('Bileşen Numarası', fontsize=12)
            ax.set_ylabel('Tekil Değer (Singular Value)', fontsize=12)
            ax.set_title('Tekil Değerler - Bileşen Önemi Analizi', fontsize=14, fontweight='bold')
            ax.grid(True, alpha=0.3)
            st.pyplot(fig)
    
    # SVD Analiz Seçenekleri (Model eğitildikten sonra her zaman göster)
    if st.session_state.get('svd_model_trained', False):
        st.subheader("🔍 SVD Analiz Seçenekleri")
        analysis_option = st.radio(
            "Hangi analizi görmek istersiniz?",
            [
                "🎯 Kullanıcı Önerileri",
                "📊 Latent Matrix (Gizli Özellik Skorları)",
                "👥 Kaynak Benzerliği (Kullanıcı Benzerliği)",
                "📚 Konu Benzerliği (Ürün Benzerliği)",
                "🔮 Tahmini Puanlar (Predicted Scores)",
                "📈 Rating Matrisi Görselleştirmesi"
            ],
            key="svd_analysis_option"
        )
        
        # Session state'ten değişkenleri al
        svd_model = st.session_state.svd_model
        singular_values = st.session_state.svd_singular_values
        rating_matrix = st.session_state.svd_rating_matrix
        n_users = st.session_state.svd_n_users
        n_items = st.session_state.svd_n_items
        n_components = st.session_state.svd_n_components
        
        # Seçeneklere göre içerik göster
        try:
                if analysis_option == "🎯 Kullanıcı Önerileri":
                    # Öneriler
                    st.subheader("🎯 Kullanıcı Önerileri")
                    with st.expander("ℹ️ Öneriler Nasıl Oluşturuluyor?", expanded=False):
                        st.markdown("""
                        **Öneri Sistemi Çalışma Prensibi**:
                        1. Model, kullanıcının geçmiş rating'lerini analiz eder
                        2. Latent faktörler kullanarak kullanıcı tercihlerini öğrenir
                        3. Benzer kullanıcıların beğendiği ürünleri bulur
                        4. Kullanıcının henüz görmediği ürünler için rating tahmin eder
                        5. En yüksek tahmin edilen rating'lere sahip ürünleri önerir
                        
                        **Tahmin Edilen Rating**: Modelin, kullanıcının bu ürüne vereceği rating tahmini (1-5 arası)
                        """)
                    
                    user_idx = st.selectbox("Kullanıcı Seçin", range(min(10, n_users)), key="svd_user_select")
                    
                    predictions = svd_model.predict_all()[user_idx]
                    # Sparse matrix desteği
                    from scipy.sparse import issparse
                    if issparse(rating_matrix):
                        # Sparse matrix için - sadece mevcut rating'leri kontrol et
                        user_row = rating_matrix.getrow(user_idx)
                        rated_items = np.zeros(rating_matrix.shape[1], dtype=bool)
                        rated_items[user_row.indices] = True
                    else:
                        rated_items = ~np.isnan(rating_matrix[user_idx])
                    unrated_items = ~rated_items
                    
                    # Debug: Tahmin çeşitliliğini kontrol et
                    if np.sum(unrated_items) > 0:
                        unrated_predictions = predictions[unrated_items]
                        unique_predictions = len(np.unique(np.round(unrated_predictions, 2)))
                        prediction_range = (np.max(unrated_predictions) - np.min(unrated_predictions))
                        prediction_std = np.std(unrated_predictions)
                    else:
                        unique_predictions = len(np.unique(np.round(predictions, 2)))
                        prediction_range = (np.max(predictions) - np.min(predictions))
                        prediction_std = np.std(predictions)
                    
                    if unique_predictions < 3 or prediction_range < 0.5:
                        st.warning(f"⚠️ Uyarı: Tahminlerde çok az çeşitlilik var!")
                        st.info(f"""
                        💡 **İpucular**:
                        - Bileşen sayısını artırın (şu an: {n_components})
                        - Daha fazla kullanıcı/ürün kullanın
                        - Farklı bir kullanıcı seçmeyi deneyin
                        
                        **Mevcut durum**: {unique_predictions} farklı tahmin, aralık: {prediction_range:.2f}
                        """)
                    
                    predictions[rated_items] = -np.inf
                    
                    top_items = np.argsort(predictions)[::-1][:10]
                    top_ratings = predictions[top_items]
                    
                    # Rating'leri 1-5 aralığına sınırla ve yuvarla
                    top_ratings = np.clip(top_ratings, 1, 5)
                    
                    # Öneri tablosu
                    recommendations_df = pd.DataFrame({
                        'Ürün ID': top_items + 1,
                        'Tahmin Edilen Rating': np.round(top_ratings, 2)
                    })
                    st.dataframe(recommendations_df, width='stretch')
                    
                    # İstatistikler
                    with st.expander("📊 Tahmin İstatistikleri", expanded=False):
                        col1, col2, col3 = st.columns(3)
                        with col1:
                            st.metric("Farklı Tahmin Sayısı", unique_predictions)
                        with col2:
                            st.metric("Tahmin Aralığı", f"{prediction_range:.2f}")
                        with col3:
                            st.metric("Tahmin Std Sapma", f"{prediction_std:.2f}")
                
                elif analysis_option == "📊 Latent Matrix (Gizli Özellik Skorları)":
                    st.subheader("📊 Latent Matrix (Gizli Özellik Skorları)")
                    with st.expander("ℹ️ Latent Matrix Nedir?", expanded=False):
                        st.markdown("""
                        **Latent Matrix** (Gizli Özellik Skorları):
                        - Her kaynağın (kullanıcının) gizli özelliklerdeki skorlarını gösterir
                        - `fit_transform()` ile oluşturulan matristir
                        - Her satır bir kullanıcıyı, her sütun bir gizli özelliği temsil eder
                        - Bu matris, kullanıcıların latent space'deki konumunu gösterir
                        - Benzer kullanıcılar benzer skorlara sahip olacaktır
                        """)
                    
                    latent_matrix = svd_model.get_latent_matrix()
                    
                    # Kullanıcı ID'leri
                    user_ids = [f"User_{i+1}" for i in range(latent_matrix.shape[0])]
                    component_ids = [f"Component_{i+1}" for i in range(latent_matrix.shape[1])]
                    
                    latent_df = pd.DataFrame(
                        latent_matrix,
                        index=user_ids,
                        columns=component_ids
                    )
                    
                    st.dataframe(latent_df, width='stretch', height=400)
                    
                    # İstatistikler
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.metric("Kullanıcı Sayısı", latent_matrix.shape[0])
                    with col2:
                        st.metric("Gizli Özellik Sayısı", latent_matrix.shape[1])
                    with col3:
                        st.metric("Toplam Değer", f"{latent_matrix.size:,}")
                    
                    st.info("💡 Bu matris, her kullanıcının gizli özelliklerdeki skorlarını gösterir. Benzer kullanıcılar benzer skorlara sahip olacaktır.")
                
                elif analysis_option == "👥 Kaynak Benzerliği (Kullanıcı Benzerliği)":
                    st.subheader("👥 Kaynak Benzerliği (Kullanıcı Benzerliği)")
                    with st.expander("ℹ️ Kaynak Benzerliği Nedir?", expanded=False):
                        st.markdown("""
                        **Kaynak Benzerliği** (Kullanıcı Benzerliği):
                        - Latent matrix üzerinden kosinüs benzerliği kullanılarak hesaplanır
                        - Hangi kullanıcıların benzer latent profillere sahip olduğunu gösterir
                        - Benzerlik değeri 1'e yakınsa kullanıcılar çok benzerdir
                        - Benzerlik değeri 0'a yakınsa kullanıcılar farklıdır
                        - Bu bilgi, "benzer kullanıcılar benzer ürünleri beğenir" prensibine dayanır
                        """)
                    
                    user_similarity = svd_model.get_user_similarity()
                    
                    # Kullanıcı ID'leri
                    user_ids = [f"User_{i+1}" for i in range(user_similarity.shape[0])]
                    
                    similarity_df = pd.DataFrame(
                        user_similarity,
                        index=user_ids,
                        columns=user_ids
                    )
                    
                    st.dataframe(similarity_df, width='stretch', height=400)
                    
                    # En benzer kullanıcı çiftleri
                    st.markdown("### 🔝 En Benzer Kullanıcı Çiftleri")
                    # Diagonal'i -1 yap (kendisiyle benzerlik hariç)
                    similarity_matrix_copy = user_similarity.copy()
                    np.fill_diagonal(similarity_matrix_copy, -1)
                    
                    # En yüksek benzerlik değerlerini bul
                    n_top = min(10, len(user_ids))
                    top_similarities = []
                    for i in range(len(user_ids)):
                        for j in range(i+1, len(user_ids)):
                            top_similarities.append((i, j, similarity_matrix_copy[i, j]))
                    
                    top_similarities.sort(key=lambda x: x[2], reverse=True)
                    top_similarities = top_similarities[:n_top]
                    
                    top_sim_df = pd.DataFrame({
                        'Kullanıcı 1': [f"User_{i+1}" for i, _, _ in top_similarities],
                        'Kullanıcı 2': [f"User_{j+1}" for _, j, _ in top_similarities],
                        'Benzerlik Skoru': [f"{sim:.4f}" for _, _, sim in top_similarities]
                    })
                    st.dataframe(top_sim_df, width='stretch')
                    
                    # Görselleştirme
                    st.markdown("### 📊 Benzerlik Matrisi Görselleştirmesi")
                    fig, ax = plt.subplots(figsize=(10, 8))
                    im = ax.imshow(user_similarity, cmap='viridis', aspect='auto')
                    ax.set_xlabel('Kullanıcılar', fontsize=12)
                    ax.set_ylabel('Kullanıcılar', fontsize=12)
                    ax.set_title('Kullanıcı Benzerlik Matrisi (Kosinüs Benzerliği)', fontsize=14, fontweight='bold')
                    plt.colorbar(im, ax=ax, label='Benzerlik Skoru')
                    st.pyplot(fig)
                
                elif analysis_option == "📚 Konu Benzerliği (Ürün Benzerliği)":
                    st.subheader("📚 Konu Benzerliği (Ürün Benzerliği)")
                    with st.expander("ℹ️ Konu Benzerliği Nedir?", expanded=False):
                        st.markdown("""
                        **Konu Benzerliği** (Ürün Benzerliği):
                        - `svd.components_` üzerinden kosinüs benzerliği kullanılarak hesaplanır
                        - Hangi ürünlerin (konuların) benzer latent profillere sahip olduğunu gösterir
                        - Benzerlik değeri 1'e yakınsa ürünler çok benzerdir
                        - Benzerlik değeri 0'a yakınsa ürünler farklıdır
                        - Bu bilgi, "benzer ürünler benzer kullanıcılar tarafından beğenilir" prensibine dayanır
                        """)
                    
                    item_similarity = svd_model.get_item_similarity()
                    
                    # Ürün ID'leri
                    item_ids = [f"Item_{i+1}" for i in range(item_similarity.shape[0])]
                    
                    similarity_df = pd.DataFrame(
                        item_similarity,
                        index=item_ids,
                        columns=item_ids
                    )
                    
                    st.dataframe(similarity_df, width='stretch', height=400)
                    
                    # En benzer ürün çiftleri
                    st.markdown("### 🔝 En Benzer Ürün Çiftleri")
                    # Diagonal'i -1 yap (kendisiyle benzerlik hariç)
                    similarity_matrix_copy = item_similarity.copy()
                    np.fill_diagonal(similarity_matrix_copy, -1)
                    
                    # En yüksek benzerlik değerlerini bul
                    n_top = min(10, len(item_ids))
                    top_similarities = []
                    for i in range(len(item_ids)):
                        for j in range(i+1, len(item_ids)):
                            top_similarities.append((i, j, similarity_matrix_copy[i, j]))
                    
                    top_similarities.sort(key=lambda x: x[2], reverse=True)
                    top_similarities = top_similarities[:n_top]
                    
                    top_sim_df = pd.DataFrame({
                        'Ürün 1': [f"Item_{i+1}" for i, _, _ in top_similarities],
                        'Ürün 2': [f"Item_{j+1}" for _, j, _ in top_similarities],
                        'Benzerlik Skoru': [f"{sim:.4f}" for _, _, sim in top_similarities]
                    })
                    st.dataframe(top_sim_df, width='stretch')
                    
                    # Görselleştirme
                    st.markdown("### 📊 Benzerlik Matrisi Görselleştirmesi")
                    fig, ax = plt.subplots(figsize=(10, 8))
                    im = ax.imshow(item_similarity, cmap='viridis', aspect='auto')
                    ax.set_xlabel('Ürünler', fontsize=12)
                    ax.set_ylabel('Ürünler', fontsize=12)
                    ax.set_title('Ürün Benzerlik Matrisi (Kosinüs Benzerliği)', fontsize=14, fontweight='bold')
                    plt.colorbar(im, ax=ax, label='Benzerlik Skoru')
                    st.pyplot(fig)
                
                elif analysis_option == "🔮 Tahmini Puanlar (Predicted Scores)":
                    st.subheader("🔮 Tahmini Puanlar (Predicted Scores)")
                    with st.expander("ℹ️ Tahmini Puanlar Nedir?", expanded=False):
                        st.markdown("""
                        **Tahmini Puanlar** (Predicted Scores):
                        - `latent_matrix @ svd.components_` ile yeniden oluşturulan matristir
                        - Orijinalde NaN olan hücreler için SVD modelinin tahmin ettiği puanları içerir
                        - Bu matris, bir kaynağa henüz içerik üretmediği konular için potansiyel puanları gösterir
                        - Orijinal rating'ler ile karşılaştırılarak model performansı değerlendirilebilir
                        """)
                    
                    # Kullanıcı ve ürün ID'leri
                    user_ids = [f"User_{i+1}" for i in range(n_users)]
                    item_ids = [f"Item_{i+1}" for i in range(n_items)]
                    
                    predicted_scores_df = svd_model.get_predicted_scores(
                        as_dataframe=True,
                        user_ids=user_ids,
                        item_ids=item_ids
                    )
                    
                    st.dataframe(predicted_scores_df, width='stretch', height=400)
                    
                    # İstatistikler
                    predicted_scores = svd_model.get_predicted_scores(as_dataframe=False)
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        st.metric("Minimum Tahmin", f"{np.min(predicted_scores):.2f}")
                    with col2:
                        st.metric("Maksimum Tahmin", f"{np.max(predicted_scores):.2f}")
                    with col3:
                        st.metric("Ortalama Tahmin", f"{np.mean(predicted_scores):.2f}")
                    with col4:
                        st.metric("Std Sapma", f"{np.std(predicted_scores):.2f}")
                    
                    # Orijinal matris ile karşılaştırma (eğer varsa)
                    st.markdown("### 📊 Orijinal vs Tahmini Puanlar Karşılaştırması")
                    from scipy.sparse import issparse
                    if issparse(rating_matrix):
                        original_ratings = rating_matrix.toarray()
                    else:
                        original_ratings = rating_matrix.copy()
                    
                    # Sadece mevcut rating'leri karşılaştır
                    mask = ~np.isnan(original_ratings)
                    if np.sum(mask) > 0:
                        original_values = original_ratings[mask]
                        predicted_values = predicted_scores[mask]
                        
                        comparison_df = pd.DataFrame({
                            'Orijinal Rating': original_values[:20],  # İlk 20 örnek
                            'Tahmini Rating': predicted_values[:20],
                            'Fark': np.abs(original_values[:20] - predicted_values[:20])
                        })
                        st.dataframe(comparison_df, width='stretch')
                        
                        # RMSE hesapla
                        rmse_comparison = np.sqrt(np.mean((original_values - predicted_values)**2))
                        st.metric("RMSE (Orijinal vs Tahmini)", f"{rmse_comparison:.4f}")
                
                elif analysis_option == "📈 Rating Matrisi Görselleştirmesi":
                    # Rating matrisi görselleştirme
                    st.subheader("📊 Rating Matrisi Görselleştirmesi")
                    with st.expander("ℹ️ Rating Matrisi Ne Gösteriyor?", expanded=False):
                        st.markdown("""
                        **Rating Matrisi**:
                        - **Satırlar**: Kullanıcılar
                        - **Sütunlar**: Ürünler
                        - **Renkler**: Rating değerleri (koyu = düşük, açık = yüksek)
                        - **Beyaz alanlar**: Eksik rating'ler (kullanıcı henüz değerlendirmemiş)
                        - Bu matris, öneri sisteminin temel girdisidir
                        - SVD, bu matrisi faktörize ederek eksik değerleri tahmin eder
                        """)
                    
                    fig2 = plot_ratings_matrix(rating_matrix)
                    st.pyplot(fig2)
        except Exception as e:
            st.error(f"❌ Analiz sırasında hata oluştu: {str(e)}")
            import traceback
            with st.expander("🔍 Detaylı Hata Mesajı"):
                st.code(traceback.format_exc())
            st.info("💡 Lütfen modelin başarıyla eğitildiğinden emin olun.")


def get_data_file_format_selector(model_name="model"):
    """
    Veri dosyası formatı seçici (veri matrisi için - rating matrisi değil)
    
    Args:
        model_name: Model adı (key için)
        
    Returns:
        Seçilen format string'i
    """
    format_options = [
        "📊 Excel/CSV Formatı (Her satır bir örnek, her sütun bir özellik)",
        "📋 Matris Formatı (Zaten matris formatında)",
        "📁 Her İkisi (Otomatik Tespit)"
    ]
    
    selected_format = st.radio(
        "📄 Dosya Formatı Seçin",
        format_options,
        key=f"{model_name}_data_format",
        help="Yükleyeceğiniz veri dosyasının formatını seçin"
    )
    
    # Format açıklaması
    if selected_format.startswith("📊"):
        st.info("💡 **Excel/CSV Formatı**: Her satır bir veri örneği, her sütun bir özellik. Sayısal veri olmalıdır.")
    elif selected_format.startswith("📋"):
        st.info("💡 **Matris Formatı**: Veri zaten matris formatında (n_samples x n_features)")
    elif selected_format.startswith("📁"):
        st.info("💡 **Otomatik Tespit**: Sistem dosyayı analiz edip uygun formatı otomatik seçer")
    
    return selected_format


def show_svd_noise_reduction():
    """SVD gürültü temizleme"""
    st.header("🔇 SVD - Gürültü Temizleme")
    
    # Info bölümü
    with st.expander("ℹ️ SVD Gürültü Temizleme Hakkında Bilgi", expanded=False):
        st.markdown("""
        ### SVD ile Gürültü Temizleme Nasıl Çalışır?
        
        **Prensip**: SVD, veriyi önemli bileşenlere ve gürültüye ayırır.
        
        ### Adımlar:
        
        1. **SVD Uygulama**: Veri matrisini tekil değerlerine ayırır
        2. **Bileşen Seçimi**: Sadece önemli bileşenleri tutar (yüksek tekil değerler)
        3. **Gürültü Filtreleme**: Düşük tekil değerli bileşenleri atar (gürültü)
        4. **Yeniden Oluşturma**: Seçilen bileşenlerle temiz veriyi yeniden oluşturur
        
        ### Varyans Eşiği
        
        - **0.95 (95%)**: Verinin %95'ini korur, %5 gürültüyü temizler
        - **0.90 (90%)**: Daha agresif temizleme, daha fazla gürültü kaldırır
        - **0.99 (99%)**: Çok az temizleme, neredeyse tüm veriyi korur
        
        ### Kullanım Alanları
        
        - Görüntü gürültü temizleme
        - Sinyal işleme
        - Veri ön işleme
        - Özellik çıkarımı
        """)
    
    st.markdown("""
    SVD kullanarak veri gürültüsünü temizleme.
    """)
    
    col1, col2 = st.columns(2)
    with col1:
        n_samples = st.slider("Örnek Sayısı", 100, 500, 200)
        n_features = st.slider("Özellik Sayısı", 50, 200, 100)
    with col2:
        noise_level = st.slider("Gürültü Seviyesi", 0.05, 0.5, 0.2)
        threshold = st.slider("Varyans Eşiği", 0.8, 0.99, 0.95)
    
    if st.button("Gürültü Temizle"):
        with st.spinner("Gürültü temizleniyor..."):
            # Veri oluştur
            X, _ = generate_sample_data(n_samples=n_samples, n_features=n_features)
            X_noisy = generate_noisy_data(X, noise_level=noise_level)
            
            # Gürültü temizle
            noise_reducer = SVDNoiseReducer(n_components=None, threshold=threshold)
            noise_reducer.fit(X_noisy)
            X_denoised = noise_reducer.denoise(X_noisy)
            
            # Metrikler
            mse_original = np.mean((X - X_noisy)**2)
            mse_denoised = np.mean((X - X_denoised)**2)
            improvement = ((mse_original - mse_denoised) / mse_original * 100)
            noise_reduction_ratio = noise_reducer.get_noise_reduction_ratio()
            
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("Gürültü MSE", f"{mse_original:.6f}")
            with col2:
                st.metric("Temizlenmiş MSE", f"{mse_denoised:.6f}")
            with col3:
                st.metric("İyileştirme", f"{improvement:.2f}%")
            with col4:
                st.metric("Varyans Korunma", f"{noise_reduction_ratio:.2%}")
            
            # Görselleştirme
            st.subheader("📊 Gürültü Temizleme Sonuçları")
            with st.expander("ℹ️ Bu Görselleştirme Ne Gösteriyor?", expanded=False):
                st.markdown("""
                **Üç Panel Karşılaştırması**:
                
                1. **Orijinal Veri**: Temiz, gürültüsüz orijinal veri
                2. **Gürültülü Veri**: Rastgele gürültü eklenmiş veri (gerçek dünya senaryosu)
                3. **Temizlenmiş Veri**: SVD ile gürültü temizlenmiş veri
                
                **Renk Haritası (Viridis)**:
                - Koyu renkler: Düşük değerler
                - Açık renkler: Yüksek değerler
                - Düzgün geçişler: Temiz veri
                - Rastgele noktalar: Gürültü
                
                **Karşılaştırma**: Temizlenmiş veri, orijinal veriye ne kadar yakınsa, algoritma o kadar başarılıdır.
                """)
            
            fig, axes = plt.subplots(1, 3, figsize=(18, 5))
            
            # Orijinal
            im1 = axes[0].imshow(X[:50, :50], cmap='viridis', aspect='auto')
            axes[0].set_title('Orijinal Veri (Temiz)', fontsize=12, fontweight='bold')
            plt.colorbar(im1, ax=axes[0])
            
            # Gürültülü
            im2 = axes[1].imshow(X_noisy[:50, :50], cmap='viridis', aspect='auto')
            axes[1].set_title('Gürültülü Veri', fontsize=12, fontweight='bold')
            plt.colorbar(im2, ax=axes[1])
            
            # Temizlenmiş
            im3 = axes[2].imshow(X_denoised[:50, :50], cmap='viridis', aspect='auto')
            axes[2].set_title('Temizlenmiş Veri (SVD ile)', fontsize=12, fontweight='bold')
            plt.colorbar(im3, ax=axes[2])
            
            st.pyplot(fig)
            
            # Varyans analizi
            st.subheader("📈 Varyans Korunma Analizi")
            with st.expander("ℹ️ Varyans Analizi Ne Anlama Geliyor?", expanded=False):
                st.markdown("""
                **Kümülatif Varyans Grafiği**:
                - **X Ekseni**: Kullanılan bileşen sayısı
                - **Y Ekseni**: Korunan veri varyansı oranı (0-1 arası)
                - **Kırmızı Çizgi**: Seçilen eşik değeri (örn: %95)
                - **Yeşil Çizgi**: Otomatik seçilen optimal bileşen sayısı
                
                **Yorumlama**:
                - İlk birkaç bileşen genellikle varyansın çoğunu açıklar
                - Eğri yataylaştığında, ek bileşenler çok az bilgi ekler
                - Optimal nokta: Eşik değerine ulaşan minimum bileşen sayısı
                - Daha fazla bileşen = daha az gürültü temizleme ama daha yavaş işlem
                """)
            
            component_counts, variance_ratios = noise_reducer.get_optimal_components(X_noisy)
            
            fig2, ax = plt.subplots(figsize=(10, 5))
            ax.plot(component_counts, variance_ratios, 'o-', linewidth=2, markersize=6)
            ax.axhline(y=threshold, color='r', linestyle='--', linewidth=2, label=f'{threshold:.0%} Eşiği')
            ax.axvline(x=noise_reducer.n_components, color='g', 
                      linestyle='--', linewidth=2, label=f'Seçilen ({noise_reducer.n_components})')
            ax.set_xlabel('Bileşen Sayısı', fontsize=12)
            ax.set_ylabel('Kümülatif Varyans Oranı', fontsize=12)
            ax.set_title('Varyans Korunma Analizi - Optimal Bileşen Seçimi', fontsize=14, fontweight='bold')
            ax.legend(fontsize=10)
            ax.grid(True, alpha=0.3)
            st.pyplot(fig2)


def show_pca_visualization():
    """PCA görselleştirme"""
    st.header("📈 PCA - Veri Görselleştirme")
    
    # Info bölümü
    with st.expander("ℹ️ PCA (Principal Component Analysis) Hakkında Bilgi", expanded=False):
        st.markdown("""
        ### PCA Nedir?
        
        **PCA**, yüksek boyutlu veriyi daha düşük boyutlu bir uzaya izdüşüren bir boyut azaltma tekniğidir.
        
        ### Nasıl Çalışır?
        
        1. **Kovaryans Matrisi**: Veri özellikleri arasındaki ilişkileri hesaplar
        2. **Özvektörler**: Verinin ana yönlerini (principal components) bulur
        3. **Özdeğerler**: Her bileşenin ne kadar varyans açıkladığını gösterir
        4. **İzdüşüm**: Veriyi yeni düşük boyutlu uzaya dönüştürür
        
        ### Kullanım Alanları
        
        - **Veri Görselleştirme**: Yüksek boyutlu veriyi 2D/3D'de görselleştirme
        - **Özellik Seçimi**: En önemli özellikleri belirleme
        - **Boyut Azaltma**: Gürültüyü azaltma ve hesaplama hızını artırma
        - **Ön İşleme**: Makine öğrenmesi modelleri için veri hazırlama
        
        ### Metrikler
        
        - **Açıklanan Varyans**: Her bileşenin verinin ne kadarını açıkladığı
        - **Kümülatif Varyans**: İlk N bileşenin toplam açıkladığı varyans
        - **%95 Varyans Kuralı**: Verinin %95'ini açıklayan minimum bileşen sayısı
        """)
    
    st.markdown("""
    PCA kullanarak veri görselleştirme ve özellik seçimi.
    """)
    
    # Optimal parametreleri al
    optimal_params = get_optimal_model_params("pca", n_samples=500, n_features=50)
    optimal_n_components = optimal_params['n_components']
    
    col1, col2 = st.columns(2)
    with col1:
        n_samples = st.slider("Örnek Sayısı", 100, 1000, 500)
        n_features = st.slider("Özellik Sayısı", 20, 100, 50)
    with col2:
        n_clusters = st.slider("Küme Sayısı", 2, 10, 5)
        n_components = st.slider(
            "Gösterilecek Bileşen", 
            5, 30, 
            optimal_n_components,
            help=f"Önerilen değer: {optimal_n_components}"
        )
    
    if st.button("PCA Uygula"):
        with st.spinner("PCA uygulanıyor..."):
            # Veri oluştur
            X, y = generate_sample_data(
                n_samples=n_samples, 
                n_features=n_features, 
                n_clusters=n_clusters
            )
            
            # PCA
            pca = PCAAnalyzer(n_components=None)
            X_transformed = pca.fit_transform(X)
            
            # Metrikler
            n_95 = pca.get_optimal_components(X, variance_threshold=0.95)
            variance_first_5 = np.sum(pca.explained_variance_ratio_[:5])
            
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Toplam Bileşen", len(pca.explained_variance_ratio_))
            with col2:
                st.metric("%95 Varyans için Bileşen", n_95)
            with col3:
                st.metric("İlk 5 Bileşen Varyansı", f"{variance_first_5:.2%}")
            
            # Açıklanan varyans grafiği
            st.subheader("📊 Açıklanan Varyans Analizi")
            with st.expander("ℹ️ Bu Grafikler Ne Anlama Geliyor?", expanded=False):
                st.markdown("""
                **Sol Grafik - Bireysel Varyans**:
                - Her bileşenin (PC) ne kadar varyans açıkladığını gösterir
                - İlk birkaç bileşen genellikle en yüksek varyansa sahiptir
                - "Elbow" noktası: Optimal bileşen sayısını belirler
                
                **Sağ Grafik - Kümülatif Varyans**:
                - İlk N bileşenin toplam açıkladığı varyans
                - %95 eşiği: Verinin %95'ini korumak için gerekli bileşen sayısı
                - Eğri yataylaştığında, ek bileşenler az bilgi ekler
                """)
            
            fig = pca.plot_explained_variance(n_components=n_components)
            st.pyplot(fig)
            
            # 2D ve 3D izdüşüm
            st.subheader("🎨 Veri Görselleştirme (2D & 3D)")
            with st.expander("ℹ️ İzdüşüm Grafikleri Ne Gösteriyor?", expanded=False):
                st.markdown("""
                **2D İzdüşüm (PC1 vs PC2)**:
                - Yüksek boyutlu veriyi 2 boyuta indirger
                - Renkler: Farklı sınıfları/kümeleri gösterir
                - Yakın noktalar: Benzer veri örnekleri
                - Ayrık gruplar: Farklı sınıflar/kümeler
                
                **3D İzdüşüm (PC1 vs PC2 vs PC3)**:
                - Daha fazla bilgi korur (3 boyut)
                - Daha iyi ayrım sağlar
                - İnteraktif olarak döndürülebilir
                
                **Yorumlama**:
                - İyi ayrılmış gruplar = PCA başarılı
                - Karışık noktalar = Veri karmaşık veya daha fazla bileşen gerekli
                """)
            
            col1, col2 = st.columns(2)
            with col1:
                fig2 = pca.plot_2d_projection(X, y=y)
                st.pyplot(fig2)
            
            # 3D izdüşüm
            with col2:
                fig3 = pca.plot_3d_projection(X, y=y)
                st.pyplot(fig3)
            
            # Özellik önemi
            st.subheader("🔍 Özellik Önemi Analizi")
            with st.expander("ℹ️ Özellik Önemi Ne Anlama Geliyor?", expanded=False):
                st.markdown("""
                **Özellik Önemi Skorları**:
                - Her özelliğin PCA bileşenlerindeki katkısını gösterir
                - **Yüksek skor**: Özellik, verinin varyansını açıklamada önemli
                - **Düşük skor**: Özellik daha az önemli veya gürültülü
                
                **Kullanım**:
                - Özellik seçimi: Yüksek skorlu özellikleri tut
                - Gürültü temizleme: Düşük skorlu özellikleri kaldır
                - Model optimizasyonu: Önemli özellikleri önceliklendir
                """)
            
            feature_importance = pca.get_feature_importance()
            top_features = np.argsort(feature_importance)[-20:][::-1]
            
            fig4, ax = plt.subplots(figsize=(10, 6))
            ax.barh(range(len(top_features)), feature_importance[top_features], alpha=0.7)
            ax.set_yticks(range(len(top_features)))
            ax.set_yticklabels([f'Özellik {i+1}' for i in top_features])
            ax.set_xlabel('Önem Skoru', fontsize=12)
            ax.set_title('En Önemli 20 Özellik - PCA Özellik Önemi', fontsize=14, fontweight='bold')
            ax.grid(True, alpha=0.3, axis='x')
            st.pyplot(fig4)


def show_nmf_image_processing():
    """NMF görüntü işleme"""
    st.header("🖼️ NMF - Görüntü İşleme")
    
    # Info bölümü
    with st.expander("ℹ️ NMF Görüntü İşleme Hakkında Bilgi", expanded=False):
        st.markdown("""
        ### NMF (Non-negative Matrix Factorization) Nedir?
        
        **NMF**, matrisleri sadece **pozitif değerlerle** faktörize eden bir yöntemdir.
        
        ### Görüntü İşlemede Kullanımı
        
        1. **Basis Images (Temel Görüntüler)**: Görüntülerin temel yapı taşlarını bulur
        2. **Katsayı Matrisi**: Her görüntünün bu temel yapı taşlarını nasıl kullandığını gösterir
        3. **Yeniden Oluşturma**: Temel görüntüler ve katsayılarla orijinal görüntüyü yeniden oluşturur
        
        ### Avantajları
        
        - **Pozitif Değerler**: Görüntü piksel değerleri doğal olarak pozitiftir
        - **Yorumlanabilirlik**: Temel görüntüler anlamlı pattern'ler içerir
        - **Sıkıştırma**: Az sayıda temel görüntü ile çok görüntüyü temsil edebilir
        - **Gürültü Azaltma**: Düşük rank yaklaşımı gürültüyü filtreler
        
        ### Sıkıştırma Oranı
        
        - **Yüksek oran**: Daha fazla sıkıştırma, daha az kalite
        - **Düşük oran**: Daha az sıkıştırma, daha yüksek kalite
        - Optimal: Kalite ve boyut arasında denge
        """)
    
    # Veri kaynağı seçimi
    data_source = st.radio(
        "Veri Kaynağı Seçin",
        ["📁 Görüntü Dosyası Yükle", "🎲 Örnek Görüntüler Kullan"],
        horizontal=True
    )
    
    # Dosya formatı seçimi (sadece görüntü yükleme seçildiyse)
    if data_source == "📁 Görüntü Dosyası Yükle":
        file_format = get_file_format_selector("nmf_image", include_image=True)
        if not file_format.startswith("🖼️"):
            st.warning("⚠️ Bu model görüntü işleme için tasarlanmıştır. Lütfen fotoğraf/görüntü formatını seçin.")
    
    uploaded_file = None
    images_flat = None
    image_shape = None
    img_array = None  # Orijinal görüntü array'i (görselleştirme için)
    
    if data_source == "📁 Görüntü Dosyası Yükle":
        st.markdown("### 📤 Görüntü Dosyası Yükle")
        
        # Dosya formatı kontrolü
        if file_format and file_format.startswith("🖼️"):
            uploaded_file = st.file_uploader(
                "Görüntü dosyası seçin (JPG, PNG, BMP)",
                type=['jpg', 'jpeg', 'png', 'bmp'],
                help="Yüklediğiniz görüntü analiz edilecek ve NMF ile temel bileşenlere ayrılacaktır."
            )
        else:
            uploaded_file = None
            st.info("💡 Lütfen yukarıda 'Fotoğraf/Görüntü' formatını seçin.")
        
        if uploaded_file is not None:
            try:
                from PIL import Image
                import io
                
                # Dosya içeriğini oku (bir kez okunabilir, bu yüzden sakla)
                file_bytes = uploaded_file.read()
                uploaded_file.seek(0)  # Reset file pointer
                
                # Görüntüyü yükle
                image = Image.open(io.BytesIO(file_bytes))
                
                # Görüntüyü göster
                st.image(image, caption="Yüklenen Görüntü", use_container_width=True)
                
                # Görüntü bilgileri
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("Genişlik", f"{image.width} px")
                with col2:
                    st.metric("Yükseklik", f"{image.height} px")
                with col3:
                    st.metric("Mod", image.mode)
                
                # Görüntüyü gri tonlamaya çevir ve normalize et
                if image.mode != 'L':
                    image = image.convert('L')
                
                # NumPy array'e çevir
                img_array = np.array(image, dtype=np.float32)
                img_array = img_array / 255.0  # Normalize et [0, 1]
                
                # Görüntüyü tek bir görüntü olarak işle (1 görüntü, düzleştirilmiş)
                image_shape = img_array.shape
                images_flat = img_array.flatten().reshape(1, -1)
                
                st.success(f"✅ Görüntü yüklendi! Boyut: {image_shape}")
                
            except Exception as e:
                st.error(f"❌ Görüntü yüklenirken hata oluştu: {str(e)}")
                st.info("💡 Lütfen geçerli bir görüntü dosyası yüklediğinizden emin olun.")
    
    else:
        st.markdown("### 🎲 Örnek Görüntüler")
        col1, col2 = st.columns(2)
        # Optimal parametreleri al
        optimal_params = get_optimal_model_params("nmf_image")
        optimal_n_components = optimal_params['n_components']
        
        with col1:
            n_images = st.slider("Görüntü Sayısı", 50, 200, 100)
        with col2:
            n_components = st.slider(
                "Bileşen Sayısı", 
                5, 50, 
                optimal_n_components,
                help=f"Önerilen değer: {optimal_n_components}"
            )
    
    # Bileşen sayısı slider'ı (görüntü yüklendiyse göster)
    if data_source == "📁 Görüntü Dosyası Yükle" and images_flat is not None:
        n_components = st.slider("Bileşen Sayısı", 5, 100, 20, key="uploaded_components")
    
    # Analiz butonu
    analyze_button = st.button("🔍 Görüntüyü Analiz Et" if data_source == "📁 Görüntü Dosyası Yükle" else "🎲 Örnek Görüntüleri İşle")
    
    if analyze_button:
        if data_source == "📁 Görüntü Dosyası Yükle":
            if uploaded_file is None:
                st.warning("⚠️ Lütfen önce bir görüntü dosyası yükleyin!")
            elif images_flat is not None and image_shape is not None and img_array is not None:
                
                with st.spinner("Görüntü analiz ediliyor..."):
                    try:
                        # NMF
                        nmf_image = NMFImageProcessor(n_components=n_components)
                        nmf_image.fit(images_flat)
                        
                        # Yeniden oluştur
                        reconstructed = nmf_image.reconstruct()
                        basis_images = nmf_image.get_basis_images(image_shape)
                        compression_ratio = nmf_image.get_compression_ratio(image_shape)
                        
                        # Metrikler
                        mse = np.mean((images_flat - reconstructed)**2)
                        psnr = 20 * np.log10(1.0 / (np.sqrt(mse) + 1e-10)) if mse > 0 else float('inf')
                        
                        # Sonuçları göster
                        st.subheader("📊 Analiz Sonuçları")
                        col1, col2, col3, col4 = st.columns(4)
                        with col1:
                            st.metric("Sıkıştırma Oranı", f"{compression_ratio:.2f}x")
                        with col2:
                            st.metric("MSE", f"{mse:.6f}")
                        with col3:
                            st.metric("PSNR (dB)", f"{psnr:.2f}")
                        with col4:
                            st.metric("Basis Görüntü Sayısı", len(basis_images))
                        
                        # Görselleştirme
                        st.subheader("🖼️ Görüntü Analizi")
                        with st.expander("ℹ️ Analiz Sonuçları Hakkında", expanded=False):
                            st.markdown("""
                            **Analiz Sonuçları**:
                            
                            1. **Orijinal Görüntü**: Yüklediğiniz görüntü
                            2. **Temel Görüntüler (Basis Images)**: NMF'nin bulduğu temel yapı taşları
                               - Her temel görüntü, görüntünüzdeki ortak pattern'leri temsil eder
                               - Örnek: Kenarlar, köşeler, dairesel şekiller, dokular vb.
                            3. **Yeniden Oluşturulmuş Görüntü**: Temel görüntüler kullanılarak yeniden oluşturulan görüntü
                               - Orijinal görüntüye ne kadar yakınsa, analiz o kadar başarılı
                            
                            **Metrikler**:
                            - **MSE (Mean Squared Error)**: Düşük = Daha iyi kalite
                            - **PSNR (Peak Signal-to-Noise Ratio)**: Yüksek = Daha iyi kalite
                            - **Sıkıştırma Oranı**: Orijinal boyut / Sıkıştırılmış boyut
                            """)
                        
                        # Orijinal görüntü
                        st.markdown("### 📸 Orijinal Görüntü")
                        fig_orig, ax_orig = plt.subplots(figsize=(8, 8))
                        ax_orig.imshow(img_array, cmap='gray')
                        ax_orig.set_title("Orijinal Görüntü", fontsize=14, fontweight='bold')
                        ax_orig.axis('off')
                        st.pyplot(fig_orig)
                        
                        # Temel görüntüler
                        st.markdown("### 🧩 Temel Görüntüler (Basis Images)")
                        with st.expander("ℹ️ Temel Görüntüler Hakkında", expanded=False):
                            st.markdown("""
                            **Temel Görüntüler (Basis Images)**:
                            - NMF algoritmasının öğrendiği temel yapı taşlarıdır
                            - Görüntünüz, bu temel görüntülerin bir kombinasyonudur
                            - Sayıları bileşen sayısına eşittir (örn: 20 bileşen = 20 temel görüntü)
                            - Anlamlı pattern'ler içerir (kenarlar, şekiller, dokular)
                            """)
                        
                        basis_flat = basis_images.reshape(len(basis_images), -1)
                        fig_basis = plot_image_grid(basis_flat, image_shape, n_cols=5)
                        st.pyplot(fig_basis)
                        
                        # Yeniden oluşturulmuş görüntü
                        st.markdown("### 🔄 Yeniden Oluşturulmuş Görüntü")
                        reconstructed_img = reconstructed[0].reshape(image_shape)
                        fig_recon, ax_recon = plt.subplots(1, 2, figsize=(16, 8))
                        
                        ax_recon[0].imshow(img_array, cmap='gray')
                        ax_recon[0].set_title("Orijinal", fontsize=14, fontweight='bold')
                        ax_recon[0].axis('off')
                        
                        ax_recon[1].imshow(reconstructed_img, cmap='gray')
                        ax_recon[1].set_title(f"Yeniden Oluşturulmuş (MSE: {mse:.6f})", fontsize=14, fontweight='bold')
                        ax_recon[1].axis('off')
                        
                        st.pyplot(fig_recon)
                        
                        # Fark görüntüsü
                        st.markdown("### 🔍 Fark Görüntüsü (Orijinal - Yeniden Oluşturulmuş)")
                        diff_img = np.abs(img_array - reconstructed_img)
                        fig_diff, ax_diff = plt.subplots(figsize=(8, 8))
                        im = ax_diff.imshow(diff_img, cmap='hot')
                        ax_diff.set_title("Fark Görüntüsü (Kırmızı = Daha Fazla Fark)", fontsize=14, fontweight='bold')
                        ax_diff.axis('off')
                        plt.colorbar(im, ax=ax_diff, label='Fark Değeri')
                        st.pyplot(fig_diff)
                        
                        st.success("✅ Analiz tamamlandı!")
                        
                    except Exception as e:
                        st.error(f"❌ Analiz sırasında hata oluştu: {str(e)}")
                        import traceback
                        st.code(traceback.format_exc())
        else:
            # Örnek görüntüler
            with st.spinner("Görüntüler işleniyor..."):
                # Görüntü yükle
                images_flat, image_shape = load_sample_images(n_images=n_images)
                
                # NMF
                nmf_image = NMFImageProcessor(n_components=n_components)
                nmf_image.fit(images_flat)
                
                # Yeniden oluştur
                reconstructed = nmf_image.reconstruct()
                basis_images = nmf_image.get_basis_images(image_shape)
                compression_ratio = nmf_image.get_compression_ratio(image_shape)
                
                # Metrikler
                mse = np.mean((images_flat - reconstructed)**2)
                
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("Sıkıştırma Oranı", f"{compression_ratio:.2f}x")
                with col2:
                    st.metric("MSE", f"{mse:.6f}")
                with col3:
                    st.metric("Basis Görüntü Sayısı", len(basis_images))
                
                # Görselleştirme
                st.subheader("🖼️ Görüntü Karşılaştırması")
                with st.expander("ℹ️ Bu Görüntüler Ne Gösteriyor?", expanded=False):
                    st.markdown("""
                    **Üç Görüntü Seti**:
                    
                    1. **Orijinal Görüntüler**: İşlenmemiş, orijinal görüntüler
                    2. **Temel Görüntüler (Basis Images)**: NMF'nin bulduğu temel yapı taşları
                       - Her temel görüntü, görüntü koleksiyonundaki ortak pattern'leri temsil eder
                       - Örnek: Kenarlar, köşeler, dairesel şekiller vb.
                    3. **Yeniden Oluşturulmuş Görüntüler**: Temel görüntüler kullanılarak yeniden oluşturulan görüntüler
                       - Orijinal görüntüye ne kadar yakınsa, sıkıştırma o kadar başarılı
                    
                    **Kalite Değerlendirmesi**:
                    - İyi yeniden oluşturma: Orijinal ve yeniden oluşturulmuş görüntüler benzer
                    - Düşük MSE: Daha yüksek kalite
                    - Yüksek sıkıştırma oranı: Daha küçük dosya boyutu
                    """)
                
                st.subheader("📸 Orijinal Görüntüler (İlk 10)")
                fig1 = plot_image_grid(images_flat[:10], image_shape, n_cols=5)
                st.pyplot(fig1)
                
                st.subheader("🧩 Temel Görüntüler (Basis Images)")
                with st.expander("ℹ️ Temel Görüntüler Hakkında", expanded=False):
                    st.markdown("""
                    **Temel Görüntüler (Basis Images)**:
                    - NMF algoritmasının öğrendiği temel yapı taşlarıdır
                    - Her görüntü, bu temel görüntülerin bir kombinasyonudur
                    - Sayıları bileşen sayısına eşittir (örn: 20 bileşen = 20 temel görüntü)
                    - Anlamlı pattern'ler içerir (kenarlar, şekiller, dokular)
                    """)
                
                basis_flat = basis_images.reshape(len(basis_images), -1)
                fig2 = plot_image_grid(basis_flat, image_shape, n_cols=5)
                st.pyplot(fig2)
                
                st.subheader("🔄 Yeniden Oluşturulmuş Görüntüler (İlk 10)")
                fig3 = plot_image_grid(reconstructed[:10], image_shape, n_cols=5)
                st.pyplot(fig3)


def extract_text_from_file(uploaded_file):
    """
    Yüklenen dosyadan metin çıkarır
    
    Args:
        uploaded_file: Streamlit file_uploader'dan gelen dosya
        
    Returns:
        Metin içeriği (string listesi - her satır bir doküman)
    """
    import io
    
    file_name = uploaded_file.name.lower()
    file_content = uploaded_file.read()
    
    documents = []
    
    try:
        if file_name.endswith(('.txt', '.text')):
            # TXT dosyası
            text = file_content.decode('utf-8', errors='ignore')
            # Paragraflara böl (boş satırlarla ayrılmış)
            paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
            if not paragraphs:
                # Eğer paragraf yoksa, satırlara böl
                paragraphs = [line.strip() for line in text.split('\n') if line.strip()]
            documents = paragraphs
            
        elif file_name.endswith(('.docx', '.doc')):
            # Word dosyası
            try:
                from docx import Document
                doc = Document(io.BytesIO(file_content))
                
                # Paragraflardan metin çıkar
                paragraphs = []
                for p in doc.paragraphs:
                    text = p.text.strip()
                    if text:
                        paragraphs.append(text)
                
                # Tablolardan da metin çıkar
                for table in doc.tables:
                    for row in table.rows:
                        row_text = ' '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                        if row_text:
                            paragraphs.append(row_text)
                
                # Başlıklar ve listeler de dahil
                if paragraphs:
                    documents = paragraphs
                else:
                    raise ValueError("Word dosyasından hiç metin çıkarılamadı. Dosya boş olabilir.")
                    
            except ImportError:
                raise ImportError("Word dosyaları için 'python-docx' kütüphanesi gerekli. Yüklemek için: pip install python-docx")
            except Exception as e:
                if file_name.endswith('.doc'):
                    raise ValueError(
                        f"Eski .doc formatındaki dosya okunamadı. "
                        f"Lütfen dosyanızı .docx formatına dönüştürün (Word'de 'Farklı Kaydet' ile .docx olarak kaydedin). "
                        f"Hata detayı: {str(e)}"
                    )
                else:
                    raise ValueError(f"Word dosyası okunamadı: {str(e)}")
        
        elif file_name.endswith('.pdf'):
            # PDF dosyası
            try:
                import PyPDF2
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                paragraphs = []
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text.strip():
                        # Sayfayı paragraflara böl
                        page_paragraphs = [p.strip() for p in page_text.split('\n\n') if p.strip()]
                        paragraphs.extend(page_paragraphs)
                documents = paragraphs if paragraphs else [pdf_reader.pages[0].extract_text()]
            except ImportError:
                try:
                    import pdfplumber
                    with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                        paragraphs = []
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text and page_text.strip():
                                page_paragraphs = [p.strip() for p in page_text.split('\n\n') if p.strip()]
                                paragraphs.extend(page_paragraphs)
                        documents = paragraphs if paragraphs else []
                except ImportError:
                    raise ImportError("PDF dosyaları için 'PyPDF2' veya 'pdfplumber' kütüphanesi gerekli. Yüklemek için: pip install PyPDF2 veya pip install pdfplumber")
        
        elif file_name.endswith('.csv'):
            # CSV dosyası
            # Farklı delimiter'ları dene
            delimiters = [',', ';', '\t']
            df = None
            for delimiter in delimiters:
                try:
                    file_bytes = io.BytesIO(file_content)
                    df = pd.read_csv(file_bytes, delimiter=delimiter, encoding='utf-8', on_bad_lines='skip')
                    if len(df.columns) > 1:  # En az 2 sütun varsa doğru delimiter bulunmuş
                        break
                except:
                    continue
            
            if df is None or df.empty:
                # Son çare olarak varsayılan ayarlarla dene
                file_bytes = io.BytesIO(file_content)
                df = pd.read_csv(file_bytes, encoding='utf-8', on_bad_lines='skip')
            
            # Her satırı bir doküman olarak al (tüm sütunları birleştir)
            documents = []
            for idx, row in df.iterrows():
                # Tüm sütunları birleştir
                row_text = ' '.join([str(val) for val in row.values if pd.notna(val) and str(val).strip()])
                if row_text.strip():
                    documents.append(row_text.strip())
        
        elif file_name.endswith(('.xlsx', '.xls')):
            # Excel dosyası
            df = pd.read_excel(io.BytesIO(file_content))
            # Her satırı bir doküman olarak al (tüm sütunları birleştir)
            documents = []
            for idx, row in df.iterrows():
                # Tüm sütunları birleştir
                row_text = ' '.join([str(val) for val in row.values if pd.notna(val) and str(val).strip()])
                if row_text.strip():
                    documents.append(row_text.strip())
        
        else:
            raise ValueError(f"Desteklenmeyen dosya formatı: {file_name}")
        
        # Boş dokümanları filtrele
        documents = [doc for doc in documents if doc and len(doc.strip()) > 10]  # En az 10 karakter
        
        if not documents:
            raise ValueError("Dosyadan hiç metin çıkarılamadı! Lütfen geçerli bir metin içeren dosya yükleyin.")
        
        return documents
        
    except Exception as e:
        raise Exception(f"Dosya okunurken hata oluştu: {str(e)}")


def show_nmf_topic_modeling():
    """NMF topic modeling"""
    st.header("📝 NMF - Topic Modeling")
    
    # Info bölümü
    with st.expander("ℹ️ NMF Topic Modeling Hakkında Bilgi", expanded=False):
        st.markdown("""
        ### Topic Modeling Nedir?
        
        **Topic Modeling**, metin dokümanlarında gizli konuları (topic) keşfetme tekniğidir.
        
        ### NMF ile Topic Modeling
        
        1. **TF-IDF Vektörizasyon**: Metinleri sayısal vektörlere dönüştürür
        2. **NMF Faktörizasyon**: Doküman-kelime matrisini faktörize eder
        3. **Topic Çıkarımı**: Her topic için önemli kelimeleri bulur
        4. **Doküman-Topic Dağılımı**: Her dokümanın hangi topic'lere ait olduğunu gösterir
        
        ### Çıktılar
        
        - **Topic'ler**: Her topic için en önemli kelimeler
        - **Topic Tutarlılığı**: Topic'lerin ne kadar tutarlı olduğu
        - **Doküman-Topic Dağılımı**: Her dokümanın topic skorları
        - **Yeni Doküman Tahmini**: Yeni bir metin için en uygun topic
        
        ### Kullanım Alanları
        
        - Haber kategorilendirme
        - Müşteri yorumları analizi
        - Araştırma makaleleri sınıflandırma
        - Sosyal medya içerik analizi
        """)
    
    # Veri kaynağı seçimi
    data_source = st.radio(
        "Veri Kaynağı Seçin",
        ["📁 Dosya Yükle (Word/PDF/Excel/CSV/TXT)", "🎲 Örnek Metinler Kullan"],
        horizontal=True
    )
    
    documents = None
    uploaded_file = None
    
    if data_source == "📁 Dosya Yükle (Word/PDF/Excel/CSV/TXT)":
        st.markdown("### 📤 Dosya Yükle")
        
        # Bilgilendirme
        st.info("""
        **📝 Desteklenen Dosya Formatları:**
        - **Word**: .docx (önerilen), .doc (eski format - .docx'e dönüştürmeniz önerilir)
        - **PDF**: .pdf
        - **Excel**: .xlsx, .xls
        - **Metin**: .txt, .csv
        
        **💡 İpucu:** Her paragraf ayrı bir doküman olarak işlenecektir.
        """)
        
        uploaded_file = st.file_uploader(
            "Dosya seçin (Word, PDF, Excel, CSV, TXT)",
            type=['xlsx', 'xls', 'csv', 'txt', 'docx', 'doc', 'pdf'],
            help="Yüklediğiniz dosyadan metin çıkarılacak ve topic modeling yapılacaktır.",
            key="nmf_topic_file_uploader"
        )
        
        if uploaded_file is not None:
            try:
                with st.spinner("Dosya okunuyor..."):
                    documents = extract_text_from_file(uploaded_file)
                    
                st.success(f"✅ Dosya yüklendi! {len(documents)} doküman bulundu.")
                
                # Dosya önizleme
                with st.expander("📄 Dosya Önizleme (İlk 3 Doküman)", expanded=False):
                    for i, doc in enumerate(documents[:3]):
                        st.markdown(f"**Doküman {i+1}:**")
                        st.text(doc[:500] + "..." if len(doc) > 500 else doc)
                        st.markdown("---")
                
                # İstatistikler
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("Toplam Doküman", len(documents))
                with col2:
                    avg_length = np.mean([len(doc.split()) for doc in documents])
                    st.metric("Ortalama Kelime Sayısı", f"{avg_length:.0f}")
                with col3:
                    total_words = sum([len(doc.split()) for doc in documents])
                    st.metric("Toplam Kelime", total_words)
                    
            except ImportError as e:
                st.error(f"❌ Gerekli kütüphane eksik: {str(e)}")
                st.info("💡 Lütfen gerekli kütüphaneyi yükleyin.")
            except Exception as e:
                st.error(f"❌ Dosya okunurken hata oluştu: {str(e)}")
                st.info("💡 Lütfen geçerli bir dosya yüklediğinizden emin olun.")
    
    else:
        st.markdown("### 🎲 Örnek Metinler")
        col1, col2 = st.columns(2)
        with col1:
            n_documents = st.slider("Doküman Sayısı", 50, 300, 200)
        with col2:
            n_topics = st.slider("Topic Sayısı", 3, 10, 5, key="example_topics")
        
        max_features = st.slider("Maksimum Kelime Sayısı", 200, 1000, 500, key="example_features")
    
    # Topic sayısı ve parametreler (dosya yüklendiyse)
    if data_source == "📁 Dosya Yükle (Word/PDF/Excel/CSV/TXT)" and documents is not None:
        # Optimal parametreleri al (doküman sayısına göre)
        optimal_params = get_optimal_model_params("nmf_topic")
        optimal_n_topics = optimal_params['n_topics']
        optimal_max_features = optimal_params['max_features']
        
        col1, col2 = st.columns(2)
        with col1:
            n_topics = st.slider(
                "Topic Sayısı", 
                3, 15, 
                optimal_n_topics, 
                key="file_topics",
                help=f"Önerilen değer: {optimal_n_topics}"
            )
        with col2:
            max_features = st.slider(
                "Maksimum Kelime Sayısı", 
                200, 2000, 
                optimal_max_features, 
                key="file_features",
                help=f"Önerilen değer: {optimal_max_features}"
            )
    
    # Analiz butonu
    analyze_button = st.button("🔍 Topic'leri Bul" if data_source == "📁 Dosya Yükle (Word/PDF/Excel/CSV/TXT)" else "🎲 Topic'leri Bul")
    
    if analyze_button:
        if data_source == "📁 Dosya Yükle (Word/PDF/Excel/CSV/TXT)":
            if uploaded_file is None or documents is None:
                st.warning("⚠️ Lütfen önce bir dosya yükleyin!")
            elif len(documents) < 3:
                st.warning("⚠️ Yeterli doküman bulunamadı! En az 3 doküman gerekli.")
            else:
                with st.spinner("Topic'ler bulunuyor..."):
                    try:
                        # NMF
                        nmf_model = NMFTopicModeler(n_topics=n_topics, max_iter=200)
                        nmf_model.fit(documents, max_features=max_features, min_df=2, max_df=0.95)
                        
                        # Topic'ler
                        topics = nmf_model.get_topics(n_words=10)
                        coherence = nmf_model.get_topic_coherence()
                        
                        # Sonuçları göster
                        st.subheader("📊 Analiz Sonuçları")
                        col1, col2, col3 = st.columns(3)
                        with col1:
                            st.metric("Topic Tutarlılığı", f"{coherence:.4f}")
                        with col2:
                            st.metric("Toplam Doküman", len(documents))
                        with col3:
                            st.metric("Bulunan Topic Sayısı", n_topics)
                        
                        # Topic'leri göster
                        st.subheader("🎯 Topic'ler ve Anahtar Kelimeler")
                        for topic_name, words_scores in topics.items():
                            with st.expander(f"{topic_name} - {', '.join([w[0] for w in words_scores[:5]])}"):
                                words_df = pd.DataFrame(words_scores, columns=['Kelime', 'Önem Skoru'])
                                st.dataframe(words_df, width='stretch', use_container_width=True)
                        
                        # Görselleştirme
                        st.subheader("📊 Topic Kelime Analizi")
                        with st.expander("ℹ️ Topic Kelime Grafikleri Ne Gösteriyor?", expanded=False):
                            st.markdown("""
                            **Topic Kelime Grafikleri**:
                            - Her topic için en önemli kelimeleri gösterir
                            - **Y Ekseni**: Kelimeler (önem sırasına göre)
                            - **X Ekseni**: Önem skoru (topic içindeki ağırlık)
                            - **Yüksek skor**: Kelime, topic'i tanımlamada çok önemli
                            - **Düşük skor**: Kelime daha az önemli
                            
                            **Yorumlama**:
                            - İyi topic: Tutarlı, anlamlı kelimeler
                            - Kötü topic: Rastgele, tutarsız kelimeler
                            - Topic tutarlılığı: Tüm topic'lerin ortalama kalitesi
                            """)
                        
                        fig = plot_topic_words(topics, n_words=10)
                        st.pyplot(fig)
                        
                        # Doküman-topic dağılımı
                        doc_topics = nmf_model.get_document_topics()
                        
                        st.subheader("🔥 Doküman-Topic Dağılımı Heatmap")
                        with st.expander("ℹ️ Heatmap Ne Gösteriyor?", expanded=False):
                            st.markdown("""
                            **Doküman-Topic Dağılımı Heatmap**:
                            - **Satırlar**: Topic'ler
                            - **Sütunlar**: Dokümanlar
                            - **Renkler**: Dokümanın topic'e ait olma skoru
                              - **Koyu renk (düşük)**: Doküman bu topic'e az ait
                              - **Açık renk (yüksek)**: Doküman bu topic'e çok ait
                            
                            **Yorumlama**:
                            - Her doküman genellikle 1-2 dominant topic'e sahiptir
                            - Koyu sütunlar: Belirsiz dokümanlar (birden fazla topic)
                            - Açık sütunlar: Net topic'lere sahip dokümanlar
                            - Dikey çizgiler: Aynı topic'teki doküman grupları
                            """)
                        
                        # Heatmap için maksimum doküman sayısı
                        max_docs_heatmap = min(50, len(documents))
                        fig2, ax = plt.subplots(figsize=(12, 8))
                        sns.heatmap(doc_topics[:max_docs_heatmap].T, cmap='YlOrRd', ax=ax, 
                                   yticklabels=[f'Topic {i+1}' for i in range(n_topics)],
                                   xticklabels=[f'Doc {i+1}' for i in range(max_docs_heatmap)])
                        ax.set_xlabel('Dokümanlar', fontsize=12)
                        ax.set_ylabel('Topic\'ler', fontsize=12)
                        ax.set_title(f'Doküman-Topic Dağılımı (İlk {max_docs_heatmap} Doküman)', fontsize=14, fontweight='bold')
                        st.pyplot(fig2)
                        
                        # Her doküman için dominant topic
                        st.subheader("📋 Doküman-Topic Eşleştirmesi")
                        dominant_topics = np.argmax(doc_topics, axis=1)
                        topic_scores = np.max(doc_topics, axis=1)
                        
                        doc_topic_df = pd.DataFrame({
                            'Doküman': [f'Doküman {i+1}' for i in range(len(documents))],
                            'Dominant Topic': [f'Topic {t+1}' for t in dominant_topics],
                            'Topic Skoru': topic_scores,
                            'Önizleme': [doc[:100] + "..." if len(doc) > 100 else doc for doc in documents]
                        })
                        
                        st.dataframe(doc_topic_df, width='stretch', use_container_width=True, height=400)
                        
                        # Yeni doküman tahmini
                        st.subheader("🔮 Yeni Doküman için Topic Tahmini")
                        new_doc = st.text_area("Doküman metni girin:", 
                                               value="",
                                               height=100,
                                               placeholder="Analiz etmek istediğiniz metni buraya yapıştırın...")
                        if new_doc and new_doc.strip():
                            top_topic, all_scores = nmf_model.predict_topic(new_doc)
                            st.write(f"**En uygun topic:** Topic {top_topic+1} (skor: {all_scores[top_topic]:.3f})")
                            
                            scores_df = pd.DataFrame({
                                'Topic': [f'Topic {i+1}' for i in range(n_topics)],
                                'Skor': all_scores
                            })
                            st.bar_chart(scores_df.set_index('Topic'))
                        
                        st.success("✅ Analiz tamamlandı!")
                        
                    except Exception as e:
                        st.error(f"❌ Analiz sırasında hata oluştu: {str(e)}")
                        import traceback
                        with st.expander("🔍 Detaylı Hata Mesajı"):
                            st.code(traceback.format_exc())
        
        else:
            # Örnek metinler
            with st.spinner("Topic'ler bulunuyor..."):
                # Korpus oluştur
                documents = generate_text_corpus(n_documents=n_documents)
                
                # NMF
                nmf_model = NMFTopicModeler(n_topics=n_topics, max_iter=200)
                nmf_model.fit(documents, max_features=max_features, min_df=2, max_df=0.95)
                
                # Topic'ler
                topics = nmf_model.get_topics(n_words=10)
                coherence = nmf_model.get_topic_coherence()
                
                st.metric("Topic Tutarlılığı", f"{coherence:.4f}")
                
                # Topic'leri göster
                st.subheader("Topic'ler ve Anahtar Kelimeler")
                for topic_name, words_scores in topics.items():
                    with st.expander(topic_name):
                        words_df = pd.DataFrame(words_scores, columns=['Kelime', 'Skor'])
                        st.dataframe(words_df, width='stretch')
                
                # Görselleştirme
                st.subheader("📊 Topic Kelime Analizi")
                with st.expander("ℹ️ Topic Kelime Grafikleri Ne Gösteriyor?", expanded=False):
                    st.markdown("""
                    **Topic Kelime Grafikleri**:
                    - Her topic için en önemli kelimeleri gösterir
                    - **Y Ekseni**: Kelimeler (önem sırasına göre)
                    - **X Ekseni**: Önem skoru (topic içindeki ağırlık)
                    - **Yüksek skor**: Kelime, topic'i tanımlamada çok önemli
                    - **Düşük skor**: Kelime daha az önemli
                    
                    **Yorumlama**:
                    - İyi topic: Tutarlı, anlamlı kelimeler
                    - Kötü topic: Rastgele, tutarsız kelimeler
                    - Topic tutarlılığı: Tüm topic'lerin ortalama kalitesi
                    """)
                
                fig = plot_topic_words(topics, n_words=10)
                st.pyplot(fig)
                
                # Doküman-topic dağılımı
                doc_topics = nmf_model.get_document_topics()
                
                st.subheader("🔥 Doküman-Topic Dağılımı Heatmap")
                with st.expander("ℹ️ Heatmap Ne Gösteriyor?", expanded=False):
                    st.markdown("""
                    **Doküman-Topic Dağılımı Heatmap**:
                    - **Satırlar**: Topic'ler
                    - **Sütunlar**: Dokümanlar
                    - **Renkler**: Dokümanın topic'e ait olma skoru
                      - **Koyu renk (düşük)**: Doküman bu topic'e az ait
                      - **Açık renk (yüksek)**: Doküman bu topic'e çok ait
                    
                    **Yorumlama**:
                    - Her doküman genellikle 1-2 dominant topic'e sahiptir
                    - Koyu sütunlar: Belirsiz dokümanlar (birden fazla topic)
                    - Açık sütunlar: Net topic'lere sahip dokümanlar
                    - Dikey çizgiler: Aynı topic'teki doküman grupları
                    """)
                
                fig2, ax = plt.subplots(figsize=(12, 8))
                sns.heatmap(doc_topics[:50].T, cmap='YlOrRd', ax=ax, 
                           yticklabels=[f'Topic {i+1}' for i in range(n_topics)],
                           xticklabels=[f'Doc {i+1}' for i in range(50)])
                ax.set_xlabel('Dokümanlar', fontsize=12)
                ax.set_ylabel('Topic\'ler', fontsize=12)
                ax.set_title('Doküman-Topic Dağılımı (İlk 50 Doküman)', fontsize=14, fontweight='bold')
                st.pyplot(fig2)
                
                # Yeni doküman tahmini
                st.subheader("Yeni Doküman için Topic Tahmini")
                new_doc = st.text_input("Doküman metni girin:", 
                                       value="computer software algorithm data network")
                if new_doc:
                    top_topic, all_scores = nmf_model.predict_topic(new_doc)
                    st.write(f"**En uygun topic:** Topic {top_topic+1} (skor: {all_scores[top_topic]:.3f})")
                    
                    scores_df = pd.DataFrame({
                        'Topic': [f'Topic {i+1}' for i in range(n_topics)],
                        'Skor': all_scores
                    })
                    st.bar_chart(scores_df.set_index('Topic'))


def show_als_recommender():
    """ALS öneri sistemi"""
    st.header("⚡ ALS - Öneri Sistemi")
    
    # Info bölümü
    with st.expander("ℹ️ ALS (Alternating Least Squares) Hakkında Bilgi", expanded=False):
        st.markdown("""
        ### ALS Nedir?
        
        **ALS**, büyük ölçekli öneri sistemleri için optimize edilmiş bir matris faktörizasyon yöntemidir.
        
        ### Nasıl Çalışır?
        
        1. **Alternatif Optimizasyon**: Kullanıcı ve ürün faktörlerini sırayla optimize eder
        2. **Paralel İşleme**: Her kullanıcı/ürün bağımsız işlenebilir (Spark uyumlu)
        3. **Regularizasyon**: Aşırı öğrenmeyi (overfitting) önler
        4. **Iteratif Güncelleme**: Her iterasyonda faktörleri iyileştirir
        
        ### SVD'den Farkları
        
        - **Paralel Çalışma**: Büyük veri setlerinde daha hızlı
        - **Eksik Veri**: Sparse matrislerde daha iyi performans
        - **Ölçeklenebilirlik**: Milyonlarca kullanıcı/ürün ile çalışabilir
        - **Regularizasyon**: Daha iyi genelleme
        
        ### Parametreler
        
        - **Latent Faktör Sayısı**: Gizli özellik sayısı (daha fazla = daha detaylı)
        - **Regularizasyon**: Aşırı öğrenmeyi önler (0.01-1.0 arası)
        - **İterasyon Sayısı**: Eğitim iterasyonu (daha fazla = daha iyi ama yavaş)
        
        ### Kullanım Alanları
        
        - Netflix, Amazon gibi büyük ölçekli öneri sistemleri
        - Spark, Hadoop gibi dağıtık sistemler
        - Gerçek zamanlı öneriler
        """)
    
    st.markdown("""
    ALS (Alternating Least Squares) kullanarak büyük ölçekli öneri sistemi.
    """)
    
    # Veri yükleme seçeneği
    data_source = st.radio(
        "Veri Kaynağı",
        ["📊 Örnek Veri Oluştur", "📁 Dosyadan Yükle"],
        horizontal=True,
        key="als_data_source"
    )
    
    # Dosya formatı seçimi (sadece dosya yükleme seçildiyse)
    file_format = None
    if data_source == "📁 Dosyadan Yükle":
        file_format = get_file_format_selector("als", include_image=False)
    
    # Session state ile veriyi koru
    if 'als_rating_matrix' not in st.session_state:
        st.session_state.als_rating_matrix = None
        st.session_state.als_user_mapping = None
        st.session_state.als_item_mapping = None
    
    rating_matrix = st.session_state.als_rating_matrix
    user_mapping = st.session_state.als_user_mapping
    item_mapping = st.session_state.als_item_mapping
    n_users = None
    n_items = None
    n_factors = 20  # Varsayılan değer
    
    if data_source == "📁 Dosyadan Yükle":
        st.markdown("### 📁 Veri Dosyası Yükle")
        st.info("""
        **Desteklenen Formatlar:** CSV, Excel (.xlsx, .xls)
        
        **Veri Formatı:** Long Format (user_id, item_id, rating) veya Matrix Format
        """)
        
        file = st.file_uploader(
            "Veri dosyasını seçin",
            type=['csv', 'xlsx', 'xls'],
            help="CSV veya Excel dosyası yükleyin",
            key="als_file"
        )
        
        if file is not None:
            try:
                # Dosya önizlemesi ve format önerisi
                # Dosya stream'i bir kez okununca tükenir, bu yüzden içeriği hafızaya al
                import io
                file_content = file.read()
                file_bytes = io.BytesIO(file_content)
                
                import pandas as pd
                if file.name.endswith('.csv'):
                    # Delimiter tespiti
                    file_bytes.seek(0)
                    first_line = file_bytes.readline().decode('utf-8', errors='ignore')
                    delimiters = [',', ';', '\t', '|']
                    detected_delimiter = ','
                    max_cols = 0
                    for delim in delimiters:
                        cols = first_line.split(delim)
                        if len(cols) > max_cols:
                            max_cols = len(cols)
                            detected_delimiter = delim
                    
                    file_bytes.seek(0)
                    preview_df = pd.read_csv(file_bytes, nrows=5, sep=detected_delimiter, engine='python')
                    # Toplam satır sayısı için dosyayı tekrar oku
                    file_bytes.seek(0)
                    total_df = pd.read_csv(file_bytes, sep=detected_delimiter, engine='python')
                elif file.name.endswith(('.xlsx', '.xls')):
                    file_bytes.seek(0)
                    preview_df = pd.read_excel(file_bytes, nrows=5)
                    # Toplam satır sayısı için dosyayı tekrar oku
                    file_bytes.seek(0)
                    total_df = pd.read_excel(file_bytes)
                else:
                    preview_df = None
                    total_df = None
                
                if preview_df is not None:
                    with st.expander("👁️ Dosya Önizleme (İlk 5 Satır)", expanded=False):
                        st.dataframe(preview_df, width='stretch')
                        st.info(f"""
                        **Dosya Bilgileri:**
                        - **Satır Sayısı**: {len(total_df) if total_df is not None else 'Bilinmiyor'} (tahmini)
                        - **Sütun Sayısı**: {len(preview_df.columns)}
                        - **Sütun İsimleri**: {', '.join(preview_df.columns.tolist()[:10])}{'...' if len(preview_df.columns) > 10 else ''}
                        
                        **Format Önerisi:**
                        - **3 sütun varsa** → Long Format seçin (user_id, item_id, rating)
                        - **10+ sütun varsa** → Matrix Format seçin (ilk sütun kullanıcı ID, diğerleri ürün ID)
                        """)
                
                # Dosya stream'ini tekrar kullanılabilir hale getir
                file.seek(0)
                
                data_format = st.radio(
                    "Veri Formatı",
                    ["Long Format (user_id, item_id, rating)", "Matrix Format (Rating Matrisi)"],
                    horizontal=True,
                    key="als_data_format"
                )
                
                if data_format == "Long Format (user_id, item_id, rating)":
                    # Manuel sütun seçimi
                    st.markdown("#### 📋 Sütun Seçimi (Opsiyonel)")
                    use_manual_cols = st.checkbox(
                        "Manuel olarak sütun seçmek istiyorum", 
                        key="als_manual_cols",
                        help="Otomatik tespit yanlış çalışıyorsa, bu seçeneği işaretleyin"
                    )
                    
                    user_col_manual = None
                    item_col_manual = None
                    rating_col_manual = None
                    
                    if use_manual_cols and preview_df is not None:
                        col1, col2, col3 = st.columns(3)
                        with col1:
                            user_col_manual = st.selectbox(
                                "Kullanıcı ID Sütunu",
                                options=preview_df.columns.tolist(),
                                key="als_user_col",
                                help="Kullanıcı ID'lerini içeren sütun"
                            )
                        with col2:
                            item_col_manual = st.selectbox(
                                "Ürün/Öğe ID Sütunu",
                                options=preview_df.columns.tolist(),
                                index=min(1, len(preview_df.columns)-1),
                                key="als_item_col",
                                help="Ürün/Öğe ID'lerini içeren sütun"
                            )
                        with col3:
                            # Sayısal sütunları bul
                            numeric_cols = preview_df.select_dtypes(include=[np.number]).columns.tolist()
                            if not numeric_cols:
                                numeric_cols = preview_df.columns.tolist()
                            
                            default_idx = min(2, len(numeric_cols)-1) if numeric_cols else 0
                            rating_col_manual = st.selectbox(
                                "Rating/Puan Sütunu",
                                options=numeric_cols if numeric_cols else preview_df.columns.tolist(),
                                index=default_idx,
                                key="als_rating_col",
                                help="Rating/Puan değerlerini içeren sütun (sayısal olmalı)"
                            )
                        
                        st.info(f"✅ Seçilen: `{user_col_manual}` (kullanıcı) + `{item_col_manual}` (ürün) + `{rating_col_manual}` (rating)")
                    
                    if st.button("📥 Veriyi Yükle", key="als_load_long"):
                        with st.spinner("Veri yükleniyor..."):
                            # Dosya bilgilerini kaydet
                            file_name = file.name
                            file_size = file.size
                            
                            # Session state'ten dosya içeriğini al (önizleme sırasında kaydedilmiş)
                            if 'als_file_content' in st.session_state:
                                # Dosya içeriğini BytesIO'ya çevir
                                import io
                                file_bytes = io.BytesIO(st.session_state.als_file_content)
                                # Dosya objesi gibi davranması için name attribute ekle
                                file_bytes.name = file_name
                                rating_matrix, user_mapping, item_mapping = load_rating_data_from_file(
                                    file_bytes, 
                                    user_col=user_col_manual, 
                                    item_col=item_col_manual, 
                                    rating_col=rating_col_manual
                                )
                            else:
                                # Fallback: dosyayı tekrar oku
                                file.seek(0)
                                rating_matrix, user_mapping, item_mapping = load_rating_data_from_file(
                                    file, 
                                    user_col=user_col_manual, 
                                    item_col=item_col_manual, 
                                    rating_col=rating_col_manual
                                )
                            
                            # Dosya bilgilerini session state'e kaydet
                            st.session_state.als_file_name = file_name
                            st.session_state.als_file_size = file_size
                            
                            # Session state'e kaydet
                            st.session_state.als_rating_matrix = rating_matrix
                            st.session_state.als_user_mapping = user_mapping
                            st.session_state.als_item_mapping = item_mapping
                            
                            # Veri istatistiklerini hesapla
                            from scipy.sparse import issparse
                            if issparse(rating_matrix):
                                n_ratings = rating_matrix.nnz
                                density = rating_matrix.nnz / (rating_matrix.shape[0] * rating_matrix.shape[1]) * 100 if rating_matrix.shape[0] * rating_matrix.shape[1] > 0 else 0
                                ratings_data = rating_matrix.data
                            else:
                                mask = ~np.isnan(rating_matrix)
                                n_ratings = np.sum(mask)
                                density = (1 - np.isnan(rating_matrix).sum() / rating_matrix.size) * 100 if rating_matrix.size > 0 else 0
                                ratings_data = rating_matrix[mask]
                            
                            # Rating istatistikleri - boş array kontrolü
                            if len(ratings_data) == 0:
                                # Eğer hiç rating yoksa varsayılan değerler
                                min_rating = 0.0
                                max_rating = 0.0
                                mean_rating = 0.0
                                median_rating = 0.0
                                st.warning("⚠️ Uyarı: Dosyada hiç rating değeri bulunamadı! Tüm değerler NaN olabilir.")
                            else:
                                min_rating = float(np.min(ratings_data))
                                max_rating = float(np.max(ratings_data))
                                mean_rating = float(np.mean(ratings_data))
                                median_rating = float(np.median(ratings_data))
                            
                            st.success(f"✅ Veri yüklendi! {rating_matrix.shape[0]} kullanıcı, {rating_matrix.shape[1]} ürün")
                            
                            # Detaylı dosya analizi bölümü
                            with st.expander("📋 Dosya Analizi - Kullanılan Veriler", expanded=True):
                                st.markdown("### 📁 Dosya Bilgileri")
                                col1, col2 = st.columns(2)
                                with col1:
                                    st.metric("Dosya Adı", file_name)
                                    st.metric("Dosya Boyutu", f"{file_size / 1024:.2f} KB")
                                with col2:
                                    st.metric("Veri Formatı", "Long Format")
                                    st.metric("Toplam Satır Sayısı", f"{n_ratings:,}")
                                
                                st.markdown("### 📈 Veri İstatistikleri")
                                col1, col2, col3, col4 = st.columns(4)
                                with col1:
                                    st.metric("Kullanıcı Sayısı", f"{rating_matrix.shape[0]:,}")
                                with col2:
                                    st.metric("Ürün Sayısı", f"{rating_matrix.shape[1]:,}")
                                with col3:
                                    st.metric("Toplam Rating", f"{n_ratings:,}")
                                with col4:
                                    st.metric("Veri Yoğunluğu", f"{density:.2f}%")
                                
                                st.markdown("### ⭐ Rating Dağılımı")
                                col1, col2, col3, col4 = st.columns(4)
                                with col1:
                                    st.metric("Minimum Rating", f"{min_rating:.2f}")
                                with col2:
                                    st.metric("Maksimum Rating", f"{max_rating:.2f}")
                                with col3:
                                    st.metric("Ortalama Rating", f"{mean_rating:.2f}")
                                with col4:
                                    st.metric("Medyan Rating", f"{median_rating:.2f}")
                                
                                st.markdown("""
                                **📝 Açıklama:**
                                - **Dosya Bilgileri**: Yüklenen dosyanın adı ve boyutu
                                - **Veri İstatistikleri**: Matris boyutları ve veri yoğunluğu
                                - **Rating Dağılımı**: Rating değerlerinin istatistiksel özeti
                                
                                Bu veriler, ALS (Alternating Least Squares) algoritması ile işlenecek ve 
                                kullanıcı-ürün etkileşimlerinden latent faktörler çıkarılacaktır.
                                """)
                            
                            st.rerun()  # Sayfayı yenile
                else:
                    if st.button("📥 Veriyi Yükle", key="als_load_matrix"):
                        with st.spinner("Veri yükleniyor..."):
                            # Dosya bilgilerini kaydet
                            file_name = file.name
                            file_size = file.size
                            
                            # Session state'ten dosya içeriğini al (önizleme sırasında kaydedilmiş)
                            if 'als_file_content' in st.session_state:
                                # Dosya içeriğini BytesIO'ya çevir
                                import io
                                file_bytes = io.BytesIO(st.session_state.als_file_content)
                                # Dosya objesi gibi davranması için name attribute ekle
                                file_bytes.name = file_name
                                try:
                                    rating_matrix = load_rating_matrix_from_file(file_bytes)
                                except Exception as e:
                                    st.error(f"❌ Hata: {str(e)}")
                                    st.info("""
                                    **💡 Matrix Format için:**
                                    - İlk sütun kullanıcı ID'leri olmalı (index)
                                    - Diğer sütunlar ürün ID'leri olmalı
                                    - Değerler rating'ler olmalı (NaN = eksik veri)
                                    - CSV dosyasında ilk sütun otomatik olarak index olarak okunur
                                    
                                    **Örnek Format:**
                                    ```
                                    user_id,item_1,item_2,item_3,...
                                    1,4.5,3.0,5.0,...
                                    2,2.5,4.0,NaN,...
                                    ```
                                    """)
                                    st.stop()
                            else:
                                # Fallback: dosyayı tekrar oku
                                file.seek(0)
                                try:
                                    rating_matrix = load_rating_matrix_from_file(file)
                                except Exception as e:
                                    st.error(f"❌ Hata: {str(e)}")
                                    st.info("""
                                    **💡 Matrix Format için:**
                                    - İlk sütun kullanıcı ID'leri olmalı (index)
                                    - Diğer sütunlar ürün ID'leri olmalı
                                    - Değerler rating'ler olmalı (NaN = eksik veri)
                                    - CSV dosyasında ilk sütun otomatik olarak index olarak okunur
                                    
                                    **Örnek Format:**
                                    ```
                                    user_id,item_1,item_2,item_3,...
                                    1,4.5,3.0,5.0,...
                                    2,2.5,4.0,NaN,...
                                    ```
                                    """)
                                    st.stop()
                            
                            # Dosya bilgilerini session state'e kaydet
                            st.session_state.als_file_name = file_name
                            st.session_state.als_file_size = file_size
                            
                            # Session state'e kaydet
                            st.session_state.als_rating_matrix = rating_matrix
                            st.session_state.als_user_mapping = None
                            st.session_state.als_item_mapping = None
                            
                            # Veri istatistiklerini hesapla
                            from scipy.sparse import issparse
                            if issparse(rating_matrix):
                                n_ratings = rating_matrix.nnz
                                density = rating_matrix.nnz / (rating_matrix.shape[0] * rating_matrix.shape[1]) * 100 if rating_matrix.shape[0] * rating_matrix.shape[1] > 0 else 0
                                ratings_data = rating_matrix.data
                            else:
                                mask = ~np.isnan(rating_matrix)
                                n_ratings = np.sum(mask)
                                density = (1 - np.isnan(rating_matrix).sum() / rating_matrix.size) * 100 if rating_matrix.size > 0 else 0
                                ratings_data = rating_matrix[mask]
                            
                            # Rating istatistikleri - boş array kontrolü
                            if len(ratings_data) == 0:
                                # Eğer hiç rating yoksa varsayılan değerler
                                min_rating = 0.0
                                max_rating = 0.0
                                mean_rating = 0.0
                                median_rating = 0.0
                                st.warning("⚠️ Uyarı: Dosyada hiç rating değeri bulunamadı! Tüm değerler NaN olabilir.")
                            else:
                                min_rating = float(np.min(ratings_data))
                                max_rating = float(np.max(ratings_data))
                                mean_rating = float(np.mean(ratings_data))
                                median_rating = float(np.median(ratings_data))
                            
                            st.success(f"✅ Veri yüklendi! {rating_matrix.shape[0]} kullanıcı, {rating_matrix.shape[1]} ürün")
                            
                            # Detaylı dosya analizi bölümü
                            with st.expander("📋 Dosya Analizi - Kullanılan Veriler", expanded=True):
                                st.markdown("### 📁 Dosya Bilgileri")
                                col1, col2 = st.columns(2)
                                with col1:
                                    st.metric("Dosya Adı", file_name)
                                    st.metric("Dosya Boyutu", f"{file_size / 1024:.2f} KB")
                                with col2:
                                    st.metric("Veri Formatı", "Matrix Format")
                                    st.metric("Toplam Rating", f"{n_ratings:,}")
                                
                                st.markdown("### 📈 Veri İstatistikleri")
                                col1, col2, col3, col4 = st.columns(4)
                                with col1:
                                    st.metric("Kullanıcı Sayısı", f"{rating_matrix.shape[0]:,}")
                                with col2:
                                    st.metric("Ürün Sayısı", f"{rating_matrix.shape[1]:,}")
                                with col3:
                                    st.metric("Toplam Rating", f"{n_ratings:,}")
                                with col4:
                                    st.metric("Veri Yoğunluğu", f"{density:.2f}%")
                                
                                st.markdown("### ⭐ Rating Dağılımı")
                                col1, col2, col3, col4 = st.columns(4)
                                with col1:
                                    st.metric("Minimum Rating", f"{min_rating:.2f}")
                                with col2:
                                    st.metric("Maksimum Rating", f"{max_rating:.2f}")
                                with col3:
                                    st.metric("Ortalama Rating", f"{mean_rating:.2f}")
                                with col4:
                                    st.metric("Medyan Rating", f"{median_rating:.2f}")
                                
                                st.markdown("""
                                **📝 Açıklama:**
                                - **Dosya Bilgileri**: Yüklenen dosyanın adı ve boyutu
                                - **Veri İstatistikleri**: Matris boyutları ve veri yoğunluğu
                                - **Rating Dağılımı**: Rating değerlerinin istatistiksel özeti
                                
                                Bu veriler, ALS (Alternating Least Squares) algoritması ile işlenecek ve 
                                kullanıcı-ürün etkileşimlerinden latent faktörler çıkarılacaktır.
                                """)
                            
                            st.rerun()  # Sayfayı yenile
                            
            except Exception as e:
                st.error(f"❌ Hata: {str(e)}")
                st.info("💡 Lütfen veri formatını kontrol edin.")
    else:
        # Örnek veri oluştur - session state'i temizle
        if st.session_state.als_rating_matrix is not None:
            st.session_state.als_rating_matrix = None
            st.session_state.als_user_mapping = None
            st.session_state.als_item_mapping = None
        
        col1, col2, col3 = st.columns(3)
        with col1:
            n_users = st.slider("Kullanıcı Sayısı", 50, 500, 100)
        with col2:
            n_items = st.slider("Ürün Sayısı", 30, 200, 50)
        with col3:
            n_factors = st.slider("Latent Faktör Sayısı", 5, 50, 20)
    
    # Model parametreleri (veri yüklendiyse)
    if rating_matrix is not None:
        n_users, n_items = rating_matrix.shape
        
        # Veri seti boş mu kontrol et
        if n_users == 0 or n_items == 0:
            st.error("❌ Hata: Veri seti boş! Lütfen geçerli bir veri dosyası yükleyin.")
            n_factors = 5  # Varsayılan değer
        else:
            max_factors = min(50, min(n_users, n_items))
            min_factors = min(5, max_factors)  # min_value max_value'dan küçük olmalı
            default_factors = min(20, max_factors)
            
            # Eğer max_factors çok küçükse veya 0 ise, slider yerine sabit değer kullan
            if max_factors <= 0 or max_factors < min_factors:
                n_factors = max(1, max_factors)  # En az 1 faktör
                if max_factors <= 0:
                    st.error("❌ Hata: Veri seti çok küçük! Latent faktör sayısı ayarlanamıyor.")
                else:
                    st.info(f"⚠️ Veri seti küçük olduğu için latent faktör sayısı otomatik olarak {n_factors} olarak ayarlandı.")
            else:
                # Optimal parametreleri al
                optimal_params = get_optimal_model_params("als", data_shape=(n_users, n_items))
                optimal_n_factors = optimal_params['n_factors']
                suggested_factors = max(min_factors, min(max_factors, optimal_n_factors))
                
                n_factors = st.slider(
                    "Latent Faktör Sayısı", 
                    min_factors, 
                    max_factors, 
                    suggested_factors,
                    key="als_n_factors_loaded",
                    help=f"Önerilen değer: {optimal_n_factors} (veri boyutuna göre otomatik hesaplandı)"
                )
                if n_factors != optimal_n_factors:
                    st.info(f"💡 Veri boyutunuza göre önerilen değer: {optimal_n_factors}")
    
    # Optimal parametreleri al (veri yüklendiyse)
    if rating_matrix is not None:
        optimal_params = get_optimal_model_params("als", data_shape=rating_matrix.shape)
        optimal_regularization = optimal_params['regularization']
        optimal_iterations = optimal_params['iterations']
    else:
        optimal_regularization = 0.1
        optimal_iterations = 15
    
    col1, col2 = st.columns(2)
    with col1:
        regularization = st.slider(
            "Regularizasyon", 
            0.01, 1.0, 
            optimal_regularization,
            help=f"Önerilen değer: {optimal_regularization}"
        )
    with col2:
        iterations = st.slider(
            "İterasyon Sayısı", 
            5, 30, 
            optimal_iterations,
            help=f"Önerilen değer: {optimal_iterations}"
        )
    
    # Sparsity sadece örnek veri için
    if rating_matrix is None:
        sparsity = st.slider("Eksik Veri Oranı", 0.3, 0.9, 0.7)
    
    if st.button("🚀 Modeli Eğit", key="als_train"):
        if rating_matrix is None and data_source == "📁 Dosyadan Yükle":
            st.warning("⚠️ Lütfen önce veri dosyasını yükleyin!")
        else:
            with st.spinner("Model eğitiliyor (bu biraz zaman alabilir)..."):
                # Veri yoksa oluştur
                if rating_matrix is None:
                    rating_matrix = generate_rating_matrix(
                        n_users=n_users, 
                        n_items=n_items, 
                        sparsity=sparsity
                    )
            
            # Train-test split (sparse matrix desteği ile)
            from scipy.sparse import issparse
            
            np.random.seed(42)
            
            if issparse(rating_matrix):
                # Sparse matrix için
                rows, cols = rating_matrix.nonzero()
                n_ratings = len(rows)
                test_size = min(int(0.2 * n_ratings), 10000)
                test_sample_indices = np.random.choice(n_ratings, size=test_size, replace=False)
                
                test_rows = rows[test_sample_indices]
                test_cols = cols[test_sample_indices]
                # Sparse matrix'ten değerleri al - matrix objesi için np.array kullan
                test_matrix_slice = rating_matrix[test_rows, test_cols]
                # matrix objesi için A property veya np.array kullan
                if hasattr(test_matrix_slice, 'A'):
                    test_values = test_matrix_slice.A.flatten()
                elif hasattr(test_matrix_slice, 'toarray'):
                    test_values = test_matrix_slice.toarray().flatten()
                else:
                    test_values = np.array(test_matrix_slice).flatten()
                
                train_matrix = rating_matrix.copy()
                train_matrix[test_rows, test_cols] = 0
                train_matrix.eliminate_zeros()
                
                test_indices = (test_rows, test_cols)
            else:
                # Dense matrix için
                mask = ~np.isnan(rating_matrix)
                n_ratings = np.sum(mask)
                test_size = min(int(0.2 * n_ratings), 10000)
                
                valid_indices = np.where(mask)
                test_sample_indices = np.random.choice(
                    len(valid_indices[0]), 
                    size=test_size, 
                    replace=False
                )
                
                test_mask = np.zeros_like(mask, dtype=bool)
                test_mask[valid_indices[0][test_sample_indices], valid_indices[1][test_sample_indices]] = True
                
                train_matrix = rating_matrix.copy()
                train_matrix[test_mask] = np.nan
                
                test_values = rating_matrix[test_mask]
                test_indices = np.where(test_mask)
            
            # Model eğit
            import time
            training_start = time.time()
            
            als_model = ALSRecommender(
                n_factors=n_factors, 
                regularization=regularization, 
                iterations=iterations
            )
            als_model.fit(train_matrix, implicit=False)
            
            training_time = time.time() - training_start
            
            # Test matrisi oluştur (evaluate için)
            if issparse(rating_matrix):
                # Sparse matrix için test matrisi oluştur
                from scipy.sparse import csr_matrix
                test_matrix = csr_matrix((test_values, (test_indices[0], test_indices[1])), 
                                        shape=rating_matrix.shape)
            else:
                # Dense matrix için test matrisi oluştur
                test_matrix = np.full_like(rating_matrix, np.nan)
                test_matrix[test_indices[0], test_indices[1]] = test_values
            
            # Değerlendirme
            test_predictions = []
            batch_size = 1000
            for i in range(0, len(test_indices[0]), batch_size):
                batch_end = min(i + batch_size, len(test_indices[0]))
                batch_users = test_indices[0][i:batch_end]
                batch_items = test_indices[1][i:batch_end]
                batch_preds = als_model.predict_all()[batch_users, batch_items]
                test_predictions.extend(batch_preds)
            
            test_predictions = np.array(test_predictions)
            rmse = np.sqrt(mean_squared_error(test_values, test_predictions))
            
            # Sonuçlar
            st.subheader("📊 Model Sonuçları")
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Test RMSE", f"{rmse:.4f}")
            with col2:
                st.metric("Eğitim İterasyonu", iterations)
            with col3:
                st.metric("Eğitim Süresi", f"{training_time:.2f} saniye")
            
            # Detaylı sonuç açıklaması
            with st.expander("📝 Sonuç Açıklaması - Ne Elde Edildi?", expanded=True):
                st.markdown("### 🔍 Kullanılan Veriler")
                if data_source == "📁 Dosyadan Yükle" and 'als_file_name' in st.session_state:
                    st.info(f"""
                    **Dosya**: {st.session_state.als_file_name} ({st.session_state.als_file_size / 1024:.2f} KB)
                    - **Kullanıcı Sayısı**: {n_users:,}
                    - **Ürün Sayısı**: {n_items:,}
                    - **Toplam Rating**: {n_ratings:,}
                    """)
                else:
                    st.info(f"""
                    **Örnek Veri**:
                    - **Kullanıcı Sayısı**: {n_users:,}
                    - **Ürün Sayısı**: {n_items:,}
                    """)
                
                st.markdown("### ⚙️ Model Parametreleri")
                st.info(f"""
                - **Latent Faktör Sayısı (n_factors)**: {n_factors}
                - **Regularizasyon**: {regularization}
                - **İterasyon Sayısı**: {iterations}
                - **Eğitim Verisi**: {n_ratings - len(test_values):,} rating
                - **Test Verisi**: {len(test_values):,} rating
                """)
                
                st.markdown("### 📈 Elde Edilen Sonuçlar")
                st.success(f"""
                **Model Başarıyla Eğitildi!**
                
                1. **Test RMSE**: {rmse:.4f}
                   - Bu değer ne kadar düşükse, model o kadar iyi tahmin yapıyor demektir
                   - RMSE, gerçek rating'ler ile tahmin edilen rating'ler arasındaki ortalama hata miktarını gösterir
                   - Örnek: RMSE = {rmse:.4f} → Ortalama {rmse:.2f} puanlık hata var
                
                2. **Eğitim İterasyonu**: {iterations}
                   - Model, {iterations} iterasyon boyunca kullanıcı ve ürün faktörlerini optimize etti
                   - Her iterasyonda faktörler daha iyi hale geldi
                
                3. **Latent Faktör Sayısı**: {n_factors}
                   - Model, veriyi {n_factors} boyutlu latent faktör uzayına indirgedi
                   - Her faktör, kullanıcı ve ürün özelliklerini temsil eden bir boyuttur
                
                4. **Eğitim Süresi**: {training_time:.2f} saniye
                   - Model eğitimi {training_time:.2f} saniyede tamamlandı
                """)
                
                st.markdown("### 🎯 Ne Yapıldı?")
                st.markdown(f"""
                **ALS (Alternating Least Squares) Algoritması** şu adımları izledi:
                
                1. **Veri Hazırlama**: Rating matrisi train ve test setlerine ayrıldı
                2. **Faktör Başlatma**: {n_factors} boyutlu kullanıcı ve ürün faktör matrisleri rastgele başlatıldı
                3. **Alternatif Optimizasyon**: {iterations} iterasyon boyunca:
                   - **Kullanıcı faktörleri sabitken**, ürün faktörleri optimize edildi
                   - **Ürün faktörleri sabitken**, kullanıcı faktörleri optimize edildi
                   - Her iterasyonda hata azaltıldı
                4. **Regularizasyon**: {regularization} değeri ile aşırı öğrenme (overfitting) önlendi
                5. **Tahmin**: Eksik rating'ler, latent faktörler kullanılarak tahmin edildi
                6. **Değerlendirme**: Test seti üzerinde RMSE hesaplandı
                
                **Sonuç**: Model, kullanıcı-ürün etkileşimlerinden öğrendiği kalıpları kullanarak 
                yeni rating'leri tahmin edebiliyor. ALS, SVD'den farklı olarak:
                - **Paralel çalışabilir**: Her kullanıcı/ürün bağımsız işlenebilir
                - **Büyük veri setlerinde daha hızlı**: Sparse matrislerde optimize edilmiştir
                - **Regularizasyon ile daha iyi genelleme**: Aşırı öğrenmeyi önler
                
                Bu sayede kullanıcılara henüz görmedikleri ürünler için kişiselleştirilmiş öneriler sunulabilir.
                """)
                
                st.markdown("### 💡 Sonraki Adımlar")
                st.info("""
                - **Öneriler**: Aşağıdaki "Kullanıcı Önerileri" bölümünden belirli bir kullanıcı için öneriler görebilirsiniz
                - **Benzer Ürünler**: Item similarity bölümünden benzer ürünleri inceleyebilirsiniz
                - **Parametre Ayarı**: Regularizasyon ve iterasyon sayısını değiştirerek model performansını artırabilirsiniz
                """)
            
            # Öneriler
            st.subheader("🎯 Kullanıcı Önerileri")
            with st.expander("ℹ️ ALS Önerileri Nasıl Çalışıyor?", expanded=False):
                st.markdown("""
                **ALS Öneri Sistemi**:
                - Kullanıcının latent faktörlerini kullanır
                - Tüm ürünler için rating tahmin eder
                - En yüksek tahmin edilen rating'lere sahip ürünleri önerir
                - Zaten rating verilen ürünleri hariç tutar
                
                **Tahmin Edilen Rating**: Modelin, kullanıcının bu ürüne vereceği rating tahmini
                - Yüksek değer: Kullanıcının beğenme olasılığı yüksek
                - Düşük değer: Kullanıcının beğenme olasılığı düşük
                """)
            
            user_idx = st.selectbox("Kullanıcı Seçin", range(min(10, n_users)))
            
            recommendations = als_model.recommend(
                user_idx, 
                n_recommendations=10, 
                exclude_rated=True, 
                rating_matrix=rating_matrix
            )
            
            recommendations_df = pd.DataFrame({
                'Ürün ID': recommendations[0] + 1,
                'Tahmin Edilen Rating': recommendations[1]
            })
            st.dataframe(recommendations_df, width='stretch')
            
            # Benzer item'lar
            st.subheader("🔗 Benzer Ürünler (Item Similarity)")
            with st.expander("ℹ️ Benzerlik Nasıl Hesaplanıyor?", expanded=False):
                st.markdown("""
                **Cosine Similarity**:
                - İki ürünün latent faktör vektörleri arasındaki açıyı ölçer
                - **1.0**: Tamamen benzer (aynı yönde)
                - **0.0**: Ortogonal (ilişkisiz)
                - **-1.0**: Tamamen zıt
                
                **Kullanım**:
                - "Bu ürüne bakanlar şunlara da baktı" özelliği
                - Ürün kategorilendirme
                - Cross-selling önerileri
                """)
            
            col1, col2 = st.columns([2, 1])
            with col1:
                item_idx = st.selectbox("Ürün Seçin", range(min(10, n_items)), key="als_item")
            with col2:
                n_similar = st.number_input("Gösterilecek Benzer Ürün Sayısı", min_value=5, max_value=50, value=10, step=5, key="als_n_similar")
            
            if st.button("🔍 Benzer Ürünleri Göster", key="als_show_similar", type="primary"):
                with st.spinner("Benzer ürünler hesaplanıyor..."):
                    similar_items = als_model.get_similar_items(item_idx, n_similar=n_similar)
                    
                    # Benzer ürünler tablosu
                    similar_df = pd.DataFrame({
                        'Benzer Ürün ID': similar_items[0] + 1,
                        'Benzerlik Skoru': [f"{score:.4f}" for score in similar_items[1]]
                    })
                    
                    st.markdown(f"### 📊 Ürün {item_idx + 1} ile Benzer {len(similar_items[0])} Ürün")
                    st.dataframe(similar_df, width='stretch')
                    
                    # Görselleştirme
                    fig, ax = plt.subplots(figsize=(12, 6))
                    bars = ax.barh(range(len(similar_items[1])), similar_items[1], alpha=0.7, 
                                  color=plt.cm.viridis(similar_items[1]))
                    ax.set_yticks(range(len(similar_items[1])))
                    ax.set_yticklabels([f"Ürün {idx + 1}" for idx in similar_items[0]])
                    ax.set_xlabel('Benzerlik Skoru (Cosine Similarity)', fontsize=12)
                    ax.set_title(f'Ürün {item_idx + 1} ile En Benzer {n_similar} Ürün', fontsize=14, fontweight='bold')
                    ax.grid(True, alpha=0.3, axis='x')
                    ax.set_xlim(0, 1.1)
                    
                    # Değerleri çubukların üzerine yaz
                    for i, (idx, score) in enumerate(zip(similar_items[0], similar_items[1])):
                        ax.text(score + 0.01, i, f'{score:.3f}', va='center', fontsize=9)
                    
                    plt.tight_layout()
                    st.pyplot(fig)
                    
                    # İstatistikler
                    with st.expander("📈 Benzerlik İstatistikleri", expanded=False):
                        col1, col2, col3, col4 = st.columns(4)
                        with col1:
                            st.metric("Ortalama Benzerlik", f"{np.mean(similar_items[1]):.4f}")
                        with col2:
                            st.metric("Maksimum Benzerlik", f"{np.max(similar_items[1]):.4f}")
                        with col3:
                            st.metric("Minimum Benzerlik", f"{np.min(similar_items[1]):.4f}")
                        with col4:
                            st.metric("Standart Sapma", f"{np.std(similar_items[1]):.4f}")
                        
                        st.info(f"""
                        **Yorumlama:**
                        - **Yüksek benzerlik (>0.7)**: Ürünler çok benzer, aynı kategoride olabilir
                        - **Orta benzerlik (0.3-0.7)**: Ürünler benzer özelliklere sahip
                        - **Düşük benzerlik (<0.3)**: Ürünler farklı kategorilerde olabilir
                        
                        **Ürün {item_idx + 1}** için en benzer ürün: **Ürün {similar_items[0][0] + 1}** (Benzerlik: {similar_items[1][0]:.4f})
                        """)


def show_performance_comparison():
    """Performans karşılaştırması"""
    st.header("📊 Performans Karşılaştırması")
    
    # Info bölümü
    with st.expander("ℹ️ Performans Karşılaştırması Hakkında", expanded=False):
        st.markdown("""
        ### Karşılaştırılan Algoritmalar
        
        **SVD (Singular Value Decomposition)**:
        - ✅ Matematiksel olarak en kesin
        - ✅ Hızlı eğitim
        - ❌ Büyük veri setlerinde yavaş
        - ❌ Paralel çalışmaya uygun değil
        
        **ALS (Alternating Least Squares)**:
        - ✅ Büyük ölçekli veri setlerinde hızlı
        - ✅ Paralel çalışmaya uygun (Spark)
        - ✅ Sparse matrislerde iyi
        - ❌ Daha fazla iterasyon gerekir
        
        ### Metrikler
        
        **RMSE (Root Mean Square Error)**:
        - Tahmin hatasının ölçüsü
        - **Düşük RMSE = Daha iyi performans**
        - Formül: √(Σ(tahmin - gerçek)² / n)
        
        **Eğitim Süresi**:
        - Modelin eğitilmesi için geçen süre
        - Büyük veri setlerinde önemli
        """)
    
    st.markdown("""
    Farklı algoritmaların performans karşılaştırması.
    """)
    
    col1, col2 = st.columns(2)
    with col1:
        n_users = st.slider("Kullanıcı Sayısı", 50, 300, 100)
    with col2:
        n_items = st.slider("Ürün Sayısı", 30, 150, 50)
    
    if st.button("Karşılaştır"):
        with st.spinner("Algoritmalar karşılaştırılıyor..."):
            # Veri oluştur
            rating_matrix = generate_rating_matrix(
                n_users=n_users, 
                n_items=n_items, 
                sparsity=0.7
            )
            
            # Train-test split (sparse matrix desteği ile)
            from scipy.sparse import issparse
            
            np.random.seed(42)
            
            if issparse(rating_matrix):
                # Sparse matrix için
                rows, cols = rating_matrix.nonzero()
                n_ratings = len(rows)
                test_size = min(int(0.2 * n_ratings), 10000)
                test_sample_indices = np.random.choice(n_ratings, size=test_size, replace=False)
                
                test_rows = rows[test_sample_indices]
                test_cols = cols[test_sample_indices]
                # Sparse matrix'ten değerleri al - matrix objesi için np.array kullan
                test_matrix_slice = rating_matrix[test_rows, test_cols]
                # matrix objesi için A property veya np.array kullan
                if hasattr(test_matrix_slice, 'A'):
                    test_values = test_matrix_slice.A.flatten()
                elif hasattr(test_matrix_slice, 'toarray'):
                    test_values = test_matrix_slice.toarray().flatten()
                else:
                    test_values = np.array(test_matrix_slice).flatten()
                
                train_matrix = rating_matrix.copy()
                train_matrix[test_rows, test_cols] = 0
                train_matrix.eliminate_zeros()
                
                # Test matrix oluştur (dense format, evaluate için)
                test_matrix = np.full_like(rating_matrix.toarray(), np.nan)
                test_matrix[test_rows, test_cols] = test_values
                
                test_indices = (test_rows, test_cols)
            else:
                # Dense matrix için
                mask = ~np.isnan(rating_matrix)
                n_ratings = np.sum(mask)
                test_size = min(int(0.2 * n_ratings), 10000)
                
                valid_indices = np.where(mask)
                test_sample_indices = np.random.choice(
                    len(valid_indices[0]), 
                    size=test_size, 
                    replace=False
                )
                
                test_mask = np.zeros_like(mask, dtype=bool)
                test_mask[valid_indices[0][test_sample_indices], valid_indices[1][test_sample_indices]] = True
                
                train_matrix = rating_matrix.copy()
                train_matrix[test_mask] = np.nan
                
                # Test matrix oluştur (evaluate için)
                test_matrix = np.full_like(rating_matrix, np.nan)
                test_matrix[test_mask] = rating_matrix[test_mask]
                
                test_values = rating_matrix[test_mask]
                test_indices = np.where(test_mask)
            
            results = {}
            
            # SVD
            with st.spinner("SVD eğitiliyor..."):
                svd_model = SVDRecommender(n_components=20)
                svd_model.fit(train_matrix)
                svd_rmse = svd_model.evaluate(test_matrix)
                results['SVD'] = svd_rmse
            
            # ALS
            with st.spinner("ALS eğitiliyor..."):
                als_model = ALSRecommender(n_factors=20, regularization=0.1, iterations=15)
                als_model.fit(train_matrix)
                als_rmse = als_model.evaluate(test_matrix)
                results['ALS'] = als_rmse
            
            # Sonuçlar
            results_df = pd.DataFrame({
                'Algoritma': list(results.keys()),
                'RMSE': list(results.values())
            })
            
            st.subheader("📊 RMSE Karşılaştırması")
            with st.expander("ℹ️ Bu Karşılaştırma Ne Gösteriyor?", expanded=False):
                st.markdown("""
                **RMSE Karşılaştırması**:
                - **RMSE (Root Mean Square Error)**: Tahmin hatasının ölçüsü
                - **Düşük RMSE**: Daha iyi tahmin, daha az hata
                - **Yüksek RMSE**: Daha kötü tahmin, daha fazla hata
                
                **Yorumlama**:
                - En düşük RMSE'ye sahip algoritma en iyi performansı gösterir
                - Fark küçükse: Algoritmalar benzer performans
                - Fark büyükse: Bir algoritma diğerinden belirgin şekilde daha iyi
                
                **Not**: Bu sonuçlar veri setine ve parametrelere bağlıdır. 
                Farklı veri setlerinde sonuçlar değişebilir.
                """)
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.dataframe(results_df, width='stretch')
            
            with col2:
                fig, ax = plt.subplots(figsize=(8, 5))
                colors = ['#1f77b4' if x == results_df['RMSE'].min() else '#ff7f0e' 
                         for x in results_df['RMSE']]
                ax.bar(results_df['Algoritma'], results_df['RMSE'], alpha=0.7, color=colors)
                ax.set_ylabel('RMSE (Düşük = Daha İyi)', fontsize=12)
                ax.set_title('Algoritma Performans Karşılaştırması', fontsize=14, fontweight='bold')
                ax.grid(True, alpha=0.3, axis='y')
                st.pyplot(fig)
            
            # En iyi algoritma
            best_algorithm = results_df.loc[results_df['RMSE'].idxmin(), 'Algoritma']
            best_rmse = results_df['RMSE'].min()
            worst_rmse = results_df['RMSE'].max()
            improvement = ((worst_rmse - best_rmse) / worst_rmse * 100)
            
            st.success(f"🏆 **En iyi performans**: **{best_algorithm}** (RMSE: {best_rmse:.4f})")
            st.info(f"💡 {best_algorithm}, diğer algoritmaya göre %{improvement:.1f} daha iyi performans gösteriyor.")


# ==================== MODERN ALGORITHM PAGES ====================

def show_ncf_recommender():
    """NCF (Neural Collaborative Filtering) sayfası"""
    st.header("🧠 Neural Collaborative Filtering (NCF)")
    
    with st.expander("ℹ️ NCF Nedir?", expanded=False):
        st.markdown("""
        **Neural Collaborative Filtering (NCF)**, SVD'nin modern Deep Learning versiyonudur.
        
        **Farkları:**
        - ✅ SVD sadece **doğrusal** (linear) ilişkileri yakalar
        - ✅ NCF **doğrusal olmayan** (non-linear) ilişkileri öğrenir
        - ✅ Kullanıcı davranışlarındaki karmaşık pattern'leri yakalar
        - ✅ Embedding + Multi-Layer Perceptron (MLP) kullanır
        
        **Kullanım Alanları:**
        - Netflix, YouTube gibi modern öneri sistemleri
        - Büyük ölçekli e-ticaret platformları
        """)
    
    # Veri yükleme seçeneği
    data_source = st.radio(
        "Veri Kaynağı",
        ["📊 Örnek Veri Oluştur", "📁 Dosyadan Yükle"],
        horizontal=True,
        key="ncf_data_source"
    )
    
    # Session state ile veriyi koru
    if 'ncf_rating_matrix' not in st.session_state:
        st.session_state.ncf_rating_matrix = None
        st.session_state.ncf_user_mapping = None
        st.session_state.ncf_item_mapping = None
    
    rating_matrix = st.session_state.ncf_rating_matrix
    user_mapping = st.session_state.ncf_user_mapping
    item_mapping = st.session_state.ncf_item_mapping
    n_users = None
    n_items = None
    
    if data_source == "📁 Dosyadan Yükle":
        st.markdown("### 📁 Veri Dosyası Yükle")
        st.info("""
        **Desteklenen Formatlar:** CSV, Excel (.xlsx, .xls)
        
        **Veri Formatı:** Long Format (user_id, item_id, rating) veya Matrix Format
        """)
        
        file = st.file_uploader(
            "Veri dosyasını seçin",
            type=['csv', 'xlsx', 'xls'],
            help="CSV veya Excel dosyası yükleyin",
            key="ncf_file"
        )
        
        if file is not None:
            try:
                # Dosya önizlemesi ve format önerisi
                import io
                file_content = file.read()
                st.session_state.ncf_file_content = file_content
                file_bytes = io.BytesIO(file_content)
                
                import pandas as pd
                if file.name.endswith('.csv'):
                    file_bytes.seek(0)
                    first_line = file_bytes.readline().decode('utf-8', errors='ignore')
                    delimiters = [',', ';', '\t', '|']
                    detected_delimiter = ','
                    max_cols = 0
                    for delim in delimiters:
                        cols = first_line.split(delim)
                        if len(cols) > max_cols:
                            max_cols = len(cols)
                            detected_delimiter = delim
                    
                    file_bytes.seek(0)
                    preview_df = pd.read_csv(file_bytes, nrows=5, sep=detected_delimiter, engine='python')
                    file_bytes.seek(0)
                    total_df = pd.read_csv(file_bytes, sep=detected_delimiter, engine='python')
                elif file.name.endswith(('.xlsx', '.xls')):
                    file_bytes.seek(0)
                    preview_df = pd.read_excel(file_bytes, nrows=5)
                    file_bytes.seek(0)
                    total_df = pd.read_excel(file_bytes)
                else:
                    preview_df = None
                    total_df = None
                
                if preview_df is not None:
                    with st.expander("👁️ Dosya Önizleme (İlk 5 Satır)", expanded=False):
                        st.dataframe(preview_df, width='stretch')
                        st.info(f"""
                        **Dosya Bilgileri:**
                        - **Satır Sayısı**: {len(total_df) if total_df is not None else 'Bilinmiyor'}
                        - **Sütun Sayısı**: {len(preview_df.columns)}
                        - **Sütun İsimleri**: {', '.join(preview_df.columns.tolist()[:10])}{'...' if len(preview_df.columns) > 10 else ''}
                        
                        **Format Önerisi:**
                        - **3 sütun varsa** → Long Format seçin (user_id, item_id, rating)
                        - **10+ sütun varsa** → Matrix Format seçin (ilk sütun kullanıcı ID, diğerleri ürün ID)
                        """)
                
                data_format = st.radio(
                    "Veri Formatı",
                    ["Long Format (user_id, item_id, rating)", "Matrix Format (Rating Matrisi)"],
                    horizontal=True,
                    key="ncf_format"
                )
                
                if data_format == "Long Format (user_id, item_id, rating)":
                    if st.button("📥 Veriyi Yükle", key="ncf_load_long"):
                        with st.spinner("Veri yükleniyor..."):
                            file_name = file.name
                            file_size = file.size
                            
                            if 'ncf_file_content' in st.session_state:
                                import io
                                file_bytes = io.BytesIO(st.session_state.ncf_file_content)
                                file_bytes.name = file_name
                                rating_matrix, user_mapping, item_mapping = load_rating_data_from_file(file_bytes)
                            else:
                                file.seek(0)
                                rating_matrix, user_mapping, item_mapping = load_rating_data_from_file(file)
                            
                            st.session_state.ncf_file_name = file_name
                            st.session_state.ncf_file_size = file_size
                            st.session_state.ncf_rating_matrix = rating_matrix
                            st.session_state.ncf_user_mapping = user_mapping
                            st.session_state.ncf_item_mapping = item_mapping
                            
                            from scipy.sparse import issparse
                            if issparse(rating_matrix):
                                n_ratings = rating_matrix.nnz
                                density = rating_matrix.nnz / (rating_matrix.shape[0] * rating_matrix.shape[1]) * 100 if rating_matrix.shape[0] * rating_matrix.shape[1] > 0 else 0
                            else:
                                mask = ~np.isnan(rating_matrix)
                                n_ratings = np.sum(mask)
                                density = (1 - np.isnan(rating_matrix).sum() / rating_matrix.size) * 100 if rating_matrix.size > 0 else 0
                            
                            st.success(f"✅ Veri yüklendi! {rating_matrix.shape[0]} kullanıcı, {rating_matrix.shape[1]} ürün")
                            st.info(f"📊 Veri yoğunluğu: {density:.2f}%")
                            st.rerun()
                else:
                    if st.button("📥 Veriyi Yükle", key="ncf_load_matrix"):
                        with st.spinner("Veri yükleniyor..."):
                            file_name = file.name
                            file_size = file.size
                            
                            if 'ncf_file_content' in st.session_state:
                                import io
                                file_bytes = io.BytesIO(st.session_state.ncf_file_content)
                                file_bytes.name = file_name
                                try:
                                    rating_matrix = load_rating_matrix_from_file(file_bytes)
                                except Exception as e:
                                    st.error(f"❌ Hata: {str(e)}")
                                    st.stop()
                            else:
                                file.seek(0)
                                try:
                                    rating_matrix = load_rating_matrix_from_file(file)
                                except Exception as e:
                                    st.error(f"❌ Hata: {str(e)}")
                                    st.stop()
                            
                            st.session_state.ncf_file_name = file_name
                            st.session_state.ncf_file_size = file_size
                            st.session_state.ncf_rating_matrix = rating_matrix
                            st.session_state.ncf_user_mapping = None
                            st.session_state.ncf_item_mapping = None
                            
                            from scipy.sparse import issparse
                            if issparse(rating_matrix):
                                n_ratings = rating_matrix.nnz
                                density = rating_matrix.nnz / (rating_matrix.shape[0] * rating_matrix.shape[1]) * 100 if rating_matrix.shape[0] * rating_matrix.shape[1] > 0 else 0
                            else:
                                mask = ~np.isnan(rating_matrix)
                                n_ratings = np.sum(mask)
                                density = (1 - np.isnan(rating_matrix).sum() / rating_matrix.size) * 100 if rating_matrix.size > 0 else 0
                            
                            st.success(f"✅ Veri yüklendi! {rating_matrix.shape[0]} kullanıcı, {rating_matrix.shape[1]} ürün")
                            st.info(f"📊 Veri yoğunluğu: {density:.2f}%")
                            st.rerun()
                            
            except Exception as e:
                st.error(f"❌ Hata: {str(e)}")
                st.info("💡 Lütfen veri formatını kontrol edin.")
    else:
        # Örnek veri oluştur - session state'i temizle
        if st.session_state.ncf_rating_matrix is not None:
            st.session_state.ncf_rating_matrix = None
            st.session_state.ncf_user_mapping = None
            st.session_state.ncf_item_mapping = None
    
    # Model parametreleri
    if rating_matrix is not None:
        n_users, n_items = rating_matrix.shape
        st.info(f"📊 Yüklenen veri: {n_users} kullanıcı, {n_items} ürün")
    else:
        col1, col2 = st.columns(2)
        with col1:
            n_users = st.slider("Kullanıcı Sayısı", 50, 500, 100, key="ncf_n_users")
            n_items = st.slider("Ürün Sayısı", 30, 300, 50, key="ncf_n_items")
        with col2:
            sparsity = st.slider("Sparsity (Eksik Rating Oranı)", 0.3, 0.9, 0.6, key="ncf_sparsity")
    
    # Optimal parametreleri al
    optimal_params = get_optimal_model_params("ncf")
    optimal_n_factors = optimal_params['n_factors']
    optimal_epochs = optimal_params['epochs']
    optimal_batch_size = optimal_params['batch_size']
    optimal_dropout = optimal_params['dropout']
    
    n_factors = st.slider(
        "Latent Faktör Sayısı", 
        10, 100, 
        optimal_n_factors, 
        key="ncf_n_factors",
        help=f"Önerilen değer: {optimal_n_factors}"
    )
    
    col3, col4 = st.columns(2)
    with col3:
        epochs = st.slider("Epochs (Eğitim İterasyonu)", 5, 50, 10, key="ncf_epochs")
        batch_size = st.slider("Batch Size", 32, 512, 256, key="ncf_batch_size")
    with col4:
        hidden_layers_str = st.text_input("Gizli Katmanlar (virgülle ayırın)", "64,32,16", key="ncf_hidden")
        dropout = st.slider(
            "Dropout Rate", 
            0.0, 0.5, 
            optimal_dropout, 
            key="ncf_dropout",
            help=f"Önerilen değer: {optimal_dropout}"
        )
    
    if st.button("🚀 Modeli Eğit ve Öneriler Üret", key="ncf_train"):
        if rating_matrix is None and data_source == "📁 Dosyadan Yükle":
            st.warning("⚠️ Lütfen önce veri dosyasını yükleyin!")
        else:
            with st.spinner("Veri hazırlanıyor..."):
                if rating_matrix is None:
                    rating_matrix = generate_rating_matrix(n_users, n_items, sparsity)
        
            # Sparse matrix'i dense'e çevir (NCF için gerekli)
            from scipy.sparse import issparse
            if issparse(rating_matrix):
                rating_matrix_dense = rating_matrix.toarray()
                # 0 değerlerini NaN'a çevir
                rating_matrix_dense = np.where(rating_matrix_dense == 0, np.nan, rating_matrix_dense)
            else:
                rating_matrix_dense = rating_matrix.copy()
            
            with st.spinner("NCF modeli eğitiliyor (bu biraz zaman alabilir)..."):
                try:
                    import time
                    training_start = time.time()
                    
                    hidden_layers = [int(x.strip()) for x in hidden_layers_str.split(',')]
                    ncf_model = NCFRecommender(
                        n_factors=n_factors,
                        hidden_layers=hidden_layers,
                        dropout_rate=dropout
                    )
                    history = ncf_model.fit(
                        rating_matrix_dense,
                        epochs=epochs,
                        batch_size=batch_size,
                        verbose=0
                    )
                    
                    training_time = time.time() - training_start
                    
                    st.success("✅ Model eğitildi!")
                    
                    # Sonuç açıklaması
                    with st.expander("📝 Sonuç Açıklaması - Ne Elde Edildi?", expanded=True):
                        st.markdown("### 🔍 Kullanılan Veriler")
                        if data_source == "📁 Dosyadan Yükle" and 'ncf_file_name' in st.session_state:
                            from scipy.sparse import issparse as issparse_check
                            if issparse_check(rating_matrix):
                                n_ratings = rating_matrix.nnz
                            else:
                                mask = ~np.isnan(rating_matrix)
                                n_ratings = np.sum(mask)
                            
                            st.info(f"""
                            **Dosya**: {st.session_state.ncf_file_name} ({st.session_state.ncf_file_size / 1024:.2f} KB)
                            - **Kullanıcı Sayısı**: {n_users:,}
                            - **Ürün Sayısı**: {n_items:,}
                            - **Toplam Rating**: {n_ratings:,}
                            """)
                        else:
                            st.info(f"""
                            **Örnek Veri**:
                            - **Kullanıcı Sayısı**: {n_users:,}
                            - **Ürün Sayısı**: {n_items:,}
                            """)
                        
                        st.markdown("### ⚙️ Model Parametreleri")
                        st.info(f"""
                        - **Latent Faktör Sayısı**: {n_factors}
                        - **Gizli Katmanlar**: {hidden_layers_str}
                        - **Dropout Rate**: {dropout}
                        - **Epochs**: {epochs}
                        - **Batch Size**: {batch_size}
                        - **Eğitim Süresi**: {training_time:.2f} saniye
                        """)
                        
                        st.markdown("### 📈 Elde Edilen Sonuçlar")
                        st.success(f"""
                        **NCF Modeli Başarıyla Eğitildi!**
                        
                        **Ne Yapıldı?**
                        1. **Embedding Katmanları**: Kullanıcı ve ürünler {n_factors} boyutlu latent space'e embed edildi
                        2. **MLP (Multi-Layer Perceptron)**: {hidden_layers_str} yapısında derin sinir ağı ile doğrusal olmayan ilişkiler öğrenildi
                        3. **Eğitim**: {epochs} epoch boyunca model optimize edildi
                        4. **Sonuç**: Model, kullanıcı-ürün etkileşimlerinden karmaşık pattern'leri öğrendi
                        
                        **SVD'den Farkı:**
                        - SVD sadece doğrusal ilişkileri yakalar
                        - NCF doğrusal olmayan, karmaşık ilişkileri öğrenir
                        - Daha güçlü özellik öğrenme kapasitesi
                        """)
                    
                    # Eğitim geçmişi
                    if history and hasattr(history, 'history'):
                        st.subheader("📊 Eğitim Geçmişi")
                        fig, ax = plt.subplots(figsize=(10, 4))
                        ax.plot(history.history['loss'], label='Training Loss', linewidth=2)
                        if 'val_loss' in history.history:
                            ax.plot(history.history['val_loss'], label='Validation Loss', linewidth=2)
                        ax.set_xlabel('Epoch', fontsize=12)
                        ax.set_ylabel('Loss', fontsize=12)
                        ax.set_title('NCF Eğitim Geçmişi - Loss Değişimi', fontsize=14, fontweight='bold')
                        ax.legend()
                        ax.grid(True, alpha=0.3)
                        st.pyplot(fig)
                    
                    # Öneriler
                    st.subheader("🎯 Kullanıcı Önerileri")
                    user_idx = st.selectbox("Kullanıcı Seçin", range(min(10, n_users)), key="ncf_user_select")
                    
                    item_indices, predicted_ratings = ncf_model.recommend(
                        user_idx, n_recommendations=10, rating_matrix=rating_matrix_dense
                    )
                    
                    recommendations_df = pd.DataFrame({
                        'Ürün ID': item_indices + 1,
                        'Tahmin Edilen Rating': np.round(np.clip(predicted_ratings, 1, 5), 2)
                    })
                    st.dataframe(recommendations_df, width='stretch')
                    
                except Exception as e:
                    st.error(f"❌ Hata: {str(e)}")
                    st.info("💡 PyTorch yüklü olduğundan emin olun: `pip install torch`")


def show_autoencoder_denoising():
    """Denoising Autoencoder sayfası"""
    st.header("🎨 Denoising Autoencoder - Gürültü Temizleme")
    
    with st.expander("ℹ️ Autoencoder Nedir?", expanded=False):
        st.markdown("""
        **Autoencoder**, SVD ve PCA'in Deep Learning karşılığıdır.
        
        **Nasıl Çalışır:**
        1. **Encoder**: Veriyi sıkıştırır (latent space'e)
        2. **Decoder**: Sıkıştırılmış veriyi tekrar genişletir
        3. **Öğrenme**: Gürültülü veriden temiz veri üretmeyi öğrenir
        
        **Kullanım Alanları:**
        - Görüntü gürültü temizleme
        - Sinyal işleme
        - Veri sıkıştırma
        """)
    
    # Veri kaynağı seçimi
    data_source = st.radio(
        "Veri Kaynağı Seçin",
        ["📁 Dosya Yükle (CSV/Excel)", "🎲 Örnek Veri Oluştur"],
        horizontal=True
    )
    
    uploaded_file = None
    original_data = None
    noisy_data = None
    data_loaded = False
    file_bytes = None
    file_name = None
    numeric_cols = None
    df_numeric = None  # DataFrame referansı (indirme için)
    
    if data_source == "📁 Dosya Yükle (CSV/Excel)":
        st.markdown("### 📤 Veri Dosyası Yükle")
        uploaded_file = st.file_uploader(
            "Veri dosyası seçin (CSV, Excel)",
            type=['csv', 'xlsx', 'xls'],
            help="Yüklediğiniz veri gürültü temizleme için kullanılacaktır. Her satır bir örnek, her sütun bir özellik olmalıdır."
        )
        
        if uploaded_file is not None:
            try:
                import io
                file_name = uploaded_file.name
                
                # Dosya içeriğini oku
                file_content = uploaded_file.read()
                file_bytes = io.BytesIO(file_content)
                file_bytes_ref = io.BytesIO(file_content)  # Referans için kopya
                
                # Dosya tipine göre yükle
                if file_name.endswith('.csv'):
                    df = pd.read_csv(file_bytes)
                elif file_name.endswith(('.xlsx', '.xls')):
                    df = pd.read_excel(file_bytes)
                else:
                    raise ValueError("Desteklenmeyen dosya formatı")
                
                # Sayısal sütunları seç
                numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
                if len(numeric_cols) == 0:
                    raise ValueError("Dosyada sayısal sütun bulunamadı!")
                
                df_numeric = df[numeric_cols]
                
                # NaN değerleri doldur (ortalama ile)
                df_numeric = df_numeric.fillna(df_numeric.mean())
                
                # Veriyi numpy array'e çevir
                df_numeric = df_numeric  # Referansı sakla
                original_data = df_numeric.values
                
                st.success(f"✅ Veri yüklendi! {original_data.shape[0]} örnek, {original_data.shape[1]} özellik")
                
                # Veri önizleme
                with st.expander("📊 Veri Önizleme", expanded=False):
                    st.dataframe(df_numeric.head(10), width='stretch')
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.metric("Örnek Sayısı", original_data.shape[0])
                    with col2:
                        st.metric("Özellik Sayısı", original_data.shape[1])
                    with col3:
                        st.metric("Veri Boyutu", f"{original_data.size:,}")
                
                data_loaded = True
                
            except Exception as e:
                st.error(f"❌ Dosya okunurken hata oluştu: {str(e)}")
                st.info("💡 Lütfen geçerli bir CSV veya Excel dosyası yüklediğinizden emin olun.")
    
    else:
        st.markdown("### 🎲 Örnek Veri")
        col1, col2 = st.columns(2)
        # Optimal parametreleri al
        optimal_params = get_optimal_model_params("autoencoder", n_samples=200, n_features=50)
        optimal_encoding_dim = optimal_params['encoding_dim']
        optimal_epochs = optimal_params['epochs']
        optimal_noise_factor = optimal_params['noise_factor']
        
        with col1:
            n_samples = st.slider("Örnek Sayısı", 100, 1000, 200)
            n_features = st.slider("Özellik Sayısı", 20, 200, 50)
        with col2:
            encoding_dim = st.slider(
                "Encoding Boyutu (Latent Space)", 
                5, 50, 
                optimal_encoding_dim,
                help=f"Önerilen değer: {optimal_encoding_dim}"
            )
            noise_factor = st.slider(
                "Gürültü Faktörü", 
                0.1, 0.5, 
                optimal_noise_factor,
                help=f"Önerilen değer: {optimal_noise_factor}"
            )
    
    # Epochs slider (her iki mod için)
    epochs = st.slider(
        "Epochs", 
        10, 100, 
        optimal_epochs if data_source == "🎲 Örnek Veri Oluştur" else 50,
        help=f"Önerilen değer: {optimal_epochs if data_source == '🎲 Örnek Veri Oluştur' else 50}"
    )
    
    # Gürültü faktörü (dosya yüklendiyse)
    if data_source == "📁 Dosya Yükle (CSV/Excel)" and data_loaded:
        # Optimal parametreleri al
        optimal_params = get_optimal_model_params("autoencoder", n_features=original_data.shape[1])
        optimal_encoding_dim = optimal_params['encoding_dim']
        optimal_noise_factor = optimal_params['noise_factor']
        
        col1, col2 = st.columns(2)
        with col1:
            encoding_dim = st.slider(
                "Encoding Boyutu (Latent Space)", 
                5, min(50, original_data.shape[1]//2), 
                optimal_encoding_dim, 
                key="file_encoding",
                help=f"Önerilen değer: {optimal_encoding_dim}"
            )
        with col2:
            noise_factor = st.slider(
                "Gürültü Faktörü", 
                0.1, 0.5, 
                optimal_noise_factor, 
                key="file_noise",
                help=f"Önerilen değer: {optimal_noise_factor}"
            )
    
    # Analiz butonu
    analyze_button = st.button("🚀 Gürültü Temizle")
    
    if analyze_button:
        if data_source == "📁 Dosya Yükle (CSV/Excel)":
            if uploaded_file is None or original_data is None:
                st.warning("⚠️ Lütfen önce bir veri dosyası yükleyin!")
            else:
                with st.spinner("Gürültülü veri oluşturuluyor..."):
                    # Gürültülü veri oluştur
                    noisy_data = generate_noisy_data(original_data, noise_level=noise_factor)
                
                with st.spinner("Autoencoder eğitiliyor (bu biraz zaman alabilir)..."):
                    try:
                        autoencoder = DenoisingAutoencoder(
                            encoding_dim=encoding_dim,
                            noise_factor=noise_factor
                        )
                        history = autoencoder.fit(
                            noisy_data,
                            epochs=epochs,
                            verbose=0
                        )
                        
                        st.success("✅ Model eğitildi!")
                        
                        # Gürültü temizleme
                        with st.spinner("Gürültü temizleniyor..."):
                            denoised_data = autoencoder.denoise(noisy_data)
                        
                        # Metrikler
                        mse_original = mean_squared_error(original_data, noisy_data)
                        mse_denoised = mean_squared_error(original_data, denoised_data)
                        improvement = ((mse_original - mse_denoised) / mse_original) * 100
                        
                        st.subheader("📊 Sonuçlar")
                        col1, col2, col3 = st.columns(3)
                        with col1:
                            st.metric("Gürültülü MSE", f"{mse_original:.6f}")
                        with col2:
                            st.metric("Temizlenmiş MSE", f"{mse_denoised:.6f}")
                        with col3:
                            st.metric("İyileştirme", f"%{improvement:.1f}")
                        
                        if improvement > 0:
                            st.success(f"🎉 %{improvement:.1f} iyileştirme sağlandı!")
                        else:
                            st.warning("⚠️ Gürültü temizleme sonucu beklenen iyileştirmeyi sağlamadı.")
                        
                        # Görselleştirme
                        st.subheader("📈 Görselleştirme")
                        with st.expander("ℹ️ Grafikler Ne Gösteriyor?", expanded=False):
                            st.markdown("""
                            **Karşılaştırma Grafikleri**:
                            - **Orijinal Veri**: Temiz, gürültüsüz orijinal veri
                            - **Gürültülü Veri**: Gürültü eklenmiş veri
                            - **Temizlenmiş Veri**: Autoencoder ile gürültü temizlenmiş veri
                            
                            **Heatmap**: Veri matrisinin görselleştirmesi (ilk 50x50)
                            """)
                        
                        # Heatmap karşılaştırması
                        max_vis = min(50, original_data.shape[0], original_data.shape[1])
                        fig, axes = plt.subplots(1, 3, figsize=(18, 5))
                        
                        # Orijinal
                        im1 = axes[0].imshow(original_data[:max_vis, :max_vis], cmap='viridis', aspect='auto')
                        axes[0].set_title('Orijinal Veri', fontsize=12, fontweight='bold')
                        axes[0].set_xlabel('Özellikler')
                        axes[0].set_ylabel('Örnekler')
                        plt.colorbar(im1, ax=axes[0])
                        
                        # Gürültülü
                        im2 = axes[1].imshow(noisy_data[:max_vis, :max_vis], cmap='viridis', aspect='auto')
                        axes[1].set_title('Gürültülü Veri', fontsize=12, fontweight='bold')
                        axes[1].set_xlabel('Özellikler')
                        axes[1].set_ylabel('Örnekler')
                        plt.colorbar(im2, ax=axes[1])
                        
                        # Temizlenmiş
                        im3 = axes[2].imshow(denoised_data[:max_vis, :max_vis], cmap='viridis', aspect='auto')
                        axes[2].set_title('Temizlenmiş Veri', fontsize=12, fontweight='bold')
                        axes[2].set_xlabel('Özellikler')
                        axes[2].set_ylabel('Örnekler')
                        plt.colorbar(im3, ax=axes[2])
                        
                        plt.tight_layout()
                        st.pyplot(fig)
                        
                        # İndirilebilir sonuçlar
                        st.subheader("💾 Sonuçları İndir")
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            # Temizlenmiş veriyi DataFrame'e çevir
                            if numeric_cols is not None:
                                denoised_df = pd.DataFrame(denoised_data, columns=numeric_cols)
                            else:
                                denoised_df = pd.DataFrame(denoised_data, columns=[f'Feature_{i+1}' for i in range(denoised_data.shape[1])])
                            csv_denoised = denoised_df.to_csv(index=False).encode('utf-8')
                            st.download_button(
                                label="📥 Temizlenmiş Veriyi İndir (CSV)",
                                data=csv_denoised,
                                file_name="denoised_data.csv",
                                mime="text/csv"
                            )
                        
                        with col2:
                            # Karşılaştırma raporu
                            comparison_df = pd.DataFrame({
                                'Metrik': ['Gürültülü MSE', 'Temizlenmiş MSE', 'İyileştirme (%)'],
                                'Değer': [mse_original, mse_denoised, improvement]
                            })
                            csv_report = comparison_df.to_csv(index=False).encode('utf-8')
                            st.download_button(
                                label="📊 Raporu İndir (CSV)",
                                data=csv_report,
                                file_name="denoising_report.csv",
                                mime="text/csv"
                            )
                        
                        # Eğitim geçmişi
                        if history and 'loss' in history:
                            st.subheader("📉 Eğitim Geçmişi")
                            fig_history, ax_history = plt.subplots(figsize=(10, 5))
                            ax_history.plot(history['loss'], label='Training Loss', linewidth=2)
                            if 'val_loss' in history:
                                ax_history.plot(history['val_loss'], label='Validation Loss', linewidth=2)
                            ax_history.set_xlabel('Epoch')
                            ax_history.set_ylabel('Loss')
                            ax_history.set_title('Model Eğitim Geçmişi', fontsize=14, fontweight='bold')
                            ax_history.legend()
                            ax_history.grid(True, alpha=0.3)
                            plt.tight_layout()
                            st.pyplot(fig_history)
                        
                    except Exception as e:
                        st.error(f"❌ Analiz sırasında hata oluştu: {str(e)}")
                        import traceback
                        with st.expander("🔍 Detaylı Hata Mesajı"):
                            st.code(traceback.format_exc())
        
        else:
            # Örnek veri
            with st.spinner("Veri oluşturuluyor..."):
                # Orijinal veri
                original_data = generate_sample_data(n_samples, n_features)[0]
                # Gürültülü veri
                noisy_data = generate_noisy_data(original_data, noise_level=noise_factor)
            
            with st.spinner("Autoencoder eğitiliyor..."):
                try:
                    autoencoder = DenoisingAutoencoder(
                        encoding_dim=encoding_dim,
                        noise_factor=noise_factor
                    )
                    history = autoencoder.fit(
                        noisy_data,
                        epochs=epochs,
                        verbose=0
                    )
                    
                    st.success("✅ Model eğitildi!")
                    
                    # Gürültü temizleme
                    denoised_data = autoencoder.denoise(noisy_data)
                    
                    # Metrikler
                    mse_original = mean_squared_error(original_data, noisy_data)
                    mse_denoised = mean_squared_error(original_data, denoised_data)
                    improvement = ((mse_original - mse_denoised) / mse_original) * 100
                    
                    st.subheader("📊 Sonuçlar")
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.metric("Gürültülü MSE", f"{mse_original:.4f}")
                    with col2:
                        st.metric("Temizlenmiş MSE", f"{mse_denoised:.4f}")
                    with col3:
                        st.metric("İyileştirme", f"%{improvement:.1f}")
                    
                    if improvement > 0:
                        st.success(f"🎉 %{improvement:.1f} iyileştirme sağlandı!")
                    
                    # Görselleştirme
                    st.subheader("📈 Görselleştirme")
                    max_vis = min(50, original_data.shape[0], original_data.shape[1])
                    fig, axes = plt.subplots(1, 3, figsize=(18, 5))
                    
                    # Orijinal
                    im1 = axes[0].imshow(original_data[:max_vis, :max_vis], cmap='viridis', aspect='auto')
                    axes[0].set_title('Orijinal Veri', fontsize=12, fontweight='bold')
                    axes[0].set_xlabel('Özellikler')
                    axes[0].set_ylabel('Örnekler')
                    plt.colorbar(im1, ax=axes[0])
                    
                    # Gürültülü
                    im2 = axes[1].imshow(noisy_data[:max_vis, :max_vis], cmap='viridis', aspect='auto')
                    axes[1].set_title('Gürültülü Veri', fontsize=12, fontweight='bold')
                    axes[1].set_xlabel('Özellikler')
                    axes[1].set_ylabel('Örnekler')
                    plt.colorbar(im2, ax=axes[1])
                    
                    # Temizlenmiş
                    im3 = axes[2].imshow(denoised_data[:max_vis, :max_vis], cmap='viridis', aspect='auto')
                    axes[2].set_title('Temizlenmiş Veri', fontsize=12, fontweight='bold')
                    axes[2].set_xlabel('Özellikler')
                    axes[2].set_ylabel('Örnekler')
                    plt.colorbar(im3, ax=axes[2])
                    
                    plt.tight_layout()
                    st.pyplot(fig)
                    
                    # Eğitim geçmişi
                    if history and 'loss' in history:
                        st.subheader("📉 Eğitim Geçmişi")
                        fig_history, ax_history = plt.subplots(figsize=(10, 5))
                        ax_history.plot(history['loss'], label='Training Loss', linewidth=2)
                        if 'val_loss' in history:
                            ax_history.plot(history['val_loss'], label='Validation Loss', linewidth=2)
                        ax_history.set_xlabel('Epoch')
                        ax_history.set_ylabel('Loss')
                        ax_history.set_title('Model Eğitim Geçmişi', fontsize=14, fontweight='bold')
                        ax_history.legend()
                        ax_history.grid(True, alpha=0.3)
                        plt.tight_layout()
                        st.pyplot(fig_history)
                    
                except Exception as e:
                    st.error(f"❌ Hata: {str(e)}")
                    import traceback
                    with st.expander("🔍 Detaylı Hata Mesajı"):
                        st.code(traceback.format_exc())


def show_vae_recommender():
    """VAE (Variational Autoencoder) sayfası"""
    st.header("🎯 Variational Autoencoder (VAE) - Öneri Sistemi")
    
    with st.expander("ℹ️ VAE Nedir?", expanded=False):
        st.markdown("""
        **Variational Autoencoder (VAE)**, SVD'nin probabilistic, Deep Learning versiyonudur.
        
        **Özellikleri:**
        - ✅ Latent space'de **probabilistic** dağılım öğrenir
        - ✅ **KL divergence** ile düzenli latent space
        - ✅ Daha iyi genelleme (generalization)
        
        **Kullanım Alanları:**
        - Öneri sistemleri
        - Görüntü üretimi
        - Anomali tespiti
        """)
    
    # Veri yükleme seçeneği
    data_source = st.radio(
        "Veri Kaynağı",
        ["📊 Örnek Veri Oluştur", "📁 Dosyadan Yükle"],
        horizontal=True,
        key="vae_data_source"
    )
    
    # Session state ile veriyi koru
    if 'vae_rating_matrix' not in st.session_state:
        st.session_state.vae_rating_matrix = None
        st.session_state.vae_user_mapping = None
        st.session_state.vae_item_mapping = None
    
    rating_matrix = st.session_state.vae_rating_matrix
    user_mapping = st.session_state.vae_user_mapping
    item_mapping = st.session_state.vae_item_mapping
    n_users = None
    n_items = None
    
    if data_source == "📁 Dosyadan Yükle":
        st.markdown("### 📁 Veri Dosyası Yükle")
        st.info("""
        **Desteklenen Formatlar:** CSV, Excel (.xlsx, .xls)
        
        **Veri Formatı:** Long Format (user_id, item_id, rating) veya Matrix Format
        """)
        
        file = st.file_uploader(
            "Veri dosyasını seçin",
            type=['csv', 'xlsx', 'xls'],
            help="CSV veya Excel dosyası yükleyin",
            key="vae_file"
        )
        
        if file is not None:
            try:
                # Dosya önizlemesi ve format önerisi
                import io
                file_content = file.read()
                st.session_state.vae_file_content = file_content
                file_bytes = io.BytesIO(file_content)
                
                import pandas as pd
                if file.name.endswith('.csv'):
                    file_bytes.seek(0)
                    first_line = file_bytes.readline().decode('utf-8', errors='ignore')
                    delimiters = [',', ';', '\t', '|']
                    detected_delimiter = ','
                    max_cols = 0
                    for delim in delimiters:
                        cols = first_line.split(delim)
                        if len(cols) > max_cols:
                            max_cols = len(cols)
                            detected_delimiter = delim
                    
                    file_bytes.seek(0)
                    preview_df = pd.read_csv(file_bytes, nrows=5, sep=detected_delimiter, engine='python')
                    file_bytes.seek(0)
                    total_df = pd.read_csv(file_bytes, sep=detected_delimiter, engine='python')
                elif file.name.endswith(('.xlsx', '.xls')):
                    file_bytes.seek(0)
                    preview_df = pd.read_excel(file_bytes, nrows=5)
                    file_bytes.seek(0)
                    total_df = pd.read_excel(file_bytes)
                else:
                    preview_df = None
                    total_df = None
                
                if preview_df is not None:
                    with st.expander("👁️ Dosya Önizleme (İlk 5 Satır)", expanded=False):
                        st.dataframe(preview_df, width='stretch')
                        st.info(f"""
                        **Dosya Bilgileri:**
                        - **Satır Sayısı**: {len(total_df) if total_df is not None else 'Bilinmiyor'}
                        - **Sütun Sayısı**: {len(preview_df.columns)}
                        - **Sütun İsimleri**: {', '.join(preview_df.columns.tolist()[:10])}{'...' if len(preview_df.columns) > 10 else ''}
                        
                        **Format Önerisi:**
                        - **3 sütun varsa** → Long Format seçin (user_id, item_id, rating)
                        - **10+ sütun varsa** → Matrix Format seçin (ilk sütun kullanıcı ID, diğerleri ürün ID)
                        """)
                
                data_format = st.radio(
                    "Veri Formatı",
                    ["Long Format (user_id, item_id, rating)", "Matrix Format (Rating Matrisi)"],
                    horizontal=True,
                    key="vae_format"
                )
                
                if data_format == "Long Format (user_id, item_id, rating)":
                    if st.button("📥 Veriyi Yükle", key="vae_load_long"):
                        with st.spinner("Veri yükleniyor..."):
                            file_name = file.name
                            file_size = file.size
                            
                            if 'vae_file_content' in st.session_state:
                                import io
                                file_bytes = io.BytesIO(st.session_state.vae_file_content)
                                file_bytes.name = file_name
                                rating_matrix, user_mapping, item_mapping = load_rating_data_from_file(file_bytes)
                            else:
                                file.seek(0)
                                rating_matrix, user_mapping, item_mapping = load_rating_data_from_file(file)
                            
                            st.session_state.vae_file_name = file_name
                            st.session_state.vae_file_size = file_size
                            st.session_state.vae_rating_matrix = rating_matrix
                            st.session_state.vae_user_mapping = user_mapping
                            st.session_state.vae_item_mapping = item_mapping
                            
                            from scipy.sparse import issparse
                            if issparse(rating_matrix):
                                n_ratings = rating_matrix.nnz
                                density = rating_matrix.nnz / (rating_matrix.shape[0] * rating_matrix.shape[1]) * 100 if rating_matrix.shape[0] * rating_matrix.shape[1] > 0 else 0
                            else:
                                mask = ~np.isnan(rating_matrix)
                                n_ratings = np.sum(mask)
                                density = (1 - np.isnan(rating_matrix).sum() / rating_matrix.size) * 100 if rating_matrix.size > 0 else 0
                            
                            st.success(f"✅ Veri yüklendi! {rating_matrix.shape[0]} kullanıcı, {rating_matrix.shape[1]} ürün")
                            st.info(f"📊 Veri yoğunluğu: {density:.2f}%")
                            st.rerun()
                else:
                    if st.button("📥 Veriyi Yükle", key="vae_load_matrix"):
                        with st.spinner("Veri yükleniyor..."):
                            file_name = file.name
                            file_size = file.size
                            
                            if 'vae_file_content' in st.session_state:
                                import io
                                file_bytes = io.BytesIO(st.session_state.vae_file_content)
                                file_bytes.name = file_name
                                try:
                                    rating_matrix = load_rating_matrix_from_file(file_bytes)
                                except Exception as e:
                                    st.error(f"❌ Hata: {str(e)}")
                                    st.stop()
                            else:
                                file.seek(0)
                                try:
                                    rating_matrix = load_rating_matrix_from_file(file)
                                except Exception as e:
                                    st.error(f"❌ Hata: {str(e)}")
                                    st.stop()
                            
                            st.session_state.vae_file_name = file_name
                            st.session_state.vae_file_size = file_size
                            st.session_state.vae_rating_matrix = rating_matrix
                            st.session_state.vae_user_mapping = None
                            st.session_state.vae_item_mapping = None
                            
                            from scipy.sparse import issparse
                            if issparse(rating_matrix):
                                n_ratings = rating_matrix.nnz
                                density = rating_matrix.nnz / (rating_matrix.shape[0] * rating_matrix.shape[1]) * 100 if rating_matrix.shape[0] * rating_matrix.shape[1] > 0 else 0
                            else:
                                mask = ~np.isnan(rating_matrix)
                                n_ratings = np.sum(mask)
                                density = (1 - np.isnan(rating_matrix).sum() / rating_matrix.size) * 100 if rating_matrix.size > 0 else 0
                            
                            st.success(f"✅ Veri yüklendi! {rating_matrix.shape[0]} kullanıcı, {rating_matrix.shape[1]} ürün")
                            st.info(f"📊 Veri yoğunluğu: {density:.2f}%")
                            st.rerun()
                            
            except Exception as e:
                st.error(f"❌ Hata: {str(e)}")
                st.info("💡 Lütfen veri formatını kontrol edin.")
    else:
        # Örnek veri oluştur - session state'i temizle
        if st.session_state.vae_rating_matrix is not None:
            st.session_state.vae_rating_matrix = None
            st.session_state.vae_user_mapping = None
            st.session_state.vae_item_mapping = None
    
    # Model parametreleri
    if rating_matrix is not None:
        n_users, n_items = rating_matrix.shape
        st.info(f"📊 Yüklenen veri: {n_users} kullanıcı, {n_items} ürün")
    else:
        col1, col2 = st.columns(2)
        with col1:
            n_users = st.slider("Kullanıcı Sayısı", 50, 500, 100, key="vae_n_users")
            n_items = st.slider("Ürün Sayısı", 30, 300, 50, key="vae_n_items")
        with col2:
            sparsity = st.slider("Sparsity", 0.3, 0.9, 0.6, key="vae_sparsity")
    
    # Optimal parametreleri al
    optimal_params = get_optimal_model_params("vae")
    optimal_latent_dim = optimal_params['latent_dim']
    optimal_epochs = optimal_params['epochs']
    
    latent_dim = st.slider(
        "Latent Dimension", 
        10, 100, 
        optimal_latent_dim, 
        key="vae_latent",
        help=f"Önerilen değer: {optimal_latent_dim}"
    )
    epochs = st.slider(
        "Epochs", 
        10, 100, 
        optimal_epochs, 
        key="vae_epochs",
        help=f"Önerilen değer: {optimal_epochs}"
    )
    
    if st.button("🚀 VAE Modeli Eğit", key="vae_train"):
        if rating_matrix is None and data_source == "📁 Dosyadan Yükle":
            st.warning("⚠️ Lütfen önce veri dosyasını yükleyin!")
        else:
            with st.spinner("Veri hazırlanıyor..."):
                if rating_matrix is None:
                    rating_matrix = generate_rating_matrix(n_users, n_items, sparsity)
            
            # Sparse matrix'i dense'e çevir (VAE için gerekli)
            from scipy.sparse import issparse
            if issparse(rating_matrix):
                rating_matrix_dense = rating_matrix.toarray()
                # 0 değerlerini NaN'a çevir
                rating_matrix_dense = np.where(rating_matrix_dense == 0, np.nan, rating_matrix_dense)
            else:
                rating_matrix_dense = rating_matrix.copy()
            
            with st.spinner("VAE modeli eğitiliyor..."):
                try:
                    import time
                    training_start = time.time()
                    
                    vae_model = VAERecommender(latent_dim=latent_dim)
                    history = vae_model.fit(
                        rating_matrix_dense,
                        epochs=epochs,
                        verbose=0
                    )
                    
                    training_time = time.time() - training_start
                    
                    st.success("✅ Model eğitildi!")
                    
                    # Sonuç açıklaması
                    with st.expander("📝 Sonuç Açıklaması - Ne Elde Edildi?", expanded=True):
                        st.markdown("### 🔍 Kullanılan Veriler")
                        if data_source == "📁 Dosyadan Yükle" and 'vae_file_name' in st.session_state:
                            from scipy.sparse import issparse as issparse_check
                            if issparse_check(rating_matrix):
                                n_ratings = rating_matrix.nnz
                            else:
                                mask = ~np.isnan(rating_matrix)
                                n_ratings = np.sum(mask)
                            
                            st.info(f"""
                            **Dosya**: {st.session_state.vae_file_name} ({st.session_state.vae_file_size / 1024:.2f} KB)
                            - **Kullanıcı Sayısı**: {n_users:,}
                            - **Ürün Sayısı**: {n_items:,}
                            - **Toplam Rating**: {n_ratings:,}
                            """)
                        else:
                            st.info(f"""
                            **Örnek Veri**:
                            - **Kullanıcı Sayısı**: {n_users:,}
                            - **Ürün Sayısı**: {n_items:,}
                            """)
                        
                        st.markdown("### ⚙️ Model Parametreleri")
                        st.info(f"""
                        - **Latent Dimension**: {latent_dim}
                        - **Epochs**: {epochs}
                        - **Eğitim Süresi**: {training_time:.2f} saniye
                        """)
                        
                        st.markdown("### 📈 Elde Edilen Sonuçlar")
                        st.success(f"""
                        **VAE Modeli Başarıyla Eğitildi!**
                        
                        **Ne Yapıldı?**
                        1. **Encoder**: Kullanıcı rating'leri {latent_dim} boyutlu latent space'e encode edildi
                        2. **Probabilistic Latent Space**: Latent space'de Gaussian dağılım öğrenildi (mean + variance)
                        3. **KL Divergence**: Latent space düzenlendi (overfitting önlendi)
                        4. **Decoder**: Latent representation'dan rating'ler decode edildi
                        5. **Eğitim**: {epochs} epoch boyunca reconstruction + KL loss minimize edildi
                        
                        **SVD'den Farkı:**
                        - SVD deterministik latent factors öğrenir
                        - VAE probabilistic latent distribution öğrenir
                        - Daha iyi genelleme (generalization) sağlar
                        - KL divergence ile düzenli latent space
                        """)
                    
                    # Öneriler
                    st.subheader("🎯 Kullanıcı Önerileri")
                    user_idx = st.selectbox("Kullanıcı Seçin", range(min(10, n_users)), key="vae_user_select")
                    
                    item_indices, predicted_ratings = vae_model.recommend(
                        user_idx, n_recommendations=10, rating_matrix=rating_matrix_dense
                    )
                    
                    recommendations_df = pd.DataFrame({
                        'Ürün ID': item_indices + 1,
                        'Tahmin Edilen Rating': np.round(np.clip(predicted_ratings, 1, 5), 2)
                    })
                    st.dataframe(recommendations_df, width='stretch')
                    
                except Exception as e:
                    st.error(f"❌ Hata: {str(e)}")
                    st.info("💡 PyTorch yüklü olduğundan emin olun: `pip install torch`")


def show_fm_recommender():
    """Factorization Machines sayfası"""
    st.header("🔗 Factorization Machines (FM)")
    
    with st.expander("ℹ️ FM Nedir?", expanded=False):
        st.markdown("""
        **Factorization Machines**, context-aware öneri sistemi sağlar.
        
        **Farkları:**
        - ✅ Sadece kullanıcı-ürün ID'sine bakmaz
        - ✅ **Yan bilgileri** (context) kullanır:
          - Saat, gün, cihaz tipi
          - Ürün özellikleri (renk, kategori)
          - Kullanıcı özellikleri (yaş, konum)
        
        **Kullanım Alanları:**
        - Reklam tıklama tahmini (CTR)
        - Context-aware öneriler
        """)
    
    # Veri yükleme seçeneği
    data_source = st.radio(
        "Veri Kaynağı",
        ["📊 Örnek Veri Oluştur", "📁 Dosyadan Yükle"],
        horizontal=True,
        key="fm_data_source"
    )
    
    # Session state ile veriyi koru
    if 'fm_rating_matrix' not in st.session_state:
        st.session_state.fm_rating_matrix = None
        st.session_state.fm_user_mapping = None
        st.session_state.fm_item_mapping = None
    
    rating_matrix = st.session_state.fm_rating_matrix
    user_mapping = st.session_state.fm_user_mapping
    item_mapping = st.session_state.fm_item_mapping
    n_users = None
    n_items = None
    
    if data_source == "📁 Dosyadan Yükle":
        st.markdown("### 📁 Veri Dosyası Yükle")
        st.info("""
        **Desteklenen Formatlar:** CSV, Excel (.xlsx, .xls)
        
        **Veri Formatı:** Long Format (user_id, item_id, rating) veya Matrix Format
        
        **Not:** FM için context features otomatik oluşturulacak (saat, cihaz tipi gibi).
        """)
        
        file = st.file_uploader(
            "Veri dosyasını seçin",
            type=['csv', 'xlsx', 'xls'],
            help="CSV veya Excel dosyası yükleyin",
            key="fm_file"
        )
        
        if file is not None:
            try:
                # Dosya önizlemesi ve format önerisi
                import io
                file_content = file.read()
                st.session_state.fm_file_content = file_content
                file_bytes = io.BytesIO(file_content)
                
                import pandas as pd
                if file.name.endswith('.csv'):
                    file_bytes.seek(0)
                    first_line = file_bytes.readline().decode('utf-8', errors='ignore')
                    delimiters = [',', ';', '\t', '|']
                    detected_delimiter = ','
                    max_cols = 0
                    for delim in delimiters:
                        cols = first_line.split(delim)
                        if len(cols) > max_cols:
                            max_cols = len(cols)
                            detected_delimiter = delim
                    
                    file_bytes.seek(0)
                    preview_df = pd.read_csv(file_bytes, nrows=5, sep=detected_delimiter, engine='python')
                    file_bytes.seek(0)
                    total_df = pd.read_csv(file_bytes, sep=detected_delimiter, engine='python')
                elif file.name.endswith(('.xlsx', '.xls')):
                    file_bytes.seek(0)
                    preview_df = pd.read_excel(file_bytes, nrows=5)
                    file_bytes.seek(0)
                    total_df = pd.read_excel(file_bytes)
                else:
                    preview_df = None
                    total_df = None
                
                if preview_df is not None:
                    with st.expander("👁️ Dosya Önizleme (İlk 5 Satır)", expanded=False):
                        st.dataframe(preview_df, width='stretch')
                        st.info(f"""
                        **Dosya Bilgileri:**
                        - **Satır Sayısı**: {len(total_df) if total_df is not None else 'Bilinmiyor'}
                        - **Sütun Sayısı**: {len(preview_df.columns)}
                        - **Sütun İsimleri**: {', '.join(preview_df.columns.tolist()[:10])}{'...' if len(preview_df.columns) > 10 else ''}
                        
                        **Format Önerisi:**
                        - **3 sütun varsa** → Long Format seçin (user_id, item_id, rating)
                        - **10+ sütun varsa** → Matrix Format seçin (ilk sütun kullanıcı ID, diğerleri ürün ID)
                        """)
                
                data_format = st.radio(
                    "Veri Formatı",
                    ["Long Format (user_id, item_id, rating)", "Matrix Format (Rating Matrisi)"],
                    horizontal=True,
                    key="fm_format"
                )
                
                if data_format == "Long Format (user_id, item_id, rating)":
                    if st.button("📥 Veriyi Yükle", key="fm_load_long"):
                        with st.spinner("Veri yükleniyor..."):
                            file_name = file.name
                            file_size = file.size
                            
                            if 'fm_file_content' in st.session_state:
                                import io
                                file_bytes = io.BytesIO(st.session_state.fm_file_content)
                                file_bytes.name = file_name
                                rating_matrix, user_mapping, item_mapping = load_rating_data_from_file(file_bytes)
                            else:
                                file.seek(0)
                                rating_matrix, user_mapping, item_mapping = load_rating_data_from_file(file)
                            
                            st.session_state.fm_file_name = file_name
                            st.session_state.fm_file_size = file_size
                            st.session_state.fm_rating_matrix = rating_matrix
                            st.session_state.fm_user_mapping = user_mapping
                            st.session_state.fm_item_mapping = item_mapping
                            
                            from scipy.sparse import issparse
                            if issparse(rating_matrix):
                                n_ratings = rating_matrix.nnz
                                density = rating_matrix.nnz / (rating_matrix.shape[0] * rating_matrix.shape[1]) * 100 if rating_matrix.shape[0] * rating_matrix.shape[1] > 0 else 0
                            else:
                                mask = ~np.isnan(rating_matrix)
                                n_ratings = np.sum(mask)
                                density = (1 - np.isnan(rating_matrix).sum() / rating_matrix.size) * 100 if rating_matrix.size > 0 else 0
                            
                            st.success(f"✅ Veri yüklendi! {rating_matrix.shape[0]} kullanıcı, {rating_matrix.shape[1]} ürün")
                            st.info(f"📊 Veri yoğunluğu: {density:.2f}%")
                            st.rerun()
                else:
                    if st.button("📥 Veriyi Yükle", key="fm_load_matrix"):
                        with st.spinner("Veri yükleniyor..."):
                            file_name = file.name
                            file_size = file.size
                            
                            if 'fm_file_content' in st.session_state:
                                import io
                                file_bytes = io.BytesIO(st.session_state.fm_file_content)
                                file_bytes.name = file_name
                                try:
                                    rating_matrix = load_rating_matrix_from_file(file_bytes)
                                except Exception as e:
                                    st.error(f"❌ Hata: {str(e)}")
                                    st.stop()
                            else:
                                file.seek(0)
                                try:
                                    rating_matrix = load_rating_matrix_from_file(file)
                                except Exception as e:
                                    st.error(f"❌ Hata: {str(e)}")
                                    st.stop()
                            
                            st.session_state.fm_file_name = file_name
                            st.session_state.fm_file_size = file_size
                            st.session_state.fm_rating_matrix = rating_matrix
                            st.session_state.fm_user_mapping = None
                            st.session_state.fm_item_mapping = None
                            
                            from scipy.sparse import issparse
                            if issparse(rating_matrix):
                                n_ratings = rating_matrix.nnz
                                density = rating_matrix.nnz / (rating_matrix.shape[0] * rating_matrix.shape[1]) * 100 if rating_matrix.shape[0] * rating_matrix.shape[1] > 0 else 0
                            else:
                                mask = ~np.isnan(rating_matrix)
                                n_ratings = np.sum(mask)
                                density = (1 - np.isnan(rating_matrix).sum() / rating_matrix.size) * 100 if rating_matrix.size > 0 else 0
                            
                            st.success(f"✅ Veri yüklendi! {rating_matrix.shape[0]} kullanıcı, {rating_matrix.shape[1]} ürün")
                            st.info(f"📊 Veri yoğunluğu: {density:.2f}%")
                            st.rerun()
                            
            except Exception as e:
                st.error(f"❌ Hata: {str(e)}")
                st.info("💡 Lütfen veri formatını kontrol edin.")
    else:
        # Örnek veri oluştur - session state'i temizle
        if st.session_state.fm_rating_matrix is not None:
            st.session_state.fm_rating_matrix = None
            st.session_state.fm_user_mapping = None
            st.session_state.fm_item_mapping = None
    
    # Model parametreleri
    if rating_matrix is not None:
        n_users, n_items = rating_matrix.shape
        st.info(f"📊 Yüklenen veri: {n_users} kullanıcı, {n_items} ürün")
    else:
        col1, col2 = st.columns(2)
        with col1:
            n_users = st.slider("Kullanıcı Sayısı", 50, 500, 100, key="fm_n_users")
            n_items = st.slider("Ürün Sayısı", 30, 300, 50, key="fm_n_items")
        with col2:
            sparsity = st.slider("Sparsity", 0.3, 0.9, 0.6, key="fm_sparsity")
    
    st.info("💡 Bu demo için basit context features oluşturulacak.")
    
    # Optimal parametreleri al
    optimal_params = get_optimal_model_params("fm")
    optimal_epochs = optimal_params['epochs']
    
    n_factors = st.slider("Faktör Sayısı", 5, 50, 10, key="fm_n_factors")
    epochs = st.slider(
        "Epochs", 
        5, 30, 
        optimal_epochs, 
        key="fm_epochs",
        help=f"Önerilen değer: {optimal_epochs}"
    )
    
    if st.button("🚀 FM Modeli Eğit", key="fm_train"):
        if rating_matrix is None and data_source == "📁 Dosyadan Yükle":
            st.warning("⚠️ Lütfen önce veri dosyasını yükleyin!")
        else:
            with st.spinner("Veri hazırlanıyor..."):
                if rating_matrix is None:
                    rating_matrix = generate_rating_matrix(n_users, n_items, sparsity)
            
            # Context features oluştur (örnek: saat, cihaz tipi)
            user_ids = []
            item_ids = []
            ratings = []
            context_features = []
            
            # Sparse matrix desteği
            from scipy.sparse import issparse
            if issparse(rating_matrix):
                # Sparse matrix için - sadece mevcut rating'leri al
                rows, cols = rating_matrix.nonzero()
                user_ids = rows.tolist()
                item_ids = cols.tolist()
                ratings = rating_matrix.data.tolist()
                # Context features ekle
                np.random.seed(42)
                for _ in range(len(user_ids)):
                    context_features.append([
                        np.random.rand(),  # Örnek: saat (normalize)
                        np.random.rand()   # Örnek: cihaz tipi
                    ])
            else:
                # Dense matrix için
                np.random.seed(42)
                for u in range(n_users):
                    for i in range(n_items):
                        if not np.isnan(rating_matrix[u, i]):
                            user_ids.append(u)
                            item_ids.append(i)
                            ratings.append(rating_matrix[u, i])
                            # Basit context: rastgele özellikler
                            context_features.append([
                                np.random.rand(),  # Örnek: saat (normalize)
                                np.random.rand()   # Örnek: cihaz tipi
                            ])
            
            user_ids = np.array(user_ids)
            item_ids = np.array(item_ids)
            ratings = np.array(ratings)
            context_features = np.array(context_features)
        
            with st.spinner("FM modeli eğitiliyor..."):
                try:
                    import time
                    training_start = time.time()
                    
                    fm_model = FactorizationMachine(n_factors=n_factors)
                    history = fm_model.fit(
                        user_ids, item_ids, ratings, context_features,
                        epochs=epochs,
                        verbose=0
                    )
                    
                    training_time = time.time() - training_start
                    
                    st.success("✅ Model eğitildi!")
                    
                    # Sonuç açıklaması
                    with st.expander("📝 Sonuç Açıklaması - Ne Elde Edildi?", expanded=True):
                        st.markdown("### 🔍 Kullanılan Veriler")
                        if data_source == "📁 Dosyadan Yükle" and 'fm_file_name' in st.session_state:
                            from scipy.sparse import issparse as issparse_check
                            if issparse_check(rating_matrix):
                                n_ratings = rating_matrix.nnz
                            else:
                                mask = ~np.isnan(rating_matrix)
                                n_ratings = np.sum(mask)
                            
                            st.info(f"""
                            **Dosya**: {st.session_state.fm_file_name} ({st.session_state.fm_file_size / 1024:.2f} KB)
                            - **Kullanıcı Sayısı**: {n_users:,}
                            - **Ürün Sayısı**: {n_items:,}
                            - **Toplam Rating**: {n_ratings:,}
                            - **Context Features**: 2 (saat, cihaz tipi - otomatik oluşturuldu)
                            """)
                        else:
                            st.info(f"""
                            **Örnek Veri**:
                            - **Kullanıcı Sayısı**: {n_users:,}
                            - **Ürün Sayısı**: {n_items:,}
                            - **Context Features**: 2 (saat, cihaz tipi - otomatik oluşturuldu)
                            """)
                        
                        st.markdown("### ⚙️ Model Parametreleri")
                        st.info(f"""
                        - **Faktör Sayısı**: {n_factors}
                        - **Epochs**: {epochs}
                        - **Eğitim Süresi**: {training_time:.2f} saniye
                        """)
                        
                        st.markdown("### 📈 Elde Edilen Sonuçlar")
                        st.success(f"""
                        **FM Modeli Başarıyla Eğitildi!**
                        
                        **Ne Yapıldı?**
                        1. **Feature Engineering**: Kullanıcı-ürün etkileşimleri + context features (saat, cihaz tipi) birleştirildi
                        2. **Factorization**: {n_factors} boyutlu latent factors öğrenildi
                        3. **Pairwise Interactions**: Tüm feature çiftleri arasındaki etkileşimler modellendi
                        4. **Eğitim**: {epochs} epoch boyunca model optimize edildi
                        
                        **SVD'den Farkı:**
                        - SVD sadece kullanıcı-ürün ID'lerini kullanır
                        - FM context features (saat, cihaz, konum vb.) kullanır
                        - Daha zengin özellik seti ile daha iyi tahmin
                        - Context-aware öneriler sağlar
                        """)
                    
                    st.info("💡 FM, context features kullanarak daha kişiselleştirilmiş öneriler sağlar.")
                    
                except Exception as e:
                    st.error(f"❌ Hata: {str(e)}")
                    st.info("💡 PyTorch yüklü olduğundan emin olun: `pip install torch`")


def show_deepfm_recommender():
    """DeepFM sayfası"""
    st.header("🚀 DeepFM - Factorization Machines + Deep Learning")
    
    with st.expander("ℹ️ DeepFM Nedir?", expanded=False):
        st.markdown("""
        **DeepFM**, FM ve Deep Learning'i birleştirir.
        
        **Avantajları:**
        - ✅ **FM Component**: Doğrusal ve çift etkileşimleri yakalar
        - ✅ **Deep Component**: Doğrusal olmayan karmaşık pattern'leri öğrenir
        - ✅ Her iki yaklaşımın güçlü yönlerini birleştirir
        
        **Kullanım Alanları:**
        - Büyük ölçekli öneri sistemleri
        - CTR tahmini
        """)
    
    # Veri yükleme seçeneği - FM ile aynı yapı
    data_source = st.radio(
        "Veri Kaynağı",
        ["📊 Örnek Veri Oluştur", "📁 Dosyadan Yükle"],
        horizontal=True,
        key="deepfm_data_source"
    )
    
    # Session state ile veriyi koru
    if 'deepfm_rating_matrix' not in st.session_state:
        st.session_state.deepfm_rating_matrix = None
    
    rating_matrix = st.session_state.deepfm_rating_matrix
    n_users = None
    n_items = None
    
    if data_source == "📁 Dosyadan Yükle":
        st.markdown("### 📁 Veri Dosyası Yükle")
        st.info("""
        **Desteklenen Formatlar:** CSV, Excel (.xlsx, .xls)
        
        **Veri Formatı:** Long Format (user_id, item_id, rating) veya Matrix Format
        
        **Not:** DeepFM için context features otomatik oluşturulacak.
        """)
        
        file = st.file_uploader(
            "Veri dosyasını seçin",
            type=['csv', 'xlsx', 'xls'],
            help="CSV veya Excel dosyası yükleyin",
            key="deepfm_file"
        )
        
        if file is not None:
            try:
                import io
                file_content = file.read()
                st.session_state.deepfm_file_content = file_content
                file_bytes = io.BytesIO(file_content)
                
                import pandas as pd
                if file.name.endswith('.csv'):
                    file_bytes.seek(0)
                    first_line = file_bytes.readline().decode('utf-8', errors='ignore')
                    delimiters = [',', ';', '\t', '|']
                    detected_delimiter = ','
                    for delim in delimiters:
                        if len(first_line.split(delim)) > len(first_line.split(detected_delimiter)):
                            detected_delimiter = delim
                    file_bytes.seek(0)
                    preview_df = pd.read_csv(file_bytes, nrows=5, sep=detected_delimiter, engine='python')
                elif file.name.endswith(('.xlsx', '.xls')):
                    file_bytes.seek(0)
                    preview_df = pd.read_excel(file_bytes, nrows=5)
                else:
                    preview_df = None
                
                if preview_df is not None:
                    with st.expander("👁️ Dosya Önizleme (İlk 5 Satır)", expanded=False):
                        st.dataframe(preview_df, width='stretch')
                
                data_format = st.radio(
                    "Veri Formatı",
                    ["Long Format (user_id, item_id, rating)", "Matrix Format (Rating Matrisi)"],
                    horizontal=True,
                    key="deepfm_format"
                )
                
                if data_format == "Long Format (user_id, item_id, rating)":
                    if st.button("📥 Veriyi Yükle", key="deepfm_load_long"):
                        with st.spinner("Veri yükleniyor..."):
                            file_name = file.name
                            file_size = file.size
                            
                            if 'deepfm_file_content' in st.session_state:
                                file_bytes = io.BytesIO(st.session_state.deepfm_file_content)
                                file_bytes.name = file_name
                                rating_matrix, _, _ = load_rating_data_from_file(file_bytes)
                            else:
                                file.seek(0)
                                rating_matrix, _, _ = load_rating_data_from_file(file)
                            
                            st.session_state.deepfm_file_name = file_name
                            st.session_state.deepfm_file_size = file_size
                            st.session_state.deepfm_rating_matrix = rating_matrix
                            st.success(f"✅ Veri yüklendi! {rating_matrix.shape[0]} kullanıcı, {rating_matrix.shape[1]} ürün")
                            st.rerun()
                else:
                    if st.button("📥 Veriyi Yükle", key="deepfm_load_matrix"):
                        with st.spinner("Veri yükleniyor..."):
                            file_name = file.name
                            file_size = file.size
                            
                            if 'deepfm_file_content' in st.session_state:
                                file_bytes = io.BytesIO(st.session_state.deepfm_file_content)
                                file_bytes.name = file_name
                                rating_matrix = load_rating_matrix_from_file(file_bytes)
                            else:
                                file.seek(0)
                                rating_matrix = load_rating_matrix_from_file(file)
                            
                            st.session_state.deepfm_file_name = file_name
                            st.session_state.deepfm_file_size = file_size
                            st.session_state.deepfm_rating_matrix = rating_matrix
                            st.success(f"✅ Veri yüklendi! {rating_matrix.shape[0]} kullanıcı, {rating_matrix.shape[1]} ürün")
                            st.rerun()
                            
            except Exception as e:
                st.error(f"❌ Hata: {str(e)}")
    else:
        if st.session_state.deepfm_rating_matrix is not None:
            st.session_state.deepfm_rating_matrix = None
    
    # Model parametreleri
    if rating_matrix is not None:
        n_users, n_items = rating_matrix.shape
        st.info(f"📊 Yüklenen veri: {n_users} kullanıcı, {n_items} ürün")
    else:
        col1, col2 = st.columns(2)
        with col1:
            n_users = st.slider("Kullanıcı Sayısı", 50, 500, 100, key="deepfm_n_users")
            n_items = st.slider("Ürün Sayısı", 30, 300, 50, key="deepfm_n_items")
        with col2:
            sparsity = st.slider("Sparsity", 0.3, 0.9, 0.6, key="deepfm_sparsity")
    
    st.info("💡 DeepFM, FM'nin gelişmiş versiyonudur.")
    
    # Optimal parametreleri al
    optimal_params = get_optimal_model_params("deepfm")
    optimal_epochs = optimal_params['epochs']
    
    n_factors = st.slider("FM Faktör Sayısı", 5, 50, 10, key="deepfm_n_factors")
    epochs = st.slider(
        "Epochs", 
        5, 30, 
        optimal_epochs, 
        key="deepfm_epochs",
        help=f"Önerilen değer: {optimal_epochs}"
    )
    
    if st.button("🚀 DeepFM Modeli Eğit", key="deepfm_train"):
        if rating_matrix is None and data_source == "📁 Dosyadan Yükle":
            st.warning("⚠️ Lütfen önce veri dosyasını yükleyin!")
        else:
            with st.spinner("Veri hazırlanıyor..."):
                if rating_matrix is None:
                    rating_matrix = generate_rating_matrix(n_users, n_items, sparsity)
            
            user_ids = []
            item_ids = []
            ratings = []
            context_features = []
            
            # Sparse matrix desteği
            from scipy.sparse import issparse
            if issparse(rating_matrix):
                # Sparse matrix için - sadece mevcut rating'leri al
                rows, cols = rating_matrix.nonzero()
                user_ids = rows.tolist()
                item_ids = cols.tolist()
                ratings = rating_matrix.data.tolist()
                # Context features ekle
                for _ in range(len(user_ids)):
                    context_features.append([np.random.rand(), np.random.rand()])
            else:
                # Dense matrix için
                for u in range(n_users):
                    for i in range(n_items):
                        if not np.isnan(rating_matrix[u, i]):
                            user_ids.append(u)
                            item_ids.append(i)
                            ratings.append(rating_matrix[u, i])
                            context_features.append([np.random.rand(), np.random.rand()])
            
            user_ids = np.array(user_ids)
            item_ids = np.array(item_ids)
            ratings = np.array(ratings)
            context_features = np.array(context_features)
        
            with st.spinner("DeepFM modeli eğitiliyor..."):
                try:
                    import time
                    training_start = time.time()
                    
                    deepfm_model = DeepFM(n_factors=n_factors)
                    history = deepfm_model.fit(
                        user_ids, item_ids, ratings, context_features,
                        epochs=epochs,
                        verbose=0
                    )
                    
                    training_time = time.time() - training_start
                    
                    st.success("✅ Model eğitildi!")
                    
                    # Sonuç açıklaması
                    with st.expander("📝 Sonuç Açıklaması - Ne Elde Edildi?", expanded=True):
                        st.markdown("### 🔍 Kullanılan Veriler")
                        if data_source == "📁 Dosyadan Yükle" and 'deepfm_file_name' in st.session_state:
                            from scipy.sparse import issparse as issparse_check
                            if issparse_check(rating_matrix):
                                n_ratings = rating_matrix.nnz
                            else:
                                mask = ~np.isnan(rating_matrix)
                                n_ratings = np.sum(mask)
                            
                            st.info(f"""
                            **Dosya**: {st.session_state.deepfm_file_name} ({st.session_state.deepfm_file_size / 1024:.2f} KB)
                            - **Kullanıcı Sayısı**: {n_users:,}
                            - **Ürün Sayısı**: {n_items:,}
                            - **Toplam Rating**: {n_ratings:,}
                            - **Context Features**: 2 (otomatik oluşturuldu)
                            """)
                        else:
                            st.info(f"""
                            **Örnek Veri**:
                            - **Kullanıcı Sayısı**: {n_users:,}
                            - **Ürün Sayısı**: {n_items:,}
                            """)
                        
                        st.markdown("### ⚙️ Model Parametreleri")
                        st.info(f"""
                        - **FM Faktör Sayısı**: {n_factors}
                        - **Epochs**: {epochs}
                        - **Eğitim Süresi**: {training_time:.2f} saniye
                        """)
                        
                        st.markdown("### 📈 Elde Edilen Sonuçlar")
                        st.success(f"""
                        **DeepFM Modeli Başarıyla Eğitildi!**
                        
                        **Ne Yapıldı?**
                        1. **FM Component**: Doğrusal ve pairwise etkileşimler öğrenildi
                        2. **Deep Component**: Derin sinir ağı ile doğrusal olmayan pattern'ler öğrenildi
                        3. **Birleşik Model**: FM + Deep Learning birleştirildi
                        4. **Eğitim**: {epochs} epoch boyunca model optimize edildi
                        
                        **FM'den Farkı:**
                        - FM sadece doğrusal ve pairwise etkileşimleri yakalar
                        - DeepFM ek olarak derin sinir ağı ile karmaşık pattern'leri öğrenir
                        - Daha güçlü özellik öğrenme kapasitesi
                        """)
                    
                    st.info("💡 DeepFM, hem doğrusal hem de doğrusal olmayan özellikleri öğrenir.")
                    
                except Exception as e:
                    st.error(f"❌ Hata: {str(e)}")
                    st.info("💡 PyTorch yüklü olduğundan emin olun: `pip install torch`")


def show_transformer_recommender():
    """Transformer-based recommendation sayfası"""
    st.header("🔄 Transformer - Sequential Recommendation")
    
    with st.expander("ℹ️ Transformer Nedir?", expanded=False):
        st.markdown("""
        **Transformer**, ChatGPT mimarisinin öneri sistemlerine uyarlanmış halidir.
        
        **Özellikleri:**
        - ✅ **Zaman ve sıra** bilgisini kullanır (SVD bunu yapamaz)
        - ✅ **Self-attention** mekanizması ile uzun mesafe bağımlılıkları yakalar
        - ✅ "Sıradaki ne?" sorusunu cevaplar
        
        **Kullanım Alanları:**
        - TikTok, YouTube gibi sequential öneriler
        - E-ticaret sepet önerileri
        """)
    
    col1, col2 = st.columns(2)
    with col1:
        n_items = st.slider("Toplam Ürün Sayısı", 50, 500, 100)
        max_seq_length = st.slider("Maksimum Sequence Uzunluğu", 10, 100, 50)
    with col2:
        d_model = st.slider("Model Boyutu", 64, 256, 128)
        n_heads = st.slider("Attention Head Sayısı", 2, 8, 4)
    
    epochs = st.slider("Epochs", 5, 30, 10)
    
    if st.button("🚀 Transformer Modeli Eğit"):
        with st.spinner("Sequential veri oluşturuluyor..."):
            # Örnek sequential veri (her kullanıcı için item sequence)
            n_users = 50
            user_sequences = []
            for u in range(n_users):
                # Her kullanıcı için rastgele item sequence
                seq_length = np.random.randint(5, max_seq_length)
                sequence = np.random.choice(n_items, size=seq_length, replace=False).tolist()
                user_sequences.append(sequence)
        
        with st.spinner("Transformer modeli eğitiliyor..."):
            try:
                transformer_model = TransformerRecommender(
                    n_items=n_items,
                    d_model=d_model,
                    n_heads=n_heads,
                    max_seq_length=max_seq_length
                )
                history = transformer_model.fit(
                    user_sequences,
                    epochs=epochs,
                    verbose=0
                )
                
                st.success("✅ Model eğitildi!")
                
                # Örnek tahmin
                st.subheader("Sıradaki Item Tahmini")
                example_user = st.selectbox("Kullanıcı Seçin", range(min(10, n_users)))
                example_sequence = user_sequences[example_user][:-1]  # Son item hariç
                
                item_indices, probabilities = transformer_model.predict_next(example_sequence)
                
                recommendations_df = pd.DataFrame({
                    'Sıradaki Ürün ID': item_indices + 1,
                    'Olasılık': np.round(probabilities, 4)
                })
                st.dataframe(recommendations_df, width='stretch')
                
                st.info(f"💡 Kullanıcının geçmiş sequence'i: {[x+1 for x in example_sequence[-5:]]}")
                
            except Exception as e:
                st.error(f"❌ Hata: {str(e)}")


def show_gnn_recommender():
    """GNN (Graph Neural Network) sayfası"""
    st.header("🕸️ Graph Neural Network (GNN) - Öneri Sistemi")
    
    with st.expander("ℹ️ GNN Nedir?", expanded=False):
        st.markdown("""
        **Graph Neural Network**, veriyi tablo değil, **ağ (graph)** olarak görür.
        
        **Özellikleri:**
        - ✅ Kullanıcı-ürün ilişkilerini graph olarak modeler
        - ✅ **İlişkisel veriyi** en iyi işleyen yöntem
        - ✅ Arkadaşlık ağları, ürün benzerlik ağları kullanır
        
        **Kullanım Alanları:**
        - Pinterest, Uber Eats
        - Sosyal ağ tabanlı öneriler
        - İlişkisel veri analizi
        """)
    
    # PyTorch kontrolü
    if not MODERN_AVAILABLE:
        st.warning("⚠️ GNN için PyTorch ve PyTorch Geometric gerekli!")
        st.info("💡 Lütfen şu komutu çalıştırın: `pip install torch torch-geometric`")
        st.stop()
    else:
        st.success("✅ PyTorch ve PyTorch Geometric yüklü - GNN kullanılabilir!")
    
    # Veri yükleme seçeneği
    data_source = st.radio(
        "Veri Kaynağı",
        ["📊 Örnek Veri Oluştur", "📁 Dosyadan Yükle"],
        horizontal=True,
        key="gnn_data_source"
    )
    
    # Session state ile veriyi koru
    if 'gnn_rating_matrix' not in st.session_state:
        st.session_state.gnn_rating_matrix = None
    
    rating_matrix = st.session_state.gnn_rating_matrix
    n_users = None
    n_items = None
    
    if data_source == "📁 Dosyadan Yükle":
        st.markdown("### 📁 Veri Dosyası Yükle")
        st.info("""
        **Desteklenen Formatlar:** CSV, Excel (.xlsx, .xls)
        
        **Veri Formatı:** Long Format (user_id, item_id, rating) veya Matrix Format
        """)
        
        file = st.file_uploader(
            "Veri dosyasını seçin",
            type=['csv', 'xlsx', 'xls'],
            help="CSV veya Excel dosyası yükleyin",
            key="gnn_file"
        )
        
        if file is not None:
            try:
                import io
                file_content = file.read()
                st.session_state.gnn_file_content = file_content
                
                data_format = st.radio(
                    "Veri Formatı",
                    ["Long Format (user_id, item_id, rating)", "Matrix Format (Rating Matrisi)"],
                    horizontal=True,
                    key="gnn_format"
                )
                
                if data_format == "Long Format (user_id, item_id, rating)":
                    if st.button("📥 Veriyi Yükle", key="gnn_load_long"):
                        with st.spinner("Veri yükleniyor..."):
                            file_name = file.name
                            file_size = file.size
                            
                            if 'gnn_file_content' in st.session_state:
                                file_bytes = io.BytesIO(st.session_state.gnn_file_content)
                                file_bytes.name = file_name
                                rating_matrix, _, _ = load_rating_data_from_file(file_bytes)
                            else:
                                file.seek(0)
                                rating_matrix, _, _ = load_rating_data_from_file(file)
                            
                            st.session_state.gnn_file_name = file_name
                            st.session_state.gnn_file_size = file_size
                            st.session_state.gnn_rating_matrix = rating_matrix
                            st.success(f"✅ Veri yüklendi! {rating_matrix.shape[0]} kullanıcı, {rating_matrix.shape[1]} ürün")
                            st.rerun()
                else:
                    if st.button("📥 Veriyi Yükle", key="gnn_load_matrix"):
                        with st.spinner("Veri yükleniyor..."):
                            file_name = file.name
                            file_size = file.size
                            
                            if 'gnn_file_content' in st.session_state:
                                file_bytes = io.BytesIO(st.session_state.gnn_file_content)
                                file_bytes.name = file_name
                                rating_matrix = load_rating_matrix_from_file(file_bytes)
                            else:
                                file.seek(0)
                                rating_matrix = load_rating_matrix_from_file(file)
                            
                            st.session_state.gnn_file_name = file_name
                            st.session_state.gnn_file_size = file_size
                            st.session_state.gnn_rating_matrix = rating_matrix
                            st.success(f"✅ Veri yüklendi! {rating_matrix.shape[0]} kullanıcı, {rating_matrix.shape[1]} ürün")
                            st.rerun()
                            
            except Exception as e:
                st.error(f"❌ Hata: {str(e)}")
    else:
        if st.session_state.gnn_rating_matrix is not None:
            st.session_state.gnn_rating_matrix = None
    
    # Model parametreleri
    if rating_matrix is not None:
        n_users, n_items = rating_matrix.shape
        st.info(f"📊 Yüklenen veri: {n_users} kullanıcı, {n_items} ürün")
    else:
        col1, col2 = st.columns(2)
        with col1:
            n_users = st.slider("Kullanıcı Sayısı", 50, 300, 100, key="gnn_n_users")
            n_items = st.slider("Ürün Sayısı", 30, 200, 50, key="gnn_n_items")
        with col2:
            sparsity = st.slider("Sparsity", 0.3, 0.9, 0.6, key="gnn_sparsity")
    
    # Optimal parametreleri al
    optimal_params = get_optimal_model_params("gnn")
    optimal_embedding_dim = optimal_params['embedding_dim']
    optimal_epochs = optimal_params['epochs']
    
    embedding_dim = st.slider(
        "Embedding Boyutu", 
        32, 128, 
        optimal_embedding_dim, 
        key="gnn_embedding",
        help=f"Önerilen değer: {optimal_embedding_dim}"
    )
    epochs = st.slider(
        "Epochs", 
        10, 100, 
        optimal_epochs, 
        key="gnn_epochs",
        help=f"Önerilen değer: {optimal_epochs}"
    )
    
    if st.button("🚀 GNN Modeli Eğit", key="gnn_train"):
        if rating_matrix is None and data_source == "📁 Dosyadan Yükle":
            st.warning("⚠️ Lütfen önce veri dosyasını yükleyin!")
        else:
            with st.spinner("Veri hazırlanıyor..."):
                if rating_matrix is None:
                    rating_matrix = generate_rating_matrix(n_users, n_items, sparsity)
            
            with st.spinner("GNN modeli eğitiliyor (bu biraz zaman alabilir)..."):
                try:
                    import time
                    training_start = time.time()
                    
                    gnn_model = GNNRecommender(embedding_dim=embedding_dim)
                    gnn_model.fit(
                        rating_matrix,
                        epochs=epochs,
                        verbose=False
                    )
                    
                    training_time = time.time() - training_start
                    
                    st.success("✅ Model eğitildi!")
                    
                    # Sonuç açıklaması
                    with st.expander("📝 Sonuç Açıklaması - Ne Elde Edildi?", expanded=True):
                        st.markdown("### 🔍 Kullanılan Veriler")
                        if data_source == "📁 Dosyadan Yükle" and 'gnn_file_name' in st.session_state:
                            from scipy.sparse import issparse as issparse_check
                            if issparse_check(rating_matrix):
                                n_ratings = rating_matrix.nnz
                            else:
                                mask = ~np.isnan(rating_matrix)
                                n_ratings = np.sum(mask)
                            
                            st.info(f"""
                            **Dosya**: {st.session_state.gnn_file_name} ({st.session_state.gnn_file_size / 1024:.2f} KB)
                            - **Kullanıcı Sayısı**: {n_users:,}
                            - **Ürün Sayısı**: {n_items:,}
                            - **Toplam Rating**: {n_ratings:,}
                            """)
                        else:
                            st.info(f"""
                            **Örnek Veri**:
                            - **Kullanıcı Sayısı**: {n_users:,}
                            - **Ürün Sayısı**: {n_items:,}
                            """)
                        
                        st.markdown("### ⚙️ Model Parametreleri")
                        st.info(f"""
                        - **Embedding Boyutu**: {embedding_dim}
                        - **Epochs**: {epochs}
                        - **Eğitim Süresi**: {training_time:.2f} saniye
                        """)
                        
                        st.markdown("### 📈 Elde Edilen Sonuçlar")
                        st.success(f"""
                        **GNN Modeli Başarıyla Eğitildi!**
                        
                        **Ne Yapıldı?**
                        1. **Graph Oluşturma**: Kullanıcı-ürün etkileşimleri graph yapısına dönüştürüldü
                        2. **Node Embedding**: Her kullanıcı ve ürün {embedding_dim} boyutlu embedding'e dönüştürüldü
                        3. **Message Passing**: Graph üzerinde mesaj geçişi ile komşu bilgileri toplandı
                        4. **Eğitim**: {epochs} epoch boyunca model optimize edildi
                        
                        **SVD'den Farkı:**
                        - SVD rating matrisini direkt faktörize eder
                        - GNN kullanıcı-ürün ilişkilerini graph olarak modeler
                        - İlişkisel veriyi daha iyi işler
                        - Graph yapısından faydalanır
                        """)
                    
                    # Öneriler
                    st.subheader("🎯 Kullanıcı Önerileri")
                    user_idx = st.selectbox("Kullanıcı Seçin", range(min(10, n_users)), key="gnn_user_select")
                    
                    item_indices, predicted_ratings = gnn_model.recommend(
                        user_idx, n_recommendations=10, rating_matrix=rating_matrix
                    )
                    
                    recommendations_df = pd.DataFrame({
                        'Ürün ID': item_indices + 1,
                        'Tahmin Edilen Rating': np.round(np.clip(predicted_ratings, 1, 5), 2)
                    })
                    st.dataframe(recommendations_df, width='stretch')
                    
                    st.info("💡 GNN, kullanıcı-ürün ilişkilerini graph olarak modelleyerek öneriler üretir.")
                    
                except ImportError:
                    st.error("❌ PyTorch ve PyTorch Geometric yüklü değil!")
                    st.info("💡 Lütfen şu komutu çalıştırın: `pip install torch torch-geometric`")
                except Exception as e:
                    st.error(f"❌ Hata: {str(e)}")
                    st.info("💡 PyTorch yüklü olduğundan emin olun: `pip install torch torch-geometric`")


def show_ai_chatbot():
    """
    AI Chatbot - Veri Asistanı sayfası
    PandasAI kullanarak doğal dil ile veri analizi yapılmasını sağlar
    """
    st.header("🤖 AI Chat - Veri Asistanı")
    
    with st.expander("ℹ️ AI Chat Nedir?", expanded=False):
        st.markdown("""
        **AI Chat - Veri Asistanı**, doğal dil kullanarak veri analizi yapmanızı sağlar.
        
        **Özellikler:**
        - ✅ Excel/CSV dosyalarını yükleyin
        - ✅ Doğal dil ile sorular sorun: "En çok puan veren kullanıcı kim?"
        - ✅ Grafikler oluşturun: "Rating dağılımını çiz"
        - ✅ Veri temizleme: "Boş verileri temizle"
        - ✅ Özel komutlar: "SVD çalıştır" gibi matrix factorization işlemleri
        
        **Kullanım:**
        1. OpenAI API Key'inizi girin (sidebar)
        2. Veri dosyanızı yükleyin
        3. Sorularınızı sorun!
        """)
    
    # Sidebar - API Key girişi
    st.sidebar.markdown("---")
    st.sidebar.subheader("🔑 API Ayarları")
    
    # API Provider seçimi
    api_provider = st.sidebar.radio(
        "AI Provider Seçin",
        ["OpenAI", "Google Gemini"],
        help="Kullanmak istediğiniz AI servisini seçin",
        key="ai_provider"
    )
    
    if api_provider == "OpenAI":
        api_key = st.sidebar.text_input(
            "OpenAI API Key",
            type="password",
            help="OpenAI API key'inizi buraya girin. https://platform.openai.com/api-keys adresinden alabilirsiniz.",
            key="openai_api_key"
        )
        
        if not api_key:
            st.warning("⚠️ Lütfen OpenAI API Key'inizi sidebar'dan girin.")
            st.info("💡 API Key almak için: https://platform.openai.com/api-keys")
            st.stop()
    else:  # Gemini
        api_key = st.sidebar.text_input(
            "Google Gemini API Key",
            type="password",
            help="Google Gemini API key'inizi buraya girin. https://aistudio.google.com/app/apikey adresinden alabilirsiniz.",
            key="gemini_api_key"
        )
        
        if not api_key:
            st.warning("⚠️ Lütfen Google Gemini API Key'inizi sidebar'dan girin.")
            st.info("💡 API Key almak için: https://aistudio.google.com/app/apikey")
            st.stop()
    
    # Dosya yükleme
    st.markdown("### 📁 Veri Dosyası Yükle")
    uploaded_file = st.file_uploader(
        "Excel veya CSV dosyası seçin",
        type=['csv', 'xlsx', 'xls'],
        help="Analiz yapmak istediğiniz veri dosyasını yükleyin"
    )
    
    # Session state ile chat geçmişini ve DataFrame'i koru
    if 'chat_messages' not in st.session_state:
        st.session_state.chat_messages = []
    
    if 'df' not in st.session_state:
        st.session_state.df = None
    
    if 'pandasai_agent' not in st.session_state:
        st.session_state.pandasai_agent = None
    
    # Dosya yüklendiğinde DataFrame'e çevir
    if uploaded_file is not None:
        try:
            import io
            import pandas as pd
            
            # Dosya tipine göre oku
            if uploaded_file.name.endswith('.csv'):
                # CSV için delimiter tespiti
                file_bytes = uploaded_file.read()
                first_line = file_bytes.decode('utf-8', errors='ignore').split('\n')[0]
                delimiters = [',', ';', '\t', '|']
                detected_delimiter = ','
                max_cols = 0
                for delim in delimiters:
                    cols = first_line.split(delim)
                    if len(cols) > max_cols:
                        max_cols = len(cols)
                        detected_delimiter = delim
                
                uploaded_file.seek(0)
                df = pd.read_csv(uploaded_file, sep=detected_delimiter, engine='python')
            elif uploaded_file.name.endswith(('.xlsx', '.xls')):
                df = pd.read_excel(uploaded_file)
            else:
                st.error("❌ Desteklenmeyen dosya formatı!")
                st.stop()
            
            # DataFrame'i session state'e kaydet
            st.session_state.df = df
            st.session_state.chat_messages = []  # Yeni dosya yüklendiğinde chat geçmişini temizle
            
            st.success(f"✅ Dosya yüklendi! {df.shape[0]} satır, {df.shape[1]} sütun")
            
            # Veri önizlemesi
            with st.expander("👁️ Veri Önizleme", expanded=False):
                st.dataframe(df.head(10), width='stretch')
                st.info(f"**Sütunlar**: {', '.join(df.columns.tolist()[:10])}{'...' if len(df.columns) > 10 else ''}")
            
        except Exception as e:
            st.error(f"❌ Dosya yüklenirken hata oluştu: {str(e)}")
            st.stop()
    
    # DataFrame yoksa uyarı göster
    if st.session_state.df is None:
        st.info("💡 Lütfen analiz yapmak için bir veri dosyası yükleyin.")
        st.stop()
    
    df = st.session_state.df
    
    # PandasAI agent'ı oluştur (sadece bir kez)
    if st.session_state.pandasai_agent is None:
        try:
            # PandasAI import - duplicate validator hatası olabilir
            try:
                from pandasai import SmartDataframe
            except (ValueError, TypeError) as import_error:
                error_str = str(import_error)
                if "duplicate validator" in error_str.lower() or "validate_llm" in error_str:
                    # Validator çakışması import sırasında oluştu
                    st.warning("⚠️ PandasAI validator çakışması (import sırasında)")
                    st.info("💡 Alternatif moda geçiliyor...")
                    st.session_state.pandasai_agent = "alternative_mode"
                    raise  # Dış exception handler'a geç
                else:
                    raise
            
            # LLM seçimi (OpenAI veya Gemini)
            if api_provider == "OpenAI":
                try:
                    from pandasai.llm import OpenAI
                    llm = OpenAI(api_token=api_key)
                except ImportError:
                    st.warning("⚠️ PandasAI OpenAI LLM import edilemedi. Alternatif moda geçiliyor...")
                    st.session_state.pandasai_agent = "alternative_mode"
            else:  # Gemini
                try:
                    from pandasai.llm import GoogleGemini
                    llm = GoogleGemini(api_key=api_key)
                except ImportError:
                    st.warning("⚠️ PandasAI Gemini LLM import edilemedi. Alternatif moda geçiliyor...")
                    st.session_state.pandasai_agent = "alternative_mode"
            
            # Eğer alternatif moda geçildiyse, burayı atla
            if st.session_state.pandasai_agent == "alternative_mode":
                pass  # Aşağıdaki alternatif mod kısmına geç
            else:
                # SmartDataframe oluştur (custom skills ile)
                # Not: Python 3.14 uyumluluk sorunları olabilir
                try:
                    smart_df = SmartDataframe(
                    df,
                    config={
                        "llm": llm,
                        "verbose": False,
                        "save_logs": False,
                        "enable_cache": False,
                        "custom_instructions": """
                        Sen bir Matris Faktörizasyon Uygulaması'nın AI asistanısın. Bu uygulama matris faktörizasyon algoritmaları için kapsamlı bir araçtır.
                        
                        UYGULAMA ÖZELLİKLERİ:
                        
                        📊 KLASİK ALGORİTMALAR:
                        1. SVD (Singular Value Decomposition) - Öneri sistemleri ve gürültü temizleme için
                        2. PCA (Principal Component Analysis) - Veri görselleştirme ve özellik seçimi için
                        3. NMF (Non-negative Matrix Factorization) - Görüntü işleme ve topic modeling için
                        4. ALS (Alternating Least Squares) - Büyük ölçekli öneri motorları için
                        
                        🚀 MODERN DEEP LEARNING ALGORİTMALAR (PyTorch gerekli):
                        5. NCF (Neural Collaborative Filtering) - Netflix/YouTube tarzı öneriler
                        6. Autoencoder (Denoising & VAE) - Gürültü temizleme ve öneriler
                        7. Factorization Machines (FM) & DeepFM - Context-aware öneriler
                        8. Transformer (BERT4Rec/SASRec) - Sequential öneriler (TikTok/YouTube tarzı)
                        9. GNN (Graph Neural Network) - Graph tabanlı öneriler (Pinterest/Uber Eats tarzı)
                        
                        KULLANILABİLİR FONKSİYONLAR:
                        - SVD analizi yapmak için: run_svd_analysis(n_components=5)
                        - Veri seti hakkında sorular sorulabilir
                        - Algoritma önerileri yapabilirsin
                        - Veri analizi ve görselleştirme önerileri sunabilirsin
                        
                        GÖREVİN:
                        - Kullanıcıların veri setleri hakkındaki sorularını yanıtla
                        - Hangi algoritmanın ne zaman kullanılacağını öner
                        - Matris faktörizasyon teknikleri hakkında bilgi ver
                        - Veri analizi için Python pandas kodları öner
                        - Gerekirse grafikler oluştur
                        - Kullanıcı "SVD çalıştır" veya "matrix factorization yap" derse, run_svd_analysis fonksiyonunu kullan
                        """
                    }
                    )
                except Exception as config_error:
                    error_str = str(config_error)
                    # Duplicate validator hatasını özel olarak yakala
                    if "duplicate validator" in error_str.lower() or "validate_llm" in error_str:
                        st.warning("⚠️ PandasAI validator çakışması tespit edildi. Çözüm deneniyor...")
                        try:
                            # Pydantic model cache'ini temizlemeyi dene
                            import pydantic
                            if hasattr(pydantic, 'BaseModel'):
                                # Pydantic v1 için
                                from pydantic import BaseModel
                                BaseModel.__config__ = None
                            # Basit config ile tekrar dene
                            smart_df = SmartDataframe(df, config={"llm": llm, "verbose": False})
                        except Exception as retry_error:
                            st.warning(f"⚠️ İkinci deneme başarısız: {str(retry_error)}")
                            st.info("💡 Alternatif moda geçiliyor...")
                            st.session_state.pandasai_agent = "alternative_mode"
                    else:
                        # Eğer config ile başka bir sorun varsa, basit config dene
                        st.warning(f"⚠️ PandasAI config hatası: {error_str}")
                        st.info("💡 Basit mod ile devam ediliyor...")
                        try:
                            smart_df = SmartDataframe(df, config={"llm": llm, "verbose": False})
                        except:
                            st.session_state.pandasai_agent = "alternative_mode"
            
                # Custom skills ekle - SVD çalıştırma
                def run_svd_analysis(n_components=5):
                    """
                    SVD (Singular Value Decomposition) analizi yapar
                    
                    Args:
                        n_components: Kullanılacak tekil değer sayısı
                    
                    Returns:
                        SVD sonuçları hakkında bilgi
                    """
                    try:
                        from algorithms.svd import SVDRecommender
                        from scipy.sparse import issparse
                        import numpy as np
                        
                        # DataFrame'i rating matrix'e çevir (sayısal sütunlar)
                        numeric_df = df.select_dtypes(include=[np.number])
                        
                        if numeric_df.empty:
                            return "Hata: DataFrame'de sayısal sütun bulunamadı!"
                        
                        # NaN değerleri 0 ile doldur
                        rating_matrix = numeric_df.fillna(0).values
                        
                        # SVD modeli oluştur ve eğit
                        svd_model = SVDRecommender(n_components=n_components)
                        svd_model.fit(rating_matrix)
                        
                        # Sonuçları döndür
                        explained_variance = svd_model.explained_variance_ratio_
                        total_variance = explained_variance.sum()
                        
                        result = f"""
                        SVD Analizi Tamamlandı!
                        
                        - Kullanılan Bileşen Sayısı: {n_components}
                        - Toplam Açıklanan Varyans: {total_variance:.2%}
                        - Bileşen Bazında Varyans:
                        """
                        for i, var in enumerate(explained_variance):
                            result += f"\n  - Bileşen {i+1}: {var:.2%}"
                        
                        return result
                    except Exception as e:
                        return f"SVD analizi sırasında hata: {str(e)}"
                
                # Custom skill'i ekle (eğer destekleniyorsa)
                try:
                    smart_df.add_skills([run_svd_analysis])
                except Exception as skill_error:
                    st.warning(f"⚠️ Custom skills eklenemedi: {str(skill_error)}")
                    st.info("💡 Temel özelliklerle devam ediliyor...")
                
                st.session_state.pandasai_agent = smart_df
                
                st.success("✅ AI Asistanı hazır!")
            
        except ImportError as import_err:
            st.error("❌ PandasAI yüklü değil veya import edilemedi!")
            st.info(f"💡 Hata: {str(import_err)}")
            st.info("💡 Alternatif mod aktif ediliyor...")
            # Alternatif moda geç
            st.session_state.pandasai_agent = "alternative_mode"
        except Exception as e:
            error_msg = str(e)
            # Duplicate validator hatasını özel olarak yakala
            if "duplicate validator" in error_msg.lower() or "validate_llm" in error_msg:
                st.warning("⚠️ PandasAI validator çakışması tespit edildi!")
                st.info("💡 Bu, Pydantic versiyon uyumsuzluğundan kaynaklanabilir.")
                st.info("💡 Alternatif moda geçiliyor...")
                st.session_state.pandasai_agent = "alternative_mode"
            elif "Python 3.14" in error_msg or "pydantic" in error_msg.lower() or "default_factory" in error_msg:
                st.warning("⚠️ PandasAI Python 3.14 ile uyumlu değil!")
                st.info("💡 Alternatif OpenAI modu aktif ediliyor...")
                # Alternatif moda geç
                st.session_state.pandasai_agent = "alternative_mode"
            else:
                st.error(f"❌ AI Asistanı oluşturulurken hata: {error_msg}")
                st.info("💡 Alternatif mod denenecek...")
                st.session_state.pandasai_agent = "alternative_mode"
        
        # Alternatif mod: Direkt API kullan (OpenAI veya Gemini)
        if st.session_state.pandasai_agent == "alternative_mode":
            try:
                if api_provider == "OpenAI":
                    from openai import OpenAI as OpenAIClient
                    client = OpenAIClient(api_key=api_key)
                    model_name = "gpt-4o-mini"
                    provider_name = "OpenAI"
                else:  # Gemini
                    try:
                        import google.generativeai as genai
                        genai.configure(api_key=api_key)
                        client = genai
                        # Güncel Gemini model adları - önce mevcut modelleri kontrol et
                        try:
                            # Mevcut modelleri listele ve uygun olanı seç
                            available_models = []
                            for model in genai.list_models():
                                if 'generateContent' in model.supported_generation_methods:
                                    # Model adını normalize et (models/ prefix'ini kaldır)
                                    model_name_clean = model.name.replace("models/", "") if model.name.startswith("models/") else model.name
                                    available_models.append(model_name_clean)
                            
                            # Öncelik sırası: flash, pro, genel
                            if available_models:
                                if any('flash' in model.lower() for model in available_models):
                                    model_name = next((m for m in available_models if 'flash' in m.lower()), "gemini-1.5-flash")
                                elif any('pro' in model.lower() for model in available_models):
                                    model_name = next((m for m in available_models if 'pro' in m.lower()), "gemini-1.5-pro")
                                else:
                                    model_name = available_models[0]
                            else:
                                model_name = "gemini-1.5-flash"  # Varsayılan
                        except Exception as list_err:
                            # Eğer model listesi alınamazsa, güncel model adlarını dene
                            model_name = "gemini-1.5-flash"  # Varsayılan
                        provider_name = "Google Gemini"
                    except ImportError:
                        st.error("❌ Google Generative AI kütüphanesi yüklü değil!")
                        st.info("💡 Lütfen şu komutu çalıştırın: `python -m pip install google-generativeai`")
                        st.stop()
                
                st.session_state.pandasai_agent = {
                    "type": "direct_api",
                    "client": client,
                    "df": df,
                    "provider": api_provider,
                    "model": model_name
                }
                st.success(f"✅ AI Asistanı hazır! ({provider_name} Direct Mode)")
                st.info(f"💡 Bu mod PandasAI yerine direkt {provider_name} API kullanır.")
            except Exception as alt_err:
                st.error(f"❌ Alternatif mod da başarısız: {str(alt_err)}")
                st.stop()
    
    # Chat arayüzü
    st.markdown("### 💬 Sohbet")
    
    # Chat geçmişini göster
    for message in st.session_state.chat_messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
            
            # Eğer grafik varsa göster
            if "figure" in message:
                st.pyplot(message["figure"])
    
    # Kullanıcı girişi
    user_query = st.chat_input("Veri hakkında bir soru sorun... (örn: 'En çok puan veren kullanıcı kim?')")
    
    if user_query:
        # Kullanıcı mesajını ekle
        st.session_state.chat_messages.append({
            "role": "user",
            "content": user_query
        })
        
        # Kullanıcı mesajını göster
        with st.chat_message("user"):
            st.markdown(user_query)
        
        # AI yanıtını oluştur
        with st.chat_message("assistant"):
            with st.spinner("🤔 Düşünüyorum..."):
                try:
                    # PandasAI veya alternatif mod kontrolü
                    if isinstance(st.session_state.pandasai_agent, dict) and st.session_state.pandasai_agent.get("type") == "direct_api":
                        # Alternatif mod: Direkt API (OpenAI veya Gemini)
                        client = st.session_state.pandasai_agent["client"]
                        df_for_analysis = st.session_state.pandasai_agent["df"]
                        provider = st.session_state.pandasai_agent["provider"]
                        model_name = st.session_state.pandasai_agent["model"]
                        
                        # DataFrame'i string formatına çevir (ilk 100 satır)
                        df_preview = df_for_analysis.head(100).to_string()
                        
                        # Uygulama hakkında kapsamlı bilgi
                        app_info = """
MATRİS FAKTÖRİZASYON UYGULAMASI - AI ASİSTANI

Bu uygulama matris faktörizasyon algoritmaları için kapsamlı bir araçtır. Aşağıdaki özelliklere sahiptir:

📊 KLASİK ALGORİTMALAR:
1. SVD (Singular Value Decomposition)
   - Kullanım: Öneri sistemleri, gürültü temizleme
   - Sınıf: SVDRecommender, SVDNoiseReducer
   - En iyi: Matematiksel olarak en kesin yöntem

2. PCA (Principal Component Analysis)
   - Kullanım: Veri görselleştirme, özellik seçimi
   - Sınıf: PCAAnalyzer
   - En iyi: Boyut indirgeme için

3. NMF (Non-negative Matrix Factorization)
   - Kullanım: Görüntü işleme, topic modeling
   - Sınıf: NMFImageProcessor, NMFTopicModeler
   - En iyi: Pozitif değerlerle çalışan veriler için

4. ALS (Alternating Least Squares)
   - Kullanım: Büyük ölçekli öneri motorları
   - Sınıf: ALSRecommender
   - En iyi: Paralel çalışma gerektiren sistemler

🚀 MODERN DEEP LEARNING ALGORİTMALAR (PyTorch gerekli):
5. NCF (Neural Collaborative Filtering)
   - Kullanım: Netflix, YouTube tarzı öneriler
   - Sınıf: NCFRecommender
   - Özellik: Doğrusal olmayan ilişkileri öğrenir

6. Autoencoder (Denoising & VAE)
   - Kullanım: Gürültü temizleme, öneriler
   - Sınıf: DenoisingAutoencoder, VAERecommender
   - Özellik: SVD ve PCA'in Deep Learning karşılığı

7. Factorization Machines (FM) & DeepFM
   - Kullanım: Context-aware öneriler, CTR tahmini
   - Sınıf: FactorizationMachine, DeepFM
   - Özellik: Yan bilgileri (saat, cihaz, vb.) kullanır

8. Transformer (BERT4Rec/SASRec)
   - Kullanım: TikTok, YouTube tarzı sequential öneriler
   - Sınıf: TransformerRecommender
   - Özellik: Zaman ve sıra bilgisini kullanır

9. GNN (Graph Neural Network)
   - Kullanım: Pinterest, Uber Eats tarzı graph tabanlı öneriler
   - Sınıf: GNNRecommender
   - Özellik: Veriyi graph olarak modelleyerek öneriler üretir

KULLANILABİLİR FONKSİYONLAR:
- SVD analizi: run_svd_analysis(n_components=5)
- Veri yükleme: generate_sample_data(), load_rating_data_from_file()
- Görselleştirme: plot_ratings_matrix(), plot_recommendations()

GÖREVİN:
- Kullanıcıların veri setleri hakkındaki sorularını yanıtla
- Hangi algoritmanın ne zaman kullanılacağını öner
- Matris faktörizasyon teknikleri hakkında detaylı bilgi ver
- Veri analizi için Python pandas kodları öner
- Kullanıcı "SVD çalıştır" veya "matrix factorization yap" derse, uygun algoritmayı öner
"""
                        
                        df_info = f"""
{app_info}

VERİ SETİ BİLGİLERİ:
- Satır sayısı: {len(df_for_analysis)}
- Sütun sayısı: {len(df_for_analysis.columns)}
- Sütunlar: {', '.join(df_for_analysis.columns.tolist())}

VERİ ÖNİZLEME (ilk 100 satır):
{df_preview}

KULLANICI SORUSU: {user_query}

Lütfen bu veri seti hakkında kullanıcının sorusunu yanıtla. Eğer kullanıcı matris faktörizasyon algoritmaları hakkında soru sorarsa, yukarıdaki bilgileri kullanarak detaylı açıklama yap. Eğer grafik çizilmesi gerekiyorsa, Python kodu öner ama grafik çizme.
"""
                        
                        # API çağrısı (OpenAI veya Gemini)
                        if provider == "OpenAI":
                            # OpenAI API çağrısı
                            response_obj = client.chat.completions.create(
                                model=model_name,
                                messages=[
                                    {"role": "system", "content": """Sen bir Matris Faktörizasyon Uygulaması'nın AI asistanısın. Bu uygulama matris faktörizasyon algoritmaları için kapsamlı bir araçtır.

UYGULAMA ÖZELLİKLERİ:
📊 KLASİK ALGORİTMALAR: SVD (öneri/gürültü temizleme), PCA (görselleştirme), NMF (görüntü/topic modeling), ALS (büyük ölçekli)
🚀 MODERN ALGORİTMALAR: NCF, Autoencoder, Factorization Machines, Transformer, GNN (PyTorch gerekli)

GÖREVİN:
- Veri setleri hakkındaki soruları yanıtla
- Hangi algoritmanın ne zaman kullanılacağını öner
- Matris faktörizasyon teknikleri hakkında bilgi ver
- Python pandas kodları öner
- Kullanıcı "SVD çalıştır" veya "matrix factorization yap" derse, uygun algoritmayı öner"""},
                                    {"role": "user", "content": df_info}
                                ],
                                temperature=0.7,
                                max_tokens=1000
                            )
                            response = response_obj.choices[0].message.content
                        else:  # Gemini
                            # Gemini API çağrısı
                            try:
                                # Model adını kullan (GenerativeModel otomatik olarak normalize eder)
                                model = client.GenerativeModel(model_name)
                                # Prompt'ta emoji kullanmamaya dikkat et (encoding sorunları için)
                                prompt = f"""Sen bir veri analizi asistanısın. Kullanıcıların veri setleri hakkındaki sorularını yanıtla. Python pandas kodları önerebilirsin.

{df_info}"""
                                
                                response_obj = model.generate_content(prompt)
                                # Response'u güvenli şekilde al
                                try:
                                    response = response_obj.text
                                except Exception as text_err:
                                    # Eğer text alınamazsa, parts'tan al
                                    if hasattr(response_obj, 'parts') and response_obj.parts:
                                        response = ''.join([part.text for part in response_obj.parts if hasattr(part, 'text')])
                                    else:
                                        response = str(response_obj)
                            except Exception as gemini_err:
                                # Eğer model bulunamazsa, alternatif model dene
                                if "404" in str(gemini_err) or "not found" in str(gemini_err).lower():
                                    # Alternatif modelleri dene - güncel model adları
                                    alternative_models = [
                                        "gemini-1.5-flash",
                                        "gemini-1.5-pro", 
                                        "gemini-pro",
                                        "gemini-1.0-pro",
                                        "models/gemini-1.5-flash",
                                        "models/gemini-1.5-pro"
                                    ]
                                    response = None
                                    for alt_model in alternative_models:
                                        try:
                                            # Her model adını hem prefix'li hem de prefix'siz dene
                                            model_variants = [alt_model]
                                            if not alt_model.startswith("models/"):
                                                model_variants.append(f"models/{alt_model}")
                                            else:
                                                model_variants.append(alt_model.replace("models/", ""))
                                            
                                            for try_model in model_variants:
                                                try:
                                                    model = client.GenerativeModel(try_model)
                                                    prompt = f"""Sen bir Matris Faktörizasyon Uygulaması'nın AI asistanısın. Bu uygulama matris faktörizasyon algoritmaları için kapsamlı bir araçtır.

UYGULAMA ÖZELLİKLERİ:
📊 KLASİK ALGORİTMALAR: SVD (öneri/gürültü temizleme), PCA (görselleştirme), NMF (görüntü/topic modeling), ALS (büyük ölçekli)
🚀 MODERN ALGORİTMALAR: NCF, Autoencoder, Factorization Machines, Transformer, GNN (PyTorch gerekli)

GÖREVİN:
- Veri setleri hakkındaki soruları yanıtla
- Hangi algoritmanın ne zaman kullanılacağını öner
- Matris faktörizasyon teknikleri hakkında bilgi ver
- Python pandas kodları öner
- Kullanıcı "SVD çalıştır" veya "matrix factorization yap" derse, uygun algoritmayı öner

{df_info}"""
                                                    response_obj = model.generate_content(prompt)
                                                    # Response'u güvenli şekilde al
                                                    try:
                                                        response = response_obj.text
                                                    except Exception:
                                                        if hasattr(response_obj, 'parts') and response_obj.parts:
                                                            response = ''.join([part.text for part in response_obj.parts if hasattr(part, 'text')])
                                                        else:
                                                            response = str(response_obj)
                                                    # Model'i güncelle
                                                    st.session_state.pandasai_agent["model"] = try_model
                                                    break
                                                except:
                                                    continue
                                            
                                            if response:
                                                break
                                        except:
                                            continue
                                    
                                    if response is None:
                                        response = f"Hata: Gemini API - Model bulunamadi. Lutfen API key'inizi kontrol edin.\n\nDetay: {str(gemini_err)}"
                                else:
                                    # Hata mesajını güvenli şekilde encode et
                                    try:
                                        error_msg = str(gemini_err).encode('utf-8', errors='replace').decode('utf-8')
                                    except:
                                        error_msg = str(gemini_err)
                                    response = f"Hata: Gemini API - {error_msg}"
                        
                        # Eğer kullanıcı SVD istiyorsa
                        if "svd" in user_query.lower() or "matrix factorization" in user_query.lower():
                            try:
                                from algorithms.svd import SVDRecommender
                                import numpy as np
                                
                                numeric_df = df_for_analysis.select_dtypes(include=[np.number])
                                if not numeric_df.empty:
                                    rating_matrix = numeric_df.fillna(0).values
                                    svd_model = SVDRecommender(n_components=5)
                                    svd_model.fit(rating_matrix)
                                    explained_variance = svd_model.explained_variance_ratio_
                                    total_variance = explained_variance.sum()
                                    
                                    svd_result = f"""

**SVD Analizi Sonuçları:**
- Kullanılan Bileşen Sayısı: 5
- Toplam Açıklanan Varyans: {total_variance:.2%}
- Bileşen Bazında Varyans:
"""
                                    for i, var in enumerate(explained_variance):
                                        svd_result += f"  - Bileşen {i+1}: {var:.2%}\n"
                                    
                                    response += svd_result
                            except Exception as svd_err:
                                response += f"\n\n⚠️ SVD analizi sırasında hata: {str(svd_err)}"
                    else:
                        # PandasAI modu
                        response = st.session_state.pandasai_agent.chat(user_query)
                    
                    # Yanıtı göster
                    st.markdown(response)
                    
                    # Mesajı geçmişe ekle
                    message_to_save = {
                        "role": "assistant",
                        "content": response
                    }
                    
                    # Eğer grafik oluşturulduysa (matplotlib figure)
                    # PandasAI genellikle grafikleri otomatik olarak gösterir
                    # Ancak manuel kontrol için:
                    try:
                        import matplotlib.pyplot as plt
                        if plt.get_fignums():
                            fig = plt.gcf()
                            st.pyplot(fig)
                            message_to_save["figure"] = fig
                            plt.close(fig)
                    except:
                        pass
                    
                    st.session_state.chat_messages.append(message_to_save)
                    
                except Exception as e:
                    error_msg = f"❌ Hata: {str(e)}"
                    st.error(error_msg)
                    st.session_state.chat_messages.append({
                        "role": "assistant",
                        "content": error_msg
                    })
    
    # Yardımcı örnek sorular
    with st.expander("💡 Örnek Sorular", expanded=False):
        st.markdown("""
        **Temel Analiz:**
        - "Veri setinde kaç satır ve sütun var?"
        - "Eksik veriler var mı?"
        - "Sütunların istatistiklerini göster"
        
        **Grafikler:**
        - "Rating dağılımını histogram olarak çiz"
        - "Kullanıcı sayılarını bar chart olarak göster"
        - "Korelasyon matrisini göster"
        
        **Özel Komutlar:**
        - "SVD çalıştır" veya "SVD analizi yap"
        - "Matrix factorization yap"
        
        **Veri İşleme:**
        - "Boş verileri temizle"
        - "En yüksek rating'e sahip 10 kullanıcıyı göster"
        - "Ortalama rating'i hesapla"
        """)


if __name__ == "__main__":
    main()

